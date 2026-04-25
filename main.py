"""
👖 BLUE JEANS Creator Engine — Main Application

버전 정보는 prompt.py의 ENGINE_VERSION / ENGINE_BUILD_DATE를 참조한다.
(단일 소스 원칙 — 버전은 한 곳에서만 관리)

아이디어 → 기획개발 패키지
단일 페이지 · 사이드바 최소 · 2단계 Brainstorm
"""

import streamlit as st
import json
import re
from datetime import datetime
import prompt as P

# ─── 버전 정보 (prompt.py에서 단일 소스로 참조) ───
ENGINE_VERSION = P.ENGINE_VERSION
ENGINE_BUILD_DATE = P.ENGINE_BUILD_DATE
ENGINE_FOOTER = f"© 2026 BLUE JEANS PICTURES · Creator Engine {ENGINE_VERSION}"

ANTHROPIC_MODEL = "claude-sonnet-4-6"           # 구조 작업 — 비용 효율
ANTHROPIC_MODEL_OPUS = "claude-opus-4-6"        # 캐릭터 바이블 · 트리트먼트 · 톤 문서 — 최고 품질


# ─── v2.4.0: 시대극 자동 감지 컨텍스트 추출 ─────────────────────────
def _get_period_scan_context(project: dict = None) -> tuple:
    """현재 세션 상태와 프로젝트 데이터에서 시대 감지용 텍스트를 추출.
    
    Idea Engine v1.0 시드(st.session_state["locked_seed"])가 있으면 활용.
    - locked_logline, locked_references, locked_genre, locked_theme 등이
      시대(SAGEUK/COLONIAL/MODERN_HIST) 자동 감지의 결정적 시그널이 됨.
    
    Returns:
        (locked_text, idea_text) 튜플.
        둘 다 빈 문자열이면 시대 OVERRIDE·Period Pack 자동 감지가 작동하지 않으며
        v2.3.10과 동일하게 동작 (하위 호환).
    """
    try:
        locked_seed = st.session_state.get("locked_seed", {}) or {}
    except Exception:
        locked_seed = {}
    return P.build_scan_text_for_period_detection(
        project=project or {},
        locked_seed=locked_seed,
    )

# ─── Page Config ───
st.set_page_config(
    page_title="BLUE JEANS · Creator Engine",
    page_icon="👖",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# ─── Sidebar: Engine Info (버전 확인용) ───
with st.sidebar:
    st.markdown(f"""
    <div style="padding:12px;background:#F0F2FF;border-radius:8px;border-left:3px solid #191970;font-family:'Pretendard',sans-serif;">
        <div style="font-size:.72rem;color:#191970;font-weight:700;letter-spacing:.05em;margin-bottom:4px;">ENGINE INFO</div>
        <div style="font-size:1.05rem;font-weight:700;color:#191970;">Creator Engine</div>
        <div style="font-size:1.25rem;font-weight:900;color:#FFCB05;background:#191970;padding:2px 8px;border-radius:4px;display:inline-block;margin-top:4px;">
            {ENGINE_VERSION}
        </div>
        <div style="font-size:.7rem;color:#666;margin-top:8px;">
            Build: {ENGINE_BUILD_DATE}<br>
            Status: {P.ENGINE_STATUS}
        </div>
    </div>
    """, unsafe_allow_html=True)
    st.caption("버전이 최신인지 확인하세요.")

# ─── Custom CSS ───
st.markdown("""
<style>
@import url('https://cdn.jsdelivr.net/gh/orioncactus/pretendard/dist/web/static/pretendard.css');
@import url('https://cdn.jsdelivr.net/gh/projectnoonnu/2408-3@latest/Paperlogy.css');
@import url('https://fonts.googleapis.com/css2?family=Playfair+Display:wght@700;900&display=swap');

:root {
    --navy: #191970;
    --y: #FFCB05;
    --bg: #F7F7F5;
    --card: #FFFFFF;
    --card-border: #E2E2E0;
    --t: #1A1A2E;
    --r: #D32F2F;
    --g: #2EC484;
    --dim: #8E8E99;
    --light-bg: #EEEEF6;
    --serif: 'Paperlogy', 'Noto Serif KR', 'Georgia', serif;
    --display: 'Playfair Display', 'Paperlogy', 'Georgia', serif;
    --body: 'Pretendard', -apple-system, sans-serif;
    --heading: 'Paperlogy', 'Pretendard', sans-serif;
}

/* ── 기본 타이포 ── */
html, body, [class*="css"] {
    font-family: var(--body);
    color: var(--t);
    -webkit-font-smoothing: antialiased;
}

/* ══ 라이트모드 강제 (전역) ══ */
.stApp, [data-testid="stAppViewContainer"], [data-testid="stMain"],
[data-testid="stMainBlockContainer"], [data-testid="stHeader"],
[data-testid="stBottom"] {
    background-color: var(--bg) !important;
    color: var(--t) !important;
}
.stMarkdown, .stText, .stCode {
    color: var(--t) !important;
}
h1, h2, h3, h4, h5, h6 { color: var(--navy) !important; font-family: var(--heading) !important; }
p, span, label, div, li { color: inherit; }

/* ── 사이드바 숨김 ── */
section[data-testid="stSidebar"] { display: none; }

/* ══ 입력 위젯 라이트 ══ */
.stTextInput input, .stTextArea textarea,
[data-testid="stTextInput"] input, [data-testid="stTextArea"] textarea {
    background-color: var(--card) !important;
    color: var(--t) !important;
    border: 1.5px solid var(--card-border) !important;
    border-radius: 8px !important;
    font-family: var(--body) !important;
    font-size: 0.9rem !important;
    padding: 0.6rem 0.8rem !important;
    transition: border-color 0.2s;
}
.stTextInput input:focus, .stTextArea textarea:focus,
[data-testid="stTextInput"] input:focus, [data-testid="stTextArea"] textarea:focus {
    border-color: var(--navy) !important;
    box-shadow: 0 0 0 2px rgba(25,25,112,0.08) !important;
}
.stTextInput input::placeholder, .stTextArea textarea::placeholder,
[data-testid="stTextInput"] input::placeholder, [data-testid="stTextArea"] textarea::placeholder {
    color: var(--dim) !important;
    font-size: 0.85rem !important;
}
/* selectbox */
.stSelectbox > div > div, [data-baseweb="select"] > div, [data-baseweb="select"] input {
    background-color: var(--card) !important;
    color: var(--t) !important;
    border-color: var(--card-border) !important;
    border-radius: 8px !important;
}
[data-baseweb="popover"], [data-baseweb="menu"], [role="listbox"], [role="option"] {
    background-color: var(--card) !important;
    color: var(--t) !important;
}
[role="option"]:hover { background-color: var(--light-bg) !important; }
/* label */
.stTextInput label, .stTextArea label, .stSelectbox label, .stRadio label {
    color: var(--t) !important;
    font-weight: 600 !important;
    font-size: 0.82rem !important;
    margin-bottom: 0.3rem !important;
}

/* ══ 버튼 ══ */
.stButton > button {
    color: var(--t) !important;
    border: 1.5px solid var(--card-border) !important;
    background-color: var(--card) !important;
    border-radius: 8px !important;
    font-family: var(--body) !important;
    font-weight: 600 !important;
    font-size: 0.85rem !important;
    padding: 0.5rem 1.2rem !important;
    transition: all 0.2s;
}
.stButton > button:hover {
    border-color: var(--navy) !important;
    box-shadow: 0 2px 8px rgba(25,25,112,0.08) !important;
}
.stButton > button[kind="primary"],
.stButton > button[data-testid="stBaseButton-primary"] {
    background-color: var(--y) !important;
    color: var(--navy) !important;
    border-color: var(--y) !important;
    font-weight: 700 !important;
}
.stButton > button[kind="primary"]:hover,
.stButton > button[data-testid="stBaseButton-primary"]:hover {
    background-color: #E8B800 !important;
    box-shadow: 0 2px 12px rgba(255,203,5,0.3) !important;
}

/* ══ Expander ══ */
.stExpander, details, details summary {
    background-color: var(--card) !important;
    color: var(--t) !important;
    border: 1px solid var(--card-border) !important;
    border-radius: 8px !important;
}
details[open] > div { background-color: var(--card) !important; }
.stExpander summary, .stExpander summary span { color: var(--t) !important; }

/* ══ Alert 박스 ══ */
.stAlert { color: var(--t) !important; border-radius: 8px !important; }

/* ══ 내부 컨테이너 투명 ══ */
[data-testid="stVerticalBlock"], [data-testid="stHorizontalBlock"],
[data-testid="stColumn"] { background-color: transparent !important; }

/* ══ Metric ══ */
[data-testid="stMetric"] { background-color: var(--card) !important; color: var(--t) !important; }
[data-testid="stMetric"] label { color: var(--dim) !important; }
.stCaption, small { color: var(--dim) !important; }
.stCheckbox label span, .stToggle label span { color: var(--t) !important; }

/* ══ 배경색별 텍스트 강제 ══ */
[style*="background:#FFCB05"] { color: var(--navy) !important; }
[style*="background:#FFCB05"] * { color: var(--navy) !important; }
[style*="background:#2EC484"] { color: #FFFFFF !important; }
[style*="background:#2EC484"] * { color: #FFFFFF !important; }

/* ═══════════════════════════════════
   브랜딩 & 커스텀 컴포넌트
   ═══════════════════════════════════ */

.header {
    font-size: 0.85rem;
    font-weight: 700;
    color: var(--navy);
    letter-spacing: 0.15em;
    margin-bottom: 0;
    font-family: var(--heading);
}

.brand-title {
    font-size: 2.6rem;
    font-weight: 900;
    color: var(--navy);
    font-family: var(--display);
    letter-spacing: -0.02em;
    margin-bottom: 0.15rem;
    position: relative;
    display: inline-block;
}
.brand-title::after {
    content: '';
    position: absolute;
    bottom: 2px;
    left: 0;
    width: 100%;
    height: 4px;
    background: var(--y);
    border-radius: 2px;
}

.sub {
    font-size: 0.7rem;
    color: var(--dim);
    letter-spacing: 0.15em;
    margin-top: 0.5rem;
    margin-bottom: 1.5rem;
}

/* ── 카드 ── */
.card {
    background: var(--card);
    border: 1px solid var(--card-border);
    border-radius: 10px;
    padding: 1.2rem;
    margin-bottom: 0.8rem;
    box-shadow: 0 1px 4px rgba(0,0,0,0.03);
    transition: all 0.2s;
}
.card:hover {
    border-color: var(--navy);
    box-shadow: 0 3px 12px rgba(25,25,112,0.07);
    transform: translateY(-1px);
}

/* ── 콜아웃 ── */
.callout {
    background: var(--light-bg);
    border-left: 4px solid var(--navy);
    padding: 0.9rem 1.1rem;
    margin: 0.5rem 0;
    border-radius: 0 8px 8px 0;
    font-size: 0.88rem;
    color: var(--t);
}

/* ── 섹션 라벨 ── */
.cl {
    color: var(--navy);
    font-weight: 700;
    font-size: 0.72rem;
    letter-spacing: 0.03em;
    margin-bottom: 0.3rem;
    text-transform: uppercase;
}

/* ── 정보 블록 ── */
.ri {
    background: var(--light-bg);
    border-radius: 8px;
    padding: 0.9rem 1rem;
    margin-bottom: 0.5rem;
    font-size: 0.88rem;
    line-height: 1.6;
}
.rl {
    color: var(--navy);
    font-weight: 700;
    font-size: 0.72rem;
    letter-spacing: 0.02em;
}

/* ── 큰 숫자 ── */
.big {
    font-size: 2.8rem;
    font-weight: 900;
    color: var(--navy);
    text-align: center;
    font-family: var(--heading);
}
.sm {
    font-size: 0.7rem;
    color: var(--dim);
    text-align: center;
}

/* ── 뱃지 ── */
.badge {
    display: inline-block;
    padding: 0.2rem 0.6rem;
    border-radius: 5px;
    font-size: 0.7rem;
    font-weight: 600;
}
.b-done { background: var(--g); color: #fff; }
.b-run  { background: var(--y); color: var(--navy); }
.b-not  { background: #E8E8F0; color: var(--dim); }
.b-fail { background: var(--r); color: #fff; }

/* ── 노란 섹션 헤더 (웹 UI) ── */
.section-header {
    background: var(--y);
    color: var(--navy);
    padding: 0.6rem 1rem;
    border-radius: 6px;
    font-weight: 800;
    font-size: 1rem;
    font-family: var(--heading);
    margin: 1.5rem 0 0.8rem 0;
    display: flex;
    justify-content: space-between;
    align-items: center;
}
.section-header .en {
    font-family: var(--display);
    font-size: 0.75rem;
    font-weight: 700;
    letter-spacing: 0.05em;
    opacity: 0.7;
}

/* ══ Stepper ══ */
.stepper {
    display: flex;
    align-items: center;
    justify-content: center;
    margin: 1.5rem 0 2rem 0;
    gap: 0;
}
.step-wrap {
    display: flex;
    flex-direction: column;
    align-items: center;
}
.step-circle {
    width: 34px;
    height: 34px;
    border-radius: 50%;
    display: flex;
    align-items: center;
    justify-content: center;
    font-size: 0.78rem;
    font-weight: 700;
    flex-shrink: 0;
    border: 2px solid transparent;
    transition: all 0.2s;
}
.step-circle.active {
    background: var(--y);
    color: var(--navy);
    border-color: var(--y);
    box-shadow: 0 0 0 4px rgba(255,203,5,0.2);
}
.step-circle.done {
    background: var(--g);
    color: #fff;
    border-color: var(--g);
}
.step-circle.upcoming {
    background: #EDEDF0;
    color: var(--dim);
    border-color: #D8D8E0;
}
.step-label {
    font-size: 0.6rem;
    margin-top: 0.35rem;
    text-align: center;
    width: 55px;
    font-weight: 500;
}
.step-label.active { color: var(--navy); font-weight: 700; }
.step-label.done { color: var(--g); font-weight: 600; }
.step-label.upcoming { color: var(--dim); }
.step-line {
    width: 30px;
    height: 2px;
    margin: 0 2px;
    flex-shrink: 0;
    border-radius: 1px;
}
.step-line.done { background: var(--g); }
.step-line.upcoming { background: #D8D8E0; }

/* ══ 프로그레스 바 (공통) ══ */
.progress-track {
    flex: 1;
    background: #E8E8F0;
    border-radius: 4px;
    height: 8px;
    margin: 0 0.5rem;
}
.progress-fill {
    height: 100%;
    background: var(--y);
    border-radius: 4px;
    transition: width 0.3s;
}
</style>
""", unsafe_allow_html=True)


# ─── Session State ───
defaults = {
    "view": "home",         # home | project | core | char_bible | structure | scene_design | treatment | tone_doc | export
    "projects": {},
    "cur": None,
    "last_research_raw": "",
    "last_cards_raw": "",
    "last_analysis_raw": "",
    "last_core_raw": "",
    "last_gate_raw": "",
    "last_char_bible_raw": "",
    "last_structure_story_raw": "",
    "last_structure_diag_raw": "",
    "last_structure_gate_raw": "",
    "last_scene_design_raw": "",
    "last_treatment_act1_raw": "",
    "last_treatment_act2_raw": "",
    "last_treatment_act3_raw": "",
    "last_tone_doc_raw": "",
    "last_error": "",
}
for k, v in defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v


# ─── Stepper (9단계) ───
STEPS = [
    ("project", "아이디어"),
    ("project", "Brainstorm"),
    ("core", "Core"),
    ("char_bible", "Bible"),
    ("structure", "Structure"),
    ("scene_design", "Scene"),
    ("treatment", "Treatment"),
    ("tone_doc", "Tone"),
    ("export", "Export"),
]

def render_stepper(current_view, project_data=None):
    """상단 단계 네비게이션 바 — 완료된 단계는 클릭 이동 가능 (v2.3.4)"""
    view_to_step = {
        "project": 1 if project_data and project_data.get("brainstorm_cards") else 0,
        "core": 2,
        "char_bible": 3,
        "structure": 4,
        "scene_design": 5,
        "treatment": 6,
        "tone_doc": 7,
        "export": 8,
    }
    current_idx = view_to_step.get(current_view, 0)

    done_idx = -1
    if project_data:
        if project_data.get("brainstorm_cards"):
            done_idx = 1
        if project_data.get("brainstorm_analysis"):
            ga = project_data["brainstorm_analysis"].get("gate_a_scores", {})
            if ga.get("average", 0) >= 7.0:
                done_idx = 1
        if project_data.get("core"):
            done_idx = 2
        if project_data.get("char_bible"):
            done_idx = 3
        if project_data.get("structure_story"):
            done_idx = 4
        if project_data.get("scene_design"):
            done_idx = 5
        if project_data.get("treatment"):
            done_idx = 6
        if project_data.get("tone_doc"):
            done_idx = 7

    # 인라인 CSS — 버튼을 stepper 형태로 보이게
    st.markdown("""
    <style>
    .stepper-nav [data-testid="column"] {
        padding: 0 !important;
    }
    .stepper-nav .stButton button {
        width: 100%;
        min-height: 3.4rem;
        padding: 0.4rem 0.2rem;
        font-size: 0.72rem;
        font-weight: 600;
        line-height: 1.15;
        border-radius: 8px;
        white-space: normal;
        word-break: keep-all;
    }
    .step-done button {
        background: #191970 !important;
        color: #FFCB05 !important;
        border: 1px solid #191970 !important;
    }
    .step-done button:hover {
        background: #FFCB05 !important;
        color: #191970 !important;
        border-color: #FFCB05 !important;
    }
    .step-active button {
        background: #FFCB05 !important;
        color: #191970 !important;
        border: 2px solid #191970 !important;
        cursor: default !important;
    }
    .step-upcoming button {
        background: #F7F7F5 !important;
        color: #B0B0B0 !important;
        border: 1px dashed #D0D0D0 !important;
        cursor: not-allowed !important;
    }
    </style>
    """, unsafe_allow_html=True)

    # 단계 버튼을 가로로 배치
    st.markdown('<div class="stepper-nav">', unsafe_allow_html=True)
    cols = st.columns(len(STEPS))

    for i, (view_key, label) in enumerate(STEPS):
        # 상태 결정
        if i < current_idx and i <= done_idx:
            state = "done"
        elif i == current_idx:
            state = "active"
        else:
            state = "upcoming"

        icon = "✓" if state == "done" else ("●" if state == "active" else str(i + 1))
        btn_label = f"{icon}\n{label}"

        with cols[i]:
            # 상태별 컨테이너 스타일 적용
            st.markdown(f'<div class="step-{state}">', unsafe_allow_html=True)

            if state == "done":
                # 완료된 단계: 클릭 시 해당 view로 점프
                if st.button(btn_label, key=f"stepper_nav_{view_key}_{i}", use_container_width=True):
                    st.session_state.view = view_key
                    st.rerun()
            elif state == "active":
                # 현재 위치: 버튼 형태로 표시되지만 클릭해도 의미 없음 (disabled)
                st.button(btn_label, key=f"stepper_nav_{view_key}_{i}", use_container_width=True, disabled=True)
            else:
                # 미완료: 비활성
                st.button(btn_label, key=f"stepper_nav_{view_key}_{i}", use_container_width=True, disabled=True)

            st.markdown('</div>', unsafe_allow_html=True)

    st.markdown('</div>', unsafe_allow_html=True)


# ─── JSON Helpers ───
# ═══════════════════════════════════════════════════
# Project JSON Save/Load (v2.3.5 신규)
# 목적: 세션 끊김 대비한 프로젝트 상태 복구용
# 원칙: 전체 저장 / 전체 복원 — 부분 저장 금지
# ═══════════════════════════════════════════════════

def _get_project_stage(project: dict) -> str:
    """프로젝트의 현재 완료 단계 파악 — 파일명 생성용"""
    if project.get("tone_doc"):
        return "Tone완료"
    if project.get("treatment"):
        return "Treatment완료"
    if project.get("scene_design"):
        return "Scene완료"
    if project.get("structure_story"):
        return "Structure완료"
    if project.get("char_bible"):
        return "Character완료"
    if project.get("core"):
        return "Core완료"
    if project.get("brainstorm_cards"):
        return "Brainstorm완료"
    return "초기"


def save_project_to_json(project: dict, engine_version: str = "v2.3.5") -> str:
    """프로젝트 전체를 JSON 문자열로 직렬화.
    
    Args:
        project: Streamlit session_state의 project dict
        engine_version: 현재 엔진 버전 (호환성 검증용)
    
    Returns:
        UTF-8 JSON 문자열
    """
    # 저장 가능한 형태로 복사
    payload = {
        "_meta": {
            "blue_jeans_engine": "Creator Engine",
            "engine_version": engine_version,
            "saved_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "stage": _get_project_stage(project),
        },
        "project": project,
    }
    return json.dumps(payload, ensure_ascii=False, indent=2)


def load_project_from_json(json_str: str) -> dict:
    """JSON 문자열에서 프로젝트 복원.
    
    Args:
        json_str: save_project_to_json 으로 저장된 JSON 문자열
    
    Returns:
        {"project": {...}, "meta": {...}, "warnings": [...]} 
    
    Raises:
        ValueError: 형식이 올바르지 않을 때
    """
    try:
        payload = json.loads(json_str)
    except json.JSONDecodeError as e:
        raise ValueError(f"유효한 JSON 파일이 아닙니다: {e}")
    
    # 구조 검증
    if not isinstance(payload, dict):
        raise ValueError("JSON 최상위가 객체가 아닙니다.")
    if "project" not in payload:
        raise ValueError("project 필드가 없습니다. Creator Engine 프로젝트 파일이 맞는지 확인하세요.")
    
    project = payload["project"]
    meta = payload.get("_meta", {})
    warnings = []
    
    # 필수 필드 검증
    required_fields = ["project_id", "title", "idea_text", "genre", "format"]
    missing = [f for f in required_fields if f not in project]
    if missing:
        raise ValueError(f"필수 필드 누락: {', '.join(missing)}")
    
    # 엔진 버전 호환성 체크 (경고만, 실패 아님)
    saved_version = meta.get("engine_version", "unknown")
    current_version = "v2.3.5"
    if saved_version != current_version:
        warnings.append(
            f"엔진 버전 차이 — 저장 시: {saved_version} / 현재: {current_version}. "
            f"대부분 호환되지만, 이후 단계 실행 시 예상치 못한 동작이 있을 수 있습니다."
        )
    
    return {
        "project": project,
        "meta": meta,
        "warnings": warnings,
    }


def extract_json_object(text: str) -> str:
    """응답 텍스트에서 JSON 객체만 추출"""
    text = text.strip()

    # 코드블록 제거
    if text.startswith("```json"):
        text = text[7:]
    elif text.startswith("```"):
        text = text[3:]
    if text.endswith("```"):
        text = text[:-3]

    text = text.strip()

    # 첫 { 부터 마지막 } 까지 추출
    start = text.find("{")
    end = text.rfind("}")

    if start == -1 or end == -1:
        raise ValueError("JSON 객체 시작/끝을 찾지 못했습니다.")

    return text[start:end + 1]


def safe_json_loads(text: str):
    """JSON 파싱 (강화된 복구 로직 — 5단계)"""
    cleaned = extract_json_object(text)
    cleaned = re.sub(r",\s*([}\]])", r"\1", cleaned)

    # 0단계: 문자열 값 내부의 raw 제어 문자 제거 (Invalid control character 방지)
    # JSON 문자열 안에 raw \n, \r, \t가 있으면 파싱 실패함
    def _sanitize_control_chars(s):
        """JSON 문자열 값 내부의 제어 문자를 공백으로 치환"""
        result = []
        in_string = False
        i = 0
        while i < len(s):
            c = s[i]
            if c == '"' and (i == 0 or s[i-1] != '\\'):
                in_string = not in_string
                result.append(c)
            elif in_string and c in '\n\r':
                result.append(' ')
            elif in_string and c == '\t':
                result.append(' ')
            elif in_string and ord(c) < 32 and c not in ('\n', '\r', '\t'):
                result.append(' ')  # 기타 제어 문자
            else:
                result.append(c)
            i += 1
        return ''.join(result)

    cleaned = _sanitize_control_chars(cleaned)

    # 1차: 그대로 파싱
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        pass

    # 2차: 문자열 내부 줄바꿈 이스케이프 + 쌍따옴표 → 작은따옴표
    try:
        result = []
        i = 0
        n = len(cleaned)
        while i < n:
            ch = cleaned[i]
            if ch == '"':
                result.append(ch)
                i += 1
                str_chars = []
                while i < n:
                    c = cleaned[i]
                    if c == '\\' and i + 1 < n:
                        next_c = cleaned[i + 1]
                        if next_c == '"':
                            str_chars.append("'")
                            i += 2
                        elif next_c in ('n', 'r', 't', '\\', '/', 'b', 'f', 'u'):
                            str_chars.append(c)
                            str_chars.append(next_c)
                            i += 2
                        else:
                            str_chars.append(next_c)
                            i += 2
                    elif c == '"':
                        # 문자열 종료인지 판단 — 다음 비공백 문자 확인
                        rest = cleaned[i + 1:] if i + 1 < n else ""
                        rest_stripped = rest.lstrip()
                        if not rest_stripped:
                            break  # EOF
                        nc = rest_stripped[0]
                        if nc == ':':
                            break  # key-value 구분자 → 문자열 종료
                        elif nc == ',':
                            # 쉼표 뒤를 확인: " 또는 } 또는 ] → JSON 구조
                            after_comma = rest_stripped[1:].lstrip()
                            if not after_comma or after_comma[0] in '"{}[]0123456789tfn':
                                break  # 문자열 종료
                            else:
                                str_chars.append("'")
                                i += 1
                        elif nc in '}]':
                            break  # 객체/배열 종료
                        elif nc == '"':
                            # 다음이 곧바로 " → 이건 빈 문자열이 아니라 키 시작
                            break
                        else:
                            str_chars.append("'")
                            i += 1
                    elif c in '\n\r':
                        str_chars.append(' ')
                        i += 1
                    elif c == '\t':
                        str_chars.append(' ')
                        i += 1
                    else:
                        str_chars.append(c)
                        i += 1
                result.append(''.join(str_chars))
                result.append('"')
                i += 1
            else:
                result.append(ch)
                i += 1

        fixed = ''.join(result)
        fixed = re.sub(r",\s*([}\]])", r"\1", fixed)
        return json.loads(fixed)
    except (json.JSONDecodeError, IndexError):
        pass

    # 3차: 에러 위치 기반 반복 수정 (최대 30회)
    current = cleaned
    for _ in range(30):
        try:
            return json.loads(current)
        except json.JSONDecodeError as e:
            pos = e.pos
            if pos is None or pos <= 0:
                break
            # 에러 위치 주변에서 문제가 되는 " 찾아서 ' 로 교체
            found = False
            for j in range(pos, max(0, pos - 100), -1):
                if j < len(current) and current[j] == '"':
                    before = current[j - 1] if j > 0 else ''
                    after_c = current[j + 1] if j + 1 < len(current) else ''
                    # 구조적 위치가 아닌 " → 내부 따옴표로 판단
                    if before not in ('{', '[', ',', ':', '\\') and after_c not in (':', ',', '}', ']'):
                        current = current[:j] + "'" + current[j + 1:]
                        found = True
                        break
                    elif before == '\\':
                        # \" → ' (이스케이프 제거)
                        current = current[:j - 1] + "'" + current[j + 1:]
                        found = True
                        break
            if not found:
                # 에러 위치 근처의 줄바꿈을 공백으로
                for j in range(max(0, pos - 5), min(len(current), pos + 5)):
                    if current[j] in '\n\r\t':
                        current = current[:j] + ' ' + current[j + 1:]
                        found = True
                        break
            if not found:
                break

    # 4차: 최후 수단 — 모든 문자열 값 내부 " → '
    try:
        aggressive = re.sub(
            r'(?<=: ")(.*?)(?="\s*[,}\]])',
            lambda m: m.group(0).replace('"', "'"),
            cleaned,
            flags=re.DOTALL
        )
        aggressive = re.sub(r",\s*([}\]])", r"\1", aggressive)
        return json.loads(aggressive)
    except json.JSONDecodeError as e:
        raise e


# ─── API Client ───
def get_client():
    try:
        import anthropic
    except ImportError:
        raise RuntimeError("anthropic 패키지가 설치되지 않았습니다. requirements.txt를 확인하세요.")

    if "ANTHROPIC_API_KEY" not in st.secrets:
        raise RuntimeError("ANTHROPIC_API_KEY가 secrets에 설정되지 않았습니다.")

    return anthropic.Anthropic(api_key=st.secrets["ANTHROPIC_API_KEY"])


def _build_project_locked_block(project: dict) -> str:
    """프로젝트의 LOCKED 항목으로 프롬프트 주입 블록 생성.
    v2.3: OPEN 필드 폐지. LOCKED 항목이 없으면 빈 문자열 반환."""
    locked = project.get("locked_items", [])
    if not locked:
        return ""
    return P.build_locked_block(locked)


# ─── API Call: Research ───
def call_research(idea, genre, market):
    try:
        client = get_client()

        system_prompt = P.SYSTEM_RESEARCH

        user_prompt = f"""[입력]
아이디어: {idea}
장르: {genre}
타겟: {market}

[JSON 스키마]
{{
  "search_keywords": [],
  "real_events": [
    {{
      "id": 1,
      "title": "",
      "summary": "",
      "source": "",
      "year": "",
      "relevance": "",
      "story_potential": ""
    }}
  ],
  "existing_works": [
    {{
      "id": 1,
      "title": "",
      "type": "",
      "country": "",
      "year": "",
      "summary": "",
      "similarity": "",
      "difference_opportunity": ""
    }}
  ],
  "research_summary": {{
    "total_real_events": 0,
    "total_existing_works": 0,
    "key_insight": ""
  }}
}}

규칙:
- real_events 3~7개
- existing_works 3~7개
- source는 가능한 한 짧게
- key_insight는 2문장 이내
"""

        response = client.messages.create(
            model=ANTHROPIC_MODEL,
            max_tokens=16000,
            temperature=0.2,
            system=system_prompt,
            messages=[{"role": "user", "content": user_prompt}]
        )

        txt = "".join(
            block.text for block in response.content if hasattr(block, "text")
        ).strip()
        st.session_state["last_research_raw"] = txt

        return safe_json_loads(txt)

    except Exception as e:
        st.error(f"리서치 실패: {e}")
        raw = st.session_state.get("last_research_raw")
        if raw:
            with st.expander("🔧 Research Raw Response (디버그)"):
                st.text_area("Raw", raw, height=300)
        return None


# ─── API Call: Brainstorm Cards (1단계) ───
def call_brainstorm_cards(idea, genre, market, fmt, research=None, locked_block=""):
    """1단계: 아이디어 카드 10개 + Top 3 생성 (토큰 집중)"""
    try:
        client = get_client()

        system_prompt = P.SYSTEM_BRAINSTORM_CARDS

        research_block = ""
        if research:
            research_block = f"""
[리서치 참고]
{json.dumps(research, ensure_ascii=False)}

규칙:
- 실화/뉴스는 소재 참고용
- 기존작은 차별화 포인트 참고용
- 동일 구조/설정 반복 금지
"""

        user_prompt = f"""[입력]
아이디어: {idea}
장르: {genre}
타겟: {market}
포맷: {fmt}
{locked_block}
{research_block}

[JSON 스키마]
{{
  "idea_type": "story|mood|hybrid",
  "idea_type_diagnosis": "",
  "idea_type_action": "",
  "idea_cards": [
    {{
      "id": 1,
      "title": "10자 이내",
      "logline_seed": "",
      "protagonist": "",
      "conflict": "",
      "hook": "",
      "visual_image": "",
      "genre": "",
      "scores": {{
        "active_hero": 0.0,
        "conflict_clarity": 0.0,
        "visual_power": 0.0,
        "genre_immediacy": 0.0,
        "originality": 0.0,
        "market_fit": 0.0
      }},
      "total_score": 0.0
    }}
  ],
  "top3": [
    {{
      "rank": 1,
      "card_id": 0,
      "reason": ""
    }}
  ]
}}

규칙:
- idea_cards는 10개 (10개 초과 금지)
- top3는 3개
- total_score는 6개 점수 평균, 소수점 1자리
- scores 각 항목은 0.0~10.0
- idea_type_diagnosis: 이 아이디어의 현재 상태를 이 이야기에만 해당하는 구체적 언어로 1문장. "~이 필요하다" 같은 추상 진단 금지. 예: "주인공의 욕망은 선명하지만 적대자가 없어 갈등이 발생하지 않는다"
- idea_type_action: idea_type_diagnosis에 대한 즉각적 처방 1문장. 반드시 "~하라" 형태로. 예: "주인공이 원하는 것을 빼앗는 구체적 인물 또는 시스템을 설정하라"
- 각 카드의 logline_seed는 1문장, 50자 이내
- protagonist는 1문장, 30자 이내
- conflict는 1문장, 40자 이내
- hook는 1문장, 30자 이내
- visual_image는 1문장, 40자 이내
- reason은 1문장, 30자 이내
- 절대로 장문 서술하지 말 것. 짧고 강하게.
"""

        response = client.messages.create(
            model=ANTHROPIC_MODEL,
            max_tokens=16000,
            temperature=0.35,
            system=system_prompt,
            messages=[{"role": "user", "content": user_prompt}]
        )

        # 토큰 한계 도달 감지
        if response.stop_reason == "max_tokens":
            st.warning("⚠️ 응답이 토큰 한계에서 잘렸습니다. 재시도합니다...")
            # 재시도: max_tokens 더 올리고 카드 수 줄이기
            retry_prompt = user_prompt.replace("idea_cards는 10개", "idea_cards는 7개")
            response = client.messages.create(
                model=ANTHROPIC_MODEL,
                max_tokens=16000,
                temperature=0.35,
                system=system_prompt,
                messages=[{"role": "user", "content": retry_prompt}]
            )

        txt = "".join(
            block.text for block in response.content if hasattr(block, "text")
        ).strip()
        st.session_state["last_cards_raw"] = txt

        return safe_json_loads(txt)

    except Exception as e:
        st.session_state["last_error"] = str(e)
        st.error(f"카드 생성 실패: {e}")
        raw = st.session_state.get("last_cards_raw")
        if raw:
            with st.expander("🔧 Cards Raw Response (디버그)"):
                st.text_area("Raw", raw, height=400)
        return None


# ─── API Call: Brainstorm Analysis (2단계) ───
def call_brainstorm_analysis(idea, genre, market, fmt, top3_cards, research=None, locked_block=""):
    """2단계: Top 3 기반 시장분석 + 차별화 + Gate A 채점"""
    try:
        client = get_client()

        system_prompt = P.SYSTEM_BRAINSTORM_ANALYSIS

        research_block = ""
        if research:
            research_block = f"""
[리서치 참고]
{json.dumps(research, ensure_ascii=False)}
"""

        user_prompt = f"""[입력]
아이디어: {idea}
장르: {genre}
타겟: {market}
포맷: {fmt}
{locked_block}

[Top 3 컨셉 카드]
{json.dumps(top3_cards, ensure_ascii=False)}
{research_block}

[JSON 스키마]
{{
  "market_context": {{
    "target_market": "",
    "market_insight": "",
    "cultural_code": "",
    "market_risk": "",
    "reference_titles": []
  }},
  "format_context": {{
    "selected_format": "",
    "format_rationale": "",
    "structure_note": ""
  }},
  "research_applied": {{
    "real_events_used": [],
    "inspiration_note": ""
  }},
  "hook_sentence": "",
  "differentiation": ["", "", ""],
  "development_priority": {{
    "recommended_direction": "",
    "next_step": "",
    "risk": ""
  }},
  "gate_a_scores": {{
    "protagonist_visible": 0.0,
    "conflict_one_line": 0.0,
    "differentiation": 0.0,
    "poster_image": 0.0,
    "market_potential": 0.0,
    "average": 0.0
  }}
}}

규칙:
- gate_a_scores.average = 5개 항목의 평균, 소수점 1자리 반올림
- 모든 점수는 0.0~10.0
- hook_sentence는 1위 컨셉의 핵심을 기반으로 작성
- market_insight, cultural_code, market_risk, format_rationale, structure_note는 각 2문장 이내
- differentiation은 정확히 3개
"""

        response = client.messages.create(
            model=ANTHROPIC_MODEL,
            max_tokens=2000,
            temperature=0.3,
            system=system_prompt,
            messages=[{"role": "user", "content": user_prompt}]
        )

        txt = "".join(
            block.text for block in response.content if hasattr(block, "text")
        ).strip()
        st.session_state["last_analysis_raw"] = txt

        return safe_json_loads(txt)

    except Exception as e:
        st.session_state["last_error"] = str(e)
        st.error(f"분석 실패: {e}")
        raw = st.session_state.get("last_analysis_raw")
        if raw:
            with st.expander("🔧 Analysis Raw Response (디버그)"):
                st.text_area("Raw", raw, height=400)
        return None


# ─── API Call: Core Build Main (1단계) ───
def call_core_build_main(idea, genre, market, fmt, selected_concept, research=None, locked_block="",
                          fact_based=False, historical=False, film_type=""):
    """Core Build 1단계: Logline + Intent + World + Character + Goal/Need/Strategy"""
    try:
        client = get_client()

        # v2.4.0: 시대극 자동 감지 컨텍스트 추출
        _locked_text, _idea_text = _get_period_scan_context({
            "idea_text": idea or "",
            "genre": genre,
            "title": (selected_concept or {}).get("title", "") if selected_concept else "",
            "locked_items": [locked_block] if locked_block else [],
        })
        system_prompt = P.build_system_core(
            genre, fact_based=fact_based, historical=historical, film_type=film_type,
            locked_text=_locked_text, idea_text=_idea_text,
        )

        research_block = ""
        if research:
            research_block = f"\n[리서치 참고]\n{json.dumps(research, ensure_ascii=False)}"

        user_prompt = f"""[입력]
아이디어: {idea}
장르: {genre}
타겟: {market}
포맷: {fmt}

[선정 컨셉]
{json.dumps(selected_concept, ensure_ascii=False)}
{research_block}
{locked_block}

[JSON 스키마]
{{
  "logline_pack": {{
    "original": "원본 로그라인 1문장",
    "washed": "서사 불순물 제거한 정제 로그라인 1문장",
    "investor": "투자자용 1문장",
    "director": "감독용 1문장",
    "character_hook": "배우용 캐릭터 훅 1문장"
  }},
  "project_intent": {{
    "subject": "소재 기획의도 — 이 소재가 왜 지금 필요한가, 어떤 현실/감정/사회적 맥락과 연결되는가 (2~3문장)",
    "genre_approach": "장르 기획의도 — 이 장르 접근이 왜 유효한가, 관객에게 어떤 체험을 주는가 (2~3문장)",
    "market_rationale": "시장 기획의도 — 타겟 시장에서 왜 통하는가, 어떤 수요/트렌드/공백에 대응하는가 (2~3문장)",
    "pitch": "엘리베이터 피치 1문장",
    "theme": "주제 1문장",
    "tone_manner": ["톤앤매너 키워드 3~5개"]
  }},
  "goal_need_strategy": {{
    "goal": "주인공의 목적/욕망 1문장",
    "need": "진짜 결핍 (본인은 모르는 것) 1문장",
    "strategy": "해결 전략/방법 1문장",
    "risk": "실패 시 잃는 것 1문장",
    "ending_payoff": "결말에서 G/N/S가 어떻게 회수되는가 1문장"
  }},
  "narrative_drive": {{
    "desire_origin": "loss 또는 lack — 주인공의 욕망이 상실에서 시작되었는가, 결핍에서 시작되었는가",
    "origin_detail": "구체적으로 무엇을 잃었는가(loss) 또는 무엇이 없었는가(lack) 1문장",
    "arc_direction": "loss면 회복/복수/대체 중 택1, lack이면 획득/성장/증명 중 택1",
    "resolution_strategy": "이 이야기만의 독특한 해결 방식 — 기존작과 뭐가 다른가 1문장",
    "goal_need_gap": "Goal(외적 욕망)과 Need(내적 필요) 사이의 간극 1문장 — 이 간극이 이야기를 끌고 간다",
    "cost": {{
      "relation": "주인공의 Strategy가 주변 관계를 어떻게 망가뜨리는가 1문장 — 2막 후반에 누적되어 실재 손상이 될 것",
      "self": "주인공의 Strategy가 주인공 자신을 어떻게 마모시키는가 1문장 — 3막에서 자기 자신과 직면하게 할 것",
      "world": "주인공의 Strategy가 주인공이 지키려던 세계(공동체/직업/가치)를 어떻게 무너뜨리는가 1문장"
    }},
    "strategy_transformation": "3막에서 Strategy_1이 어떻게 Strategy_2로 전환되는가 1문장 — 외적 선택이 아닌 내적 전환으로 기술"
  }},
  "bjnd_four_axis_check": {{
    "necessity": {{
      "question": "이 인물이 왜 지금 이렇게 살아야만 했는가?",
      "answer": "Lack/Loss의 필연성을 입증하는 1~2문장 — 막연한 설정이 아닌 구체적 사건/기억 기반",
      "life_permeation": "이 Lack/Loss이 인물의 일상을 관통하는 3가지 방식 (각 1문장)"
    }},
    "authenticity": {{
      "question": "이 Strategy는 이 인물이 아니면 불가능한가?",
      "substitution_test": "이 Strategy를 다른 유명 작품 주인공에게 붙여도 성립하는가? PASS(다른 인물에게 붙일 수 없음) 또는 FAIL(누구나 쓸 수 있음)",
      "organic_link": "직업 + Lack + Strategy를 한 문장으로 유기적으로 연결",
      "author_justification": "왜 이 인물은 이런 방식으로 해결하려 하는가? 작가 시선이 담긴 1문장"
    }},
    "empathy": {{
      "question": "관객이 이 인물의 Cost를 보며 자기 자신을 떠올릴 것인가?",
      "universal_wound": "이 Cost가 건드리는 동시대 보편 상처 1문장",
      "audience_mirror": "관객이 이 인물에서 자기를 발견할 지점 1문장",
      "resonance_target": "특히 공감할 관객층 (예: 30대 여성 직장인, 40대 남성 아버지 등)"
    }},
    "potency": {{
      "question": "이 인물의 Strategy 전환이 관객 마음에 얼마나 큰 흔적을 남길 것인가?",
      "strategy_arc": "1막 Strategy_1 → 3막 Strategy_2, 한 문장 대비",
      "signature_moment": "변화가 완성되는 결정적 한 순간 1문장",
      "afterimage": "영화 끝난 뒤에도 관객이 떠올릴 잔상 이미지 1컷"
    }},
    "audience_scenes": {{
      "laughter_scene": "관객이 웃을 장면 1~2개 구체 (로코/코미디 아닐 경우 빈 문자열)",
      "heartache_scene": "관객이 가슴 먹먹해질 장면 1~2개 구체",
      "self_reflection_scene": "관객이 '나였다면 어떻게 했을까'를 고민할 장면 1개 구체",
      "memorable_line_or_image": "영화 끝난 후 관객 입에 남을 장면 1개"
    }}
  }},
  "genre_expectation_check": {{
    "genre_fun_alive": "true 또는 false — 이 장르의 재미 본질이 살아있는가",
    "genre_fun_diagnosis": "이 작품에서 장르 재미가 어떻게 살아있는지 3줄 진단",
    "checklist": {{
      "item_label_1": "YES 또는 NO + 1문장 근거 — 해당 장르 체크리스트 항목 1",
      "item_label_2": "YES 또는 NO + 1문장 근거 — 해당 장르 체크리스트 항목 2",
      "item_label_3": "YES 또는 NO + 1문장 근거 — 해당 장르 체크리스트 항목 3",
      "item_label_4": "YES 또는 NO + 1문장 근거 — 해당 장르 체크리스트 항목 4",
      "item_label_5": "YES 또는 NO + 1문장 근거 — 해당 장르 체크리스트 항목 5",
      "item_label_6": "YES 또는 NO + 1문장 근거 — 해당 장르 체크리스트 항목 6 (있으면)",
      "item_label_7": "YES 또는 NO + 1문장 근거 — 해당 장르 체크리스트 항목 7 (있으면)"
    }},
    "weak_zones": ["장르 기대가 약할 것으로 예상되는 비트 구간 명시 (예: '클라이맥스가 로맨스 완성 아님', '2막 Fun and Games 부족')"],
    "climax_verdict": "CLIMAX_PASS 또는 CLIMAX_FAIL — 장르 약속에 맞는 클라이맥스인가. FAIL이면 3막 전체 재설계 강제"
  }},
  "opening_strategy": {{
    "opening_type": "Action Drop / Cold Open / Tease & Reveal / In Media Res / Character Reveal Action / Hook Dialogue 중 택1 — 시스템 프롬프트의 오프닝 DNA 매핑에 따른 권장 기법을 우선하되, 작품 특성에 맞게 택1",
    "opening_intent": "이 오프닝으로 무엇을 잡는가 — 캐릭터 / 세계 / 사건 / 톤 / 테마 중 1~2개 구체화",
    "dopamine_point": "첫 3분 안에 관객이 느낄 구체적 감각 — 충격 / 웃음 / 긴장 / 경이 / 호기심 / 감정 울림 중 1개 선택 + 어떤 방식으로 터뜨리는가 1문장",
    "opening_to_act1_link": "오프닝과 1막 본편의 연결 논리 — 동일 시점 진입 / 플래시백 회수 / 톤 브리지 / 병렬 구조 중 어떤 방식인가 1문장",
    "opening_hook_line": "오프닝의 기억에 남을 한 줄 대사 또는 한 이미지 — 실제 시나리오에 쓸 수 있는 수준으로 구체적으로",
    "genre_dna_check": "이 오프닝이 장르 DNA 자가검증 질문에 어떻게 대답하는가 — 예) 호러면 '첫 3분 안에 공포의 규칙/흔적이 심어졌는가?'에 어떻게 YES인가 1문장"
  }},
  "world_build": {{
    "time": "시간 배경",
    "space": "공간 배경",
    "rules": "세계관 규칙 1문장",
    "taboo": "금기 1문장",
    "power_structure": "권력 구조 1문장",
    "visual_keywords": ["시각 이미지 키워드 3~5개"],
    "conflict_points": ["세계관 충돌 포인트 3개"]
  }},
  "characters": [
    {{
      "role": "protagonist",
      "name": "",
      "description": "인물 소개 1문장",
      "goal": "이 인물의 욕망 1문장",
      "need": "결핍 1문장",
      "strategy": "행동 방식 1문장",
      "flaw": "결점 1문장",
      "arc": "변화 1문장",
      "dialogue_tone": "대사 톤 키워드"
    }},
    {{
      "role": "antagonist",
      "name": "",
      "description": "",
      "goal": "",
      "need": "",
      "strategy": "",
      "flaw": "",
      "arc": "",
      "dialogue_tone": ""
    }},
    {{
      "role": "ally",
      "name": "",
      "description": "",
      "goal": "",
      "flaw": "",
      "dialogue_tone": ""
    }},
    {{
      "role": "mirror",
      "name": "",
      "description": "",
      "goal": "",
      "flaw": "",
      "dialogue_tone": ""
    }}
  ],
  "extended_characters": [
    {{
      "role": "역할명 (catalyst/subplot_lead/mentor/rival/informant/love_interest 등 자유)",
      "name": "",
      "description": "이 인물이 이야기에 왜 필요한가 1문장",
      "goal": "",
      "flaw": "",
      "dialogue_tone": ""
    }}
  ],
  "relationship_map": [
    "주인공↔적대자: 관계 1문장",
    "주인공↔조력자: 관계 1문장",
    "주인공↔거울: 관계 1문장",
    "확장 캐릭터 간 핵심 관계: 1문장씩"
  ],
  "attraction_design": {{
    "opening_hook": {{
      "type": "A1(In Medias Res) / A2(충격적 사건) / A3(결함 드러남) 중 선택",
      "description": "첫 장면 구체적 묘사 — 어떤 장면으로 시작하는가 2~3문장",
      "forbidden_check": "설명/일상/내레이션으로 시작하지 않는다는 확인 1문장"
    }},
    "twist_point": {{
      "expected_direction": "관객이 예상할 전개 1문장",
      "betrayal": "그 예상을 깨는 배반 포인트 1문장 — 이 이야기만의 방향",
      "why_more_true": "배반이 단순 반전이 아니라 더 진실된 방향인 이유 1문장"
    }},
    "water_cooler_moment": {{
      "scene_or_setup": "관객이 다음 날 누군가에게 말하고 싶어지는 장면 또는 설정 1문장",
      "why_memorable": "왜 기억에 남는가 1문장"
    }},
    "korean_specificity": [
      "이 이야기의 한국적 구체성 — 추상어 금지, 공간/제도/관계로 표현 (3개)"
    ],
    "villain_logic": "빌런이 자신만의 논리 안에서 옳다고 믿는 이유 1문장",
    "emotional_explosion": {{
      "suppression": "어떤 감정을 언제까지 억누르는가 1문장",
      "explosion_moment": "어느 장면에서 터지는가 1문장"
    }},
    "forbidden_directions": [
      "이 이야기가 절대 가면 안 되는 방향 3개 — 뻔해지는 순간들"
    ],
    "planting_payoff": [
      {{
        "type": "character / relationship / world 중 선택",
        "plant": "1막에 심는 것 — 대사/소품/습관/장소/이미지 1문장",
        "plant_scene": "어느 장면에서 심는가 (Beat 번호 또는 상황)",
        "payoff": "2막 후반~3막에서 회수 — 새로운 의미로 돌아오는 순간 1문장",
        "payoff_scene": "어느 장면에서 회수하는가 (Beat 번호 또는 상황)"
      }}
    ]
  }},
  "research_applied": {{
    "references_used": [
      {{
        "source_type": "real_event / existing_work",
        "source_id": 1,
        "source_title": "리서치 JSON에서 참조한 아이템의 제목",
        "applied_to": "character_background / world_rule / plot_event / motif / dialogue_reference / visual_detail 중 선택",
        "how_applied": "이 리서치 요소를 이 작품 어디에 어떻게 녹였는가 1문장 — 추상 금지, 구체적으로",
        "character_name": "해당 캐릭터 이름 (캐릭터에 적용된 경우, 아니면 빈 문자열)"
      }}
    ],
    "verisimilitude_anchors": [
      "핍진성 앵커 1 — Treatment·Scene·시나리오에서 반드시 유지할 현실 기반 디테일 1문장",
      "핍진성 앵커 2",
      "핍진성 앵커 3"
    ],
    "research_absorption_note": "리서치를 어떻게 소화했는지 작가 시점 1문장 — 단순 나열이 아닌 '이 작품만의 방식으로 어떻게 흡수되었는가'"
  }}
}}

규칙:
- logline_pack 각 버전은 관점만 다르고 같은 이야기를 가리켜야 한다.
- attraction_design은 이 이야기의 매력 설계도다. 가장 구체적으로 작성할 것.
- attraction_design.opening_hook.description은 실제 첫 장면을 그릴 수 있을 만큼 구체적으로.
- attraction_design.water_cooler_moment는 오징어게임의 달고나, 기생충의 냄새처럼 누군가에게 말하고 싶어지는 것.
- attraction_design.korean_specificity는 추상어(가난, 차별, 압박) 금지. 반드시 구체적 명사(반지하, 수능, 연습생 계약서)로.
- attraction_design.planting_payoff는 최소 3개. type별 최소 1개씩 (character/relationship/world). Plant 없는 Payoff(데우스 엑스 마키나) 금지.
- goal_need_strategy는 이 작품의 서사 엔진이다. 가장 정밀하게 작성할 것.
- narrative_drive는 BLUE JEANS 고유 서사동력 프레임워크다. desire_origin이 loss인지 lack인지를 정확히 진단하라. 이것이 이야기 전체의 방향을 결정한다.
- characters는 필수 4명(protagonist/antagonist/ally/mirror). extended_characters는 이야기가 필요로 하는 만큼 0~4명 추가 (최대 총 8명). 영화는 4~5명, 미니시리즈는 6~8명이 적정. 각 인물의 goal이 서로 달라야 한다.
- extended_characters의 role은 자유. catalyst(촉매자), subplot_lead(서브플롯 리드), mentor(멘토), rival(라이벌), informant(정보원), love_interest(연인) 등 이야기에 맞는 역할명을 직접 지정.
- world_build의 conflict_points는 세계관이 만들어내는 갈등이어야 한다.
- [v2.3.8 신규] research_applied는 '이 작품에 실제로 응용된 리서치 요소만' 선별 기록. 리서치에 있었으나 이 작품에 녹지 못한 것은 제외.
  references_used는 최소 2개, 최대 6개. 모든 항목의 how_applied는 '어디에/어떻게'가 구체적으로 드러나야 한다.
  verisimilitude_anchors는 정확히 3~5개. Treatment·시나리오에서 Writer Engine이 반드시 유지할 핵심 디테일.
  리서치가 제공되지 않은 경우에도 references_used는 빈 배열로 두되, verisimilitude_anchors는 작품 자체의 내적 설정 기반으로 3개 생성.
"""

        response = client.messages.create(
            model=ANTHROPIC_MODEL,
            max_tokens=16000,
            temperature=0.3,
            system=system_prompt,
            messages=[{"role": "user", "content": user_prompt}]
        )

        if response.stop_reason == "max_tokens":
            st.warning("⚠️ Core Build 응답 잘림. 재시도...")
            response = client.messages.create(
                model=ANTHROPIC_MODEL,
                max_tokens=16000,
                temperature=0.3,
                system=system_prompt,
                messages=[{"role": "user", "content": user_prompt}]
            )

        txt = "".join(
            block.text for block in response.content if hasattr(block, "text")
        ).strip()
        st.session_state["last_core_raw"] = txt

        return safe_json_loads(txt)

    except Exception as e:
        st.session_state["last_error"] = str(e)
        st.error(f"Core Build 실패: {e}")
        raw = st.session_state.get("last_core_raw", "")
        if raw:
            with st.expander("🔧 Core Build Raw Response (디버그)"):
                st.text_area("Raw", raw, height=400)
        return None


# ─── API Call: Core Gate (2단계) ───
def call_core_gate(core_data):
    """Core Build 2단계: Gate B (Drive) + Gate C (Character) 채점"""
    try:
        client = get_client()

        system_prompt = P.SYSTEM_CORE_GATE

        user_prompt = f"""[Core Build 결과]
{json.dumps(core_data, ensure_ascii=False)}

[JSON 스키마]
{{
  "gate_b_drive": {{
    "goal_clarity": 0.0,
    "need_from_loss": 0.0,
    "strategy_creative": 0.0,
    "failure_cost": 0.0,
    "average": 0.0,
    "feedback": "Gate B 종합 피드백 1문장"
  }},
  "gate_c_character": {{
    "protagonist_antagonist_logic": 0.0,
    "supporting_not_functional": 0.0,
    "relationship_produces_conflict": 0.0,
    "average": 0.0,
    "feedback": "Gate C 종합 피드백 1문장"
  }},
  "five_axis_scores": {{
    "goal": 0.0,
    "need": 0.0,
    "strategy": 0.0,
    "structure": 0.0,
    "character_concept": 0.0,
    "final_score": 0.0,
    "verdict": "개발 진행 | 개발 보류 | 구조 재설계"
  }}
}}

규칙:
- gate_b average = 4항목 평균
- gate_c average = 3항목 평균
- five_axis: structure는 아직 Structure Build 전이므로 goal/need/strategy 기반으로 잠정 추정
- final_score = 0.20*goal + 0.20*need + 0.20*strategy + 0.25*structure + 0.15*character_concept
- verdict: 8.0 이상 → 개발 진행, 6.0~7.9 → 개발 보류, 5.9 이하 → 구조 재설계
"""

        response = client.messages.create(
            model=ANTHROPIC_MODEL,
            max_tokens=2000,
            temperature=0.2,
            system=system_prompt,
            messages=[{"role": "user", "content": user_prompt}]
        )

        txt = "".join(
            block.text for block in response.content if hasattr(block, "text")
        ).strip()
        st.session_state["last_gate_raw"] = txt

        return safe_json_loads(txt)

    except Exception as e:
        st.session_state["last_error"] = str(e)
        st.error(f"Gate 채점 실패: {e}")
        raw = st.session_state.get("last_gate_raw", "")
        if raw:
            with st.expander("🔧 Gate Raw Response (디버그)"):
                st.text_area("Raw", raw, height=400)
        return None


# ─── API Call: Character Bible ───
def call_character_bible_single(char_data, all_chars_names, core_data, genre, fmt, locked_block="",
                                 fact_based=False, historical=False, film_type="",
                                 prior_chars_block=""):
    """캐릭터 바이블 — 캐릭터 1인 단위 호출
    
    Args:
        prior_chars_block: 이미 생성된 다른 캐릭터들의 일관성 사실 블록 (v2.3.1 신규).
                          빈 문자열이면 일관성 참조 없이 생성 (첫 번째 캐릭터).
    """
    try:
        client = get_client()
        gns = core_data.get("goal_need_strategy", {})
        nd = core_data.get("narrative_drive", {})
        lp = core_data.get("logline_pack", {})
        wb = core_data.get("world_build", {})

        char_json = json.dumps(char_data, ensure_ascii=False)
        gns_json = json.dumps(gns, ensure_ascii=False)
        wb_json = json.dumps(wb, ensure_ascii=False)

        role = char_data.get("role", "")
        name = char_data.get("name", "")
        other_names = [n for n in all_chars_names if n != name]
        others_str = ", ".join(other_names) if other_names else "없음"

        # v2.4.0: 시대극 자동 감지 컨텍스트 추출
        _locked_text, _idea_text = _get_period_scan_context({
            "genre": genre,
            "core_data": core_data or {},
            "locked_items": [locked_block] if locked_block else [],
        })
        system_prompt = P.build_system_char_bible(genre, fmt, others_str,
                                                    fact_based=fact_based, historical=historical, film_type=film_type,
                                                    locked_text=_locked_text, idea_text=_idea_text)

        schema = P.CHAR_BIBLE_SCHEMA.replace("ROLE_PLACEHOLDER", role).replace("NAME_PLACEHOLDER", name)

        user_prompt = f"""[확장할 캐릭터]
{char_json}

[작품 Goal/Need/Strategy]
{gns_json}

[세계관]
{wb_json}

[로그라인]
{lp.get("washed","")}

[다른 캐릭터 이름 (관계별 태도 작성용)]
{others_str}
{locked_block}
{prior_chars_block}
[JSON 스키마 — 이 캐릭터 1인분만 출력]
{schema}

{P.CHAR_BIBLE_RULES}"""

        response = client.messages.create(
            model=ANTHROPIC_MODEL_OPUS, max_tokens=16000, temperature=0.3,
            system=system_prompt,
            messages=[{"role": "user", "content": user_prompt}]
        )
        # max_tokens 잘림 시 분량 줄여서 재시도
        if response.stop_reason == "max_tokens":
            retry_prompt = user_prompt.replace(
                "외형·첫인상 3문장", "외형·첫인상 2문장"
            ).replace(
                "과거사 5문장", "과거사 3문장"
            ).replace(
                "1막 끝 상태 2문장", "1막 끝 상태 1문장"
            ).replace(
                "미드포인트 상태 2문장", "미드포인트 상태 1문장"
            ).replace(
                "클라이맥스 상태 2문장", "클라이맥스 상태 1문장"
            )
            response = client.messages.create(
                model=ANTHROPIC_MODEL_OPUS, max_tokens=16000, temperature=0.3,
                system=system_prompt,
                messages=[{"role": "user", "content": retry_prompt}]
            )
        txt = "".join(b.text for b in response.content if hasattr(b, "text")).strip()
        st.session_state[f"last_char_bible_{role}_raw"] = txt

        # JSON 파싱 시도 — 실패 시 자동 재시도 1회
        try:
            return safe_json_loads(txt)
        except json.JSONDecodeError:
            st.warning(f"⚠️ {name} JSON 파싱 실패 — 자동 재시도 중...")
            # 재시도: JSON 규칙을 더 강하게 강조
            retry_system = system_prompt + "\n\n★ 최우선 규칙: 반드시 유효한 JSON만 출력. 대사 안의 따옴표는 작은따옴표만. 줄바꿈 금지. 역슬래시 금지. ★"
            response2 = client.messages.create(
                model=ANTHROPIC_MODEL_OPUS, max_tokens=16000, temperature=0.2,
                system=retry_system,
                messages=[{"role": "user", "content": user_prompt}]
            )
            txt2 = "".join(b.text for b in response2.content if hasattr(b, "text")).strip()
            st.session_state[f"last_char_bible_{role}_raw"] = txt2
            return safe_json_loads(txt2)

    except Exception as e:
        st.error(f"Character Bible ({name}) 생성 실패: {e}")
        raw = st.session_state.get(f"last_char_bible_{role}_raw", "")
        if raw:
            with st.expander(f"🔧 {name} Raw (디버그)"):
                st.text_area("Raw", raw, height=400)
        return None


def call_character_bible(core_data, genre, fmt, locked_block="",
                         fact_based=False, historical=False, film_type=""):
    """캐릭터 바이블 — characters + extended_characters 모두 처리
    
    v2.3.1: 순차 생성 시 이전 캐릭터의 핵심 사실을 다음 캐릭터 생성 프롬프트에 주입하여
            가족 관계·학력·나이 등의 일관성 모순을 차단한다.
    """
    chars = core_data.get("characters", [])
    ext_chars = core_data.get("extended_characters", [])
    all_chars = chars + ext_chars
    all_names = [c.get("name", f"캐릭터{i+1}") for i, c in enumerate(all_chars)]

    result_chars = []
    prior_facts = []  # v2.3.1: 이전 캐릭터들의 일관성 사실 누적

    for i, ch in enumerate(all_chars):
        name = ch.get("name", f"캐릭터{i+1}")
        role = ch.get("role", "")

        # v2.3.1: 이전 캐릭터 맥락 블록 생성
        prior_chars_block = P.build_prior_chars_consistency_block(prior_facts)

        if i == 0:
            st.info(f"📖 {i+1}/{len(all_chars)} — {name} ({role}) 바이블 생성 중... (최초 캐릭터)")
        else:
            st.info(f"📖 {i+1}/{len(all_chars)} — {name} ({role}) 바이블 생성 중... "
                    f"(이전 {len(prior_facts)}명과 일관성 검증 중)")

        char_result = call_character_bible_single(
            ch, all_names, core_data, genre, fmt,
            locked_block=locked_block,
            fact_based=fact_based, historical=historical, film_type=film_type,
            prior_chars_block=prior_chars_block,  # v2.3.1 신규
        )

        if char_result:
            result_chars.append(char_result)
            # v2.3.1: 성공한 캐릭터의 일관성 사실을 추출하여 다음 캐릭터 생성 시 참조
            facts = P.extract_char_consistency_facts(char_result)
            prior_facts.append(facts)
        else:
            st.warning(f"⚠️ {name} 바이블 생성 실패. 건너뜁니다.")

    if result_chars:
        return {"characters": result_chars}
    return None


# ─── API Call: Structure Build Story (1단계) ───
def call_structure_story(core_data, genre, market, fmt, locked_block="",
                          fact_based=False, historical=False, film_type=""):
    """Structure 1단계: Synopsis 1P + Storyline"""
    try:
        client = get_client()
        gns = core_data.get("goal_need_strategy", {})
        nd = core_data.get("narrative_drive", {})
        lp = core_data.get("logline_pack", {})
        chars = core_data.get("characters", []) + core_data.get("extended_characters", [])

        # v2.4.0: 시대극 자동 감지 컨텍스트 추출
        _locked_text, _idea_text = _get_period_scan_context({
            "genre": genre,
            "core_data": core_data or {},
            "locked_items": [locked_block] if locked_block else [],
        })
        system_prompt = P.build_system_structure_story(
            fact_based=fact_based, historical=historical, film_type=film_type,
            locked_text=_locked_text, idea_text=_idea_text,
        )

        user_prompt = f"""[Core Build 요약]
로그라인: {lp.get("washed", lp.get("original", ""))}
장르: {genre} / 타겟: {market} / 포맷: {fmt}
{locked_block}
Goal: {gns.get("goal","")}
Need: {gns.get("need","")}
Strategy: {gns.get("strategy","")}
캐릭터: {json.dumps(chars, ensure_ascii=False)}

[JSON 스키마]
{{
  "synopsis_1p": {{
    "opening": "시작 상황",
    "catalyst": "촉발 사건",
    "development": "전개",
    "midpoint": "미드포인트 전환",
    "collapse": "붕괴/위기",
    "climax": "결전",
    "ending": "결말"
  }},
  "storyline": [
    {{
      "seq": 1,
      "label": "시퀀스 라벨 (동사→동사 형식)",
      "pages": "p.1~15",
      "summary": "이 구간 전개 요약",
      "conflict": "이 구간의 갈등",
      "emotion": "감정선 키워드",
      "hook": "다음 구간으로 당기는 힘"
    }}
  ]
}}

규칙:
- synopsis_1p 각 항목 2문장 이내
- storyline은 8개 시퀀스 (영화 기준)
- 시리즈면 시퀀스 대신 에피소드 단위 6~8개
- 각 시퀀스의 summary, conflict, hook 1문장씩
"""

        response = client.messages.create(
            model=ANTHROPIC_MODEL, max_tokens=16000, temperature=0.3,
            system=system_prompt,
            messages=[{"role": "user", "content": user_prompt}]
        )
        if response.stop_reason == "max_tokens":
            response = client.messages.create(
                model=ANTHROPIC_MODEL, max_tokens=16000, temperature=0.3,
                system=system_prompt,
                messages=[{"role": "user", "content": user_prompt}]
            )
        txt = "".join(b.text for b in response.content if hasattr(b, "text")).strip()
        st.session_state["last_structure_story_raw"] = txt
        return safe_json_loads(txt)
    except Exception as e:
        st.error(f"Story 생성 실패: {e}")
        raw = st.session_state.get("last_structure_story_raw", "")
        if raw:
            with st.expander("🔧 Story Raw (디버그)"):
                st.text_area("Raw", raw, height=400)
        return None


# ─── API Call: Structure Build Diagnosis + Character Arcs (2단계) ───
def call_structure_diagnosis(core_data, story_data, genre, fmt, locked_block="",
                              fact_based=False, historical=False, film_type=""):
    """Structure 2단계: 3막 구조 진단 + 15비트 + 캐릭터 관계 변화표"""
    try:
        client = get_client()
        gns = core_data.get("goal_need_strategy", {})
        nd = core_data.get("narrative_drive", {})
        chars = core_data.get("characters", []) + core_data.get("extended_characters", [])

        # v2.4.0: 시대극 자동 감지 컨텍스트 추출
        _locked_text, _idea_text = _get_period_scan_context({
            "genre": genre,
            "core_data": core_data or {},
            "locked_items": [locked_block] if locked_block else [],
        })
        system_prompt = P.build_system_structure_diagnosis(
            fact_based=fact_based, historical=historical, film_type=film_type,
            locked_text=_locked_text, idea_text=_idea_text,
        )

        user_prompt = f"""[입력]
장르: {genre} / 포맷: {fmt}
{locked_block}
Goal: {gns.get("goal","")} / Need: {gns.get("need","")} / Strategy: {gns.get("strategy","")}
서사동력: {nd.get("desire_origin","")}({nd.get("origin_detail","")}) → {nd.get("arc_direction","")} / 해결전략: {nd.get("resolution_strategy","")}

[캐릭터]
{json.dumps(chars, ensure_ascii=False)}

[시놉시스]
{json.dumps(story_data.get("synopsis_1p", {}), ensure_ascii=False)}

[스토리라인]
{json.dumps(story_data.get("storyline", []), ensure_ascii=False)}

[JSON 스키마]
{{
  "three_act": {{
    "act1_end": "1막 끝 전환점",
    "act2_midpoint": "미드포인트",
    "act2_end": "2막 끝 전환점 (All Is Lost)",
    "act3_climax": "클라이맥스",
    "diagnosis": "3막 구조 종합 진단 1문장"
  }},
  "beat_sheet": [
    {{
      "beat": "Opening Image",
      "status": "있음|약함|없음",
      "note": "근거 1문장"
    }}
  ],
  "character_arcs": [
    {{
      "name": "캐릭터 이름",
      "role": "protagonist|antagonist|ally|mirror|catalyst|subplot_lead",
      "act1_state": "1막에서의 상태/태도 1문장",
      "turning_point": "변화를 촉발하는 사건 1문장",
      "act3_state": "3막에서의 변화된 상태 1문장",
      "arc_type": "성장|몰락|각성|반전|정체"
    }}
  ],
  "relationship_changes": [
    {{
      "pair": "캐릭터A ↔ 캐릭터B",
      "act1": "1막 관계 상태",
      "midpoint": "미드포인트 관계 전환",
      "act3": "3막 관계 도착점"
    }}
  ]
}}

규칙:
- beat_sheet는 15비트 전체 (Opening Image ~ Final Image)
- beat status는 반드시 있음/약함/없음 중 하나
- character_arcs는 Core Build의 캐릭터 전원
- relationship_changes는 핵심 관계 3~4쌍
- 모든 필드 1문장 이내로 압축
"""

        response = client.messages.create(
            model=ANTHROPIC_MODEL, max_tokens=16000, temperature=0.25,
            system=system_prompt,
            messages=[{"role": "user", "content": user_prompt}]
        )
        if response.stop_reason == "max_tokens":
            response = client.messages.create(
                model=ANTHROPIC_MODEL, max_tokens=16000, temperature=0.25,
                system=system_prompt,
                messages=[{"role": "user", "content": user_prompt}]
            )
        txt = "".join(b.text for b in response.content if hasattr(b, "text")).strip()
        st.session_state["last_structure_diag_raw"] = txt
        return safe_json_loads(txt)
    except Exception as e:
        st.error(f"구조 진단 실패: {e}")
        raw = st.session_state.get("last_structure_diag_raw", "")
        if raw:
            with st.expander("🔧 Diagnosis Raw (디버그)"):
                st.text_area("Raw", raw, height=400)
        return None


# ─── API Call: Structure Gate D (3단계) ───
def call_structure_gate(story_data, diag_data):
    """Structure 3단계: Gate D 채점"""
    try:
        client = get_client()
        system_prompt = P.SYSTEM_STRUCTURE_GATE

        user_prompt = f"""[시놉시스]
{json.dumps(story_data.get("synopsis_1p",{}), ensure_ascii=False)}

[3막 구조]
{json.dumps(diag_data.get("three_act",{}), ensure_ascii=False)}

[15비트]
{json.dumps(diag_data.get("beat_sheet",[]), ensure_ascii=False)}

[JSON 스키마]
{{
  "gate_d_structure": {{
    "turning_points_valid": 0.0,
    "midpoint_redirects": 0.0,
    "all_is_lost_works": 0.0,
    "ending_inevitable_surprising": 0.0,
    "average": 0.0,
    "feedback": "Gate D 종합 피드백 1문장"
  }}
}}
규칙: average = 4항목 평균."""

        response = client.messages.create(
            model=ANTHROPIC_MODEL, max_tokens=1500, temperature=0.2,
            system=system_prompt,
            messages=[{"role": "user", "content": user_prompt}]
        )
        txt = "".join(b.text for b in response.content if hasattr(b, "text")).strip()
        st.session_state["last_structure_gate_raw"] = txt
        return safe_json_loads(txt)
    except Exception as e:
        st.error(f"Gate D 실패: {e}")
        return None


# ─── API Call: 기승전결 1pg 줄글 ───
def call_structure_prose(core_data, story_data):
    """기승전결 구조의 1페이지 줄글 시놉시스"""
    try:
        client = get_client()
        lp = core_data.get("logline_pack", {})
        gns = core_data.get("goal_need_strategy", {})
        nd = core_data.get("narrative_drive", {})
        syn = story_data.get("synopsis_1p", {})
        storyline = story_data.get("storyline", [])

        user_prompt = f"""[작품 정보]
로그라인: {lp.get("washed", lp.get("original",""))}
Goal: {gns.get("goal","")} / Need: {gns.get("need","")} / Strategy: {gns.get("strategy","")}
서사동력: {nd.get("desire_origin","")}({nd.get("origin_detail","")}) → {nd.get("arc_direction","")} / 해결전략: {nd.get("resolution_strategy","")}

[시놉시스]
{json.dumps(syn, ensure_ascii=False)}

[스토리라인]
{json.dumps(storyline, ensure_ascii=False)}

[요청]
위 정보를 기반으로 기-승-전-결 구조의 1페이지 분량 줄글 시놉시스를 작성하라.

규칙:
- 기(起): 상황 설정과 주인공 소개
- 승(承): 사건 전개와 갈등 심화
- 전(轉): 반전과 위기
- 결(結): 해결과 마무리
- 전체 500~700자 분량
- 줄글로 작성. 번호나 구분 기호 없이 자연스러운 서술체.
- 한국어로 작성.
- JSON 형식: {{"prose": "줄글 텍스트"}}
- JSON만 출력. JSON 외 텍스트 금지."""

        response = client.messages.create(
            model=ANTHROPIC_MODEL, max_tokens=2000, temperature=0.3,
            system=P.SYSTEM_STRUCTURE_PROSE,
            messages=[{"role": "user", "content": user_prompt}]
        )
        txt = "".join(b.text for b in response.content if hasattr(b, "text")).strip()
        return safe_json_loads(txt)
    except Exception as e:
        st.error(f"줄글 시놉시스 생성 실패: {e}")
        return None


# ─── API Call: Scene Design (장면화) ───
def call_scene_design(core_data, story_data, diag_data, genre, fmt, locked_block="",
                      fact_based=False, historical=False, film_type=""):
    """Scene Design: 핵심 장면 15~18개 설계 (Show, don't tell)"""
    try:
        client = get_client()
        gns = core_data.get("goal_need_strategy", {})
        nd = core_data.get("narrative_drive", {})
        lp = core_data.get("logline_pack", {})
        chars = core_data.get("characters", []) + core_data.get("extended_characters", [])
        storyline = story_data.get("storyline", [])

        # v2.4.0: 시대극 자동 감지 컨텍스트 추출
        _locked_text, _idea_text = _get_period_scan_context({
            "genre": genre,
            "core_data": core_data or {},
            "locked_items": [locked_block] if locked_block else [],
        })
        system_prompt = P.build_system_scene_design(
            genre, fact_based=fact_based, historical=historical, film_type=film_type,
            locked_text=_locked_text, idea_text=_idea_text,
        )

        chars_simple = json.dumps(
            [{"name": c.get("name",""), "role": c.get("role","")} for c in chars],
            ensure_ascii=False
        )
        storyline_json = json.dumps(storyline, ensure_ascii=False)

        # ── 매력 설계도 추출 ──
        ad = core_data.get("attraction_design", {})
        scene_attraction_block = ""
        if ad:
            oh = ad.get("opening_hook", {})
            tp = ad.get("twist_point", {})
            wc = ad.get("water_cooler_moment", {})
            ee = ad.get("emotional_explosion", {})
            fd = ad.get("forbidden_directions", [])
            scene_attraction_block = f"""
[⚡ 매력 설계도 — 장면 설계 최우선 명령]
첫 장면 유형: {oh.get("type", "")}
첫 장면 (scene_no 1 반드시 이것으로): {oh.get("description", "")}
배반 포인트: {tp.get("betrayal", "")}
Water Cooler Moment (반드시 key_scenes 안에 포함): {wc.get("scene_or_setup", "")}
감정 폭발 장면: {ee.get("explosion_moment", "")}
절대 금지 방향: {" / ".join(fd)}
"""

        # ── 오프닝 전략 추출 (v2.2 신규) ──
        os_data = core_data.get("opening_strategy", {})
        opening_strategy_block = ""
        if os_data:
            opening_strategy_block = f"""
[🎬 오프닝 전략 — S#1 절대 명령 (v2.2)]
오프닝 타입: {os_data.get("opening_type", "")}
오프닝 의도: {os_data.get("opening_intent", "")}
도파민 포인트: {os_data.get("dopamine_point", "")}
1막 연결 방식: {os_data.get("opening_to_act1_link", "")}
훅 라인/이미지: {os_data.get("opening_hook_line", "")}
장르 DNA 체크: {os_data.get("genre_dna_check", "")}

★ S#1(scene_no 1)은 반드시 위 '오프닝 타입'에 부합하는 장면이어야 한다. ★
★ S#1의 dramatic_action, visual_direction, key_line에 '훅 라인/이미지'가 구체적으로 드러나야 한다. ★
★ S#1의 emotional_beat는 '도파민 포인트'에 명시된 감각(충격/웃음/긴장/경이/호기심/감정 울림)을 만드는 방향이어야 한다. ★
★ S#1을 '평범한 일상 소개' '풍경 묘사 후 인물 줌인' '주인공 배경 설명'으로 열면 실패다. 첫 장면에서 장르 DNA가 선언되어야 한다. ★
"""

        # ── BJND 씬 레벨 집행 블록 (v2.3 신규) — Strategy/Cost를 씬에 강제 반영 ──
        bjnd_scene_block = P.build_bjnd_scene_block(core_data)

        user_prompt = f"""[Core]
로그라인: {lp.get("washed","")}
Goal: {gns.get("goal","")} / Need: {gns.get("need","")} / Strategy: {gns.get("strategy","")}
서사동력: {nd.get("desire_origin","")}({nd.get("origin_detail","")}) → {nd.get("arc_direction","")} / 해결전략: {nd.get("resolution_strategy","")}
캐릭터: {chars_simple}
{scene_attraction_block}
{opening_strategy_block}
{bjnd_scene_block}
{locked_block}
[Storyline]
{storyline_json}

[JSON 스키마]
{P.SCENE_DESIGN_SCHEMA}

{P.SCENE_DESIGN_RULES}
"""

        response = client.messages.create(
            model=ANTHROPIC_MODEL, max_tokens=16000, temperature=0.3,
            system=system_prompt,
            messages=[{"role": "user", "content": user_prompt}]
        )
        if response.stop_reason == "max_tokens":
            retry = user_prompt.replace("15~18개", "12~15개")
            response = client.messages.create(
                model=ANTHROPIC_MODEL, max_tokens=16000, temperature=0.3,
                system=system_prompt,
                messages=[{"role": "user", "content": retry}]
            )
        txt = "".join(b.text for b in response.content if hasattr(b, "text")).strip()
        st.session_state["last_scene_design_raw"] = txt
        return safe_json_loads(txt)
    except Exception as e:
        st.error(f"Scene Design 실패: {e}")
        raw = st.session_state.get("last_scene_design_raw", "")
        if raw:
            with st.expander("🔧 Scene Design Raw (디버그)"):
                st.text_area("Raw", raw, height=400)
        return None


# ─── 16비트 구조 정의 ───
def _get_series_act_mapping(fmt):
    """미니시리즈의 에피소드를 3막에 매핑"""
    bs = P.get_beat_structure(fmt)
    eps = list(bs.keys())
    n = len(eps)
    if n <= 3:
        return {1: eps[:1], 2: eps[1:2], 3: eps[2:]}
    elif n <= 6:
        return {1: eps[:2], 2: eps[2:4], 3: eps[4:]}
    else:  # 8화
        return {1: eps[:3], 2: eps[3:6], 3: eps[6:]}


def _build_b_story_context(core_data):
    """Core Build의 세계관/GNS에서 B-Story 컨텍스트 추출"""
    wb = core_data.get("world_build", {})
    gns = core_data.get("goal_need_strategy", {})
    parts = []
    # 시간축
    time_ctx = wb.get("time", wb.get("시간", ""))
    if time_ctx:
        parts.append(f"B-Story 시간축: {time_ctx}")
    # 공간/사회
    space_ctx = wb.get("space", wb.get("공간", ""))
    if space_ctx:
        parts.append(f"B-Story 공간: {space_ctx}")
    # 권력구조
    power_ctx = wb.get("power_structure", wb.get("권력구조", ""))
    if power_ctx:
        parts.append(f"권력구조: {power_ctx}")
    # Risk (실패 시 B-Story 영향)
    risk = gns.get("risk", "")
    if risk:
        parts.append(f"실패 시 B-Story 결과: {risk}")
    return "\n".join(parts) if parts else ""


# ─── API Call: Treatment Build (16비트 줄글) ───
def call_treatment_beats(core_data, story_data, scene_data, genre, fmt, act_number, locked_block="",
                          fact_based=False, historical=False, film_type="", research=None):
    """Treatment Build: 막별 비트 줄글 생성 — 영화/시리즈 자동 분기.
    
    [v2.3.8 추가 전파 필드]
    Core Build에서 설계했지만 Treatment에서 휘발되던 4개 카테고리를 전파:
      1. project_intent (theme + pitch + tone_manner) — 주제·톤 일관성 유지
      2. world_build (time/space/rules) — 시대 고증·세계관 일관성 (생성 단계)
      3. genre_expectation_check.weak_zones — Core Build 자가 진단 경고
      4. research.real_events (상위 3~5개) + existing_works (표절 방지) — 핍진성
    """
    try:
        client = get_client()
        gns = core_data.get("goal_need_strategy", {})
        nd = core_data.get("narrative_drive", {})
        lp = core_data.get("logline_pack", {})
        chars = core_data.get("characters", []) + core_data.get("extended_characters", [])
        syn = story_data.get("synopsis_1p", {})
        storyline = story_data.get("storyline", [])
        
        # v2.3.8 신규 — 휘발되던 필드 4종 추출
        project_intent = core_data.get("project_intent", {})
        world_build = core_data.get("world_build", {})
        genre_check = core_data.get("genre_expectation_check", {})

        is_series = P.is_series_format(fmt)

        # ── 비트 목록 구성 ──
        if is_series:
            act_ep_map = _get_series_act_mapping(fmt)
            bs = P.get_beat_structure(fmt)
            all_beats = []
            for ep_key in act_ep_map.get(act_number, []):
                for beat_tuple in bs.get(ep_key, []):
                    all_beats.append((beat_tuple[0], f"[{ep_key}] {beat_tuple[1]}", beat_tuple[2] if len(beat_tuple) > 2 else ""))
            beats = all_beats
        else:
            bs = P.BEAT_STRUCTURE_FILM
            beats = bs[act_number]

        act_labels = {1: "1막", 2: "2막", 3: "3막"}
        act_label = act_labels[act_number]

        chars_simple = json.dumps(
            [{"name": c.get("name",""), "role": c.get("role",""), "goal": c.get("goal",""), "flaw": c.get("flaw","")} for c in chars],
            ensure_ascii=False
        )

        scene_ref = ""
        if scene_data:
            scenes = scene_data.get("key_scenes", [])
            if scenes:
                scene_ref = f"\n[Scene Design 참고]\n{json.dumps(scenes, ensure_ascii=False)}"

        beat_list = "\n".join([f"Beat {b[0]}. {b[1]}" + (f" — {b[2]}" if len(b) > 2 and b[2] else "") for b in beats])

        # ── B-Story 컨텍스트 구성 ──
        b_story_context = _build_b_story_context(core_data)

        # ── 매력 설계도 추출 ──
        attraction_design = core_data.get("attraction_design", {})
        attraction_block = ""
        if attraction_design:
            oh = attraction_design.get("opening_hook", {})
            tp = attraction_design.get("twist_point", {})
            wc = attraction_design.get("water_cooler_moment", {})
            ee = attraction_design.get("emotional_explosion", {})
            fd = attraction_design.get("forbidden_directions", [])
            attraction_block = f"""
[⚡ 매력 설계도 — 최우선 명령. 반드시 반영하라]
첫 장면 유형: {oh.get("type", "")}
첫 장면: {oh.get("description", "")}
배반 포인트: {tp.get("betrayal", "")} (관객 예상: {tp.get("expected_direction", "")})
Water Cooler Moment: {wc.get("scene_or_setup", "")}
감정 억압 → 폭발: {ee.get("suppression", "")} → {ee.get("explosion_moment", "")}
절대 금지 방향: {" / ".join(fd)}
"""

        # ── 오프닝 전략 블록 (v2.2 신규) — 1막일 때만 Beat 1 강제 규칙 적용 ──
        os_data = core_data.get("opening_strategy", {})
        opening_strategy_block = ""
        if os_data and act_number == 1:
            opening_strategy_block = f"""
[🎬 오프닝 전략 — Beat 1 절대 명령 (v2.2)]
오프닝 타입: {os_data.get("opening_type", "")}
오프닝 의도: {os_data.get("opening_intent", "")}
도파민 포인트: {os_data.get("dopamine_point", "")}
1막 연결 방식: {os_data.get("opening_to_act1_link", "")}
훅 라인/이미지: {os_data.get("opening_hook_line", "")}
장르 DNA 체크: {os_data.get("genre_dna_check", "")}

★ Beat 1 (1막 첫 비트) 작성 규칙 — 절대 지킬 것 ★
1. 비트의 첫 3줄 안에 오프닝 타입에 해당하는 장면이 즉시 시작되어야 한다.
   - Action Drop: 이미 진행 중인 행동으로 시작
   - Cold Open: 본편과 다른 시점/장면으로 시작
   - Tease & Reveal: 수수께끼 이미지로 시작
   - In Media Res: 미래의 결정적 순간으로 시작
   - Character Reveal Action: 주인공의 정의적 행동으로 시작
   - Hook Dialogue: 강렬한 한 줄 대사로 시작
2. '훅 라인/이미지'는 Beat 1 narrative에 구체적으로 실제 장면으로 등장해야 한다.
3. '도파민 포인트'에 명시된 감각(충격/웃음/긴장/경이/호기심/감정 울림)이 Beat 1 안에서 최소 1회 터져야 한다.
4. 절대 금지 AI 오프닝 습관:
   ❌ '서울. 2024년 봄.' 같은 배경 설명 시작
   ❌ '어느 날 아침, 주인공은 눈을 떴다.' 일상 시작
   ❌ '오늘도 ○○은 평소처럼...' 관습 서술
   ❌ 주인공의 외형·직업·가족관계를 첫 3줄에 나열
   ❌ 풍경 묘사로 시작해서 인물로 줌인하는 구조
5. Beat 1의 마지막 문장은 Beat 2로 넘어가는 연결 장치가 되어야 한다(오프닝→1막 연결 방식 준수).
"""

        # ── 시스템 프롬프트 (fmt + b_story 전달) ──
        # v2.4.0: 시대극 자동 감지 컨텍스트 추출
        _locked_text, _idea_text = _get_period_scan_context({
            "genre": genre,
            "core_data": core_data or {},
            "locked_items": [locked_block] if locked_block else [],
        })
        system_prompt = P.build_system_treatment(genre, act_label, fmt=fmt, b_story_context=b_story_context,
                                                   fact_based=fact_based, historical=historical, film_type=film_type,
                                                   locked_text=_locked_text, idea_text=_idea_text)

        # ── 시리즈용 추가 정보 ──
        series_info = ""
        if is_series:
            ep_list = act_ep_map.get(act_number, [])
            series_info = f"\n[에피소드 구성] 이 {act_label}은 {', '.join(ep_list)}에 해당합니다. 각 에피소드의 마지막 비트는 반드시 클리프행어로 끝내세요.\n"

        # ── BJND 씬 레벨 집행 블록 (v2.3 신규) — 모든 막에 주입 (막별 Cost 단계 규칙 내장) ──
        bjnd_scene_block = P.build_bjnd_scene_block(core_data)

        # ─────────────────────────────────────────────
        # v2.3.8 신규 전파 블록 4종 — Core Build에서 휘발되던 필드 주입
        # ─────────────────────────────────────────────
        
        # [1] 기획의도 + 주제 + 톤앤매너 (project_intent 전파)
        project_intent_block = ""
        if project_intent:
            theme = project_intent.get("theme", "")
            pitch = project_intent.get("pitch", "")
            tone_manner = project_intent.get("tone_manner", [])
            if theme or pitch or tone_manner:
                tone_str = " / ".join(tone_manner) if tone_manner else ""
                project_intent_block = f"""
[🎯 기획의도 — 모든 비트가 수렴해야 할 주제]
주제: {theme}
엘리베이터 피치: {pitch}
톤앤매너: {tone_str}

★ 매 비트 작성 시 이 주제에서 이탈하지 말 것. 대사·행동·디테일 모두 주제를 향해 수렴.
★ 톤앤매너 키워드를 씬 분위기·대사 결·디테일 선택에 일관되게 반영.
"""
        
        # [2] 세계관 (world_build 전파) — 시대/공간/규칙
        world_build_block = ""
        if world_build:
            wb_time = world_build.get("time", "")
            wb_space = world_build.get("space", "")
            wb_rules = world_build.get("rules", "")
            wb_taboo = world_build.get("taboo", "")
            wb_visual = world_build.get("visual_keywords", [])
            if any([wb_time, wb_space, wb_rules]):
                visual_str = " / ".join(wb_visual) if wb_visual else ""
                world_build_block = f"""
[🌍 세계관 — 시대·공간·규칙 절대 준수]
시간 배경: {wb_time}
공간 배경: {wb_space}
세계관 규칙: {wb_rules}
금기/터부: {wb_taboo}
시각 키워드: {visual_str}

★ 위 시간 배경에 존재하지 않는 기술·제도·문화·언어를 Treatment에 포함하지 말 것.
  (예: 1980년대 설정인데 스마트폰·SNS·AI, 조선시대인데 사진·기차 등)
★ 세계관 규칙은 비트 전체에 일관되게 적용.
★ 금기가 있다면 그것이 파괴되는 순간이 긴장의 축이 되어야 함.
"""
        
        # [3] 장르 자가 진단 경고 (genre_expectation_check.weak_zones 전파)
        genre_warning_block = ""
        if genre_check:
            weak_zones = genre_check.get("weak_zones", [])
            climax_verdict = genre_check.get("climax_verdict", "")
            genre_diag = genre_check.get("genre_fun_diagnosis", "")
            if weak_zones or climax_verdict == "CLIMAX_FAIL":
                weak_str = "\n".join(f"  · {w}" for w in weak_zones) if weak_zones else "  · (없음)"
                genre_warning_block = f"""
[⚠️ 장르 자가 진단 경고 — Core Build에서 식별된 약점]
장르 재미 진단: {genre_diag}
예상 약점 구간:
{weak_str}
클라이맥스 판정: {climax_verdict}

★ 위 약점 구간은 Core Build가 스스로 식별한 위험 영역. Treatment 집필 시 이 구간을 특히 강하게 설계하라.
★ CLIMAX_FAIL 판정이 있으면 3막 마지막 비트를 장르 약속에 맞게 재설계하라.
  (예: 로맨틱 코미디 CLIMAX_FAIL이면 → 부녀 화해/가족 서사 금지, 반드시 로맨스 완성 씬)
"""
        
        # [4] v2.3.8 재설계 — research_applied (Core Build 선별) 우선 + 원본 research 보조
        # Mr. MOON 원칙: "리서치 내용 전체를 받을 필요는 없고. 설정된 캐릭터에 사용할 수 있는 것만 응용"
        research_block_treatment = ""
        
        # (A) Core Build가 이미 선별한 research_applied를 최우선으로 전달
        research_applied = core_data.get("research_applied", {}) if core_data else {}
        if research_applied:
            refs = research_applied.get("references_used", [])
            anchors = research_applied.get("verisimilitude_anchors", [])
            absorption = research_applied.get("research_absorption_note", "")
            
            # references_used를 간결하게 포맷
            refs_lines = []
            for r in refs[:6]:
                src_title = r.get("source_title", "")
                applied_to = r.get("applied_to", "")
                how = r.get("how_applied", "")
                char_name = r.get("character_name", "")
                char_suffix = f" [→ {char_name}]" if char_name else ""
                refs_lines.append(f"  · [{applied_to}] {src_title}: {how}{char_suffix}")
            refs_str = "\n".join(refs_lines) if refs_lines else "  · (Core Build에서 선별한 참조 없음)"
            
            anchors_lines = [f"  {i+1}. {a}" for i, a in enumerate(anchors[:5])]
            anchors_str = "\n".join(anchors_lines) if anchors_lines else "  (없음)"
            
            if refs_lines or anchors_lines or absorption:
                research_block_treatment = f"""
[📚 리서치 선별 응용 — Core Build가 이미 이 작품에 녹인 것만 전달]
작가 흡수 노트: {absorption}

[이 작품에 응용된 리서치 요소 — 매 비트에서 유지]
{refs_str}

[핍진성 앵커 — Treatment 집필 시 반드시 유지할 현실 기반 디테일]
{anchors_str}

★ 위 요소들은 Core Build가 '이 작품의 캐릭터·설정·사건에 실제로 응용된 것'만 선별한 결과.
★ 응용된 요소는 해당 캐릭터·장면에서 구체 디테일로 드러나야 한다. 추상적 언급 금지.
★ 핍진성 앵커는 Writer Engine까지 전달될 핵심이므로 Treatment에서 휘발시키지 말 것.
★ 리서치에 있었으나 research_applied에 없는 요소는 이 작품에 해당 안 됨 — 끌어다 쓰지 말 것.
"""
        
        # (B) 원본 research는 '표절 회피 레퍼런스'로만 보조 사용 (상위 3개 existing_works만)
        if research and not research_applied:
            # research_applied가 없는 경우에만 원본 research 최소한 참고 (구버전 프로젝트 호환)
            existing_works = research.get("existing_works", [])
            existing_titles = []
            for w in existing_works[:5]:
                t = w.get("title", "")
                y = w.get("year", "")
                if t:
                    existing_titles.append(f"{t}({y})" if y else t)
            existing_str = ", ".join(existing_titles) if existing_titles else "(없음)"
            
            if existing_titles:
                research_block_treatment = f"""
[📚 리서치 — 표절 회피 참조]
유사 작품: {existing_str}
★ 위 작품들과 유사한 설정·구조·클라이맥스를 피하라. 이 작품만의 차별점을 유지.
※ Core Build의 research_applied가 비어있어 최소 참고만 제공. 다음 재생성 시 research_applied 작성 권장.
"""

        user_prompt = f"""[작품 정보]
로그라인: {lp.get("washed","")}
장르: {genre} / 포맷: {fmt}
Goal: {gns.get("goal","")} / Need: {gns.get("need","")} / Strategy: {gns.get("strategy","")}
서사동력: {nd.get("desire_origin","")}({nd.get("origin_detail","")}) → {nd.get("arc_direction","")} / 해결전략: {nd.get("resolution_strategy","")}
캐릭터: {chars_simple}
{locked_block}
{series_info}
{project_intent_block}
{world_build_block}
{genre_warning_block}
{research_block_treatment}
{attraction_block}
{opening_strategy_block}
{bjnd_scene_block}
[Synopsis]
{json.dumps(syn, ensure_ascii=False)}

[Storyline]
{json.dumps(storyline, ensure_ascii=False)}
{scene_ref}

[{act_label} — 작성할 비트]
{beat_list}

[JSON 스키마]
{P.TREATMENT_BEAT_SCHEMA_TEMPLATE.format(act_number=act_number)}

{P.TREATMENT_BEAT_RULES_TEMPLATE.format(beat_count=len(beats))}
"""

        response = client.messages.create(
            model=ANTHROPIC_MODEL_OPUS, max_tokens=16000, temperature=0.4,
            system=system_prompt,
            messages=[{"role": "user", "content": user_prompt}]
        )
        if response.stop_reason == "max_tokens":
            retry = user_prompt.replace("2500~4000자", "2000~3000자").replace("4000~6000자", "3000~4000자")
            response = client.messages.create(
                model=ANTHROPIC_MODEL_OPUS, max_tokens=16000, temperature=0.4,
                system=system_prompt,
                messages=[{"role": "user", "content": retry}]
            )
        txt = "".join(b.text for b in response.content if hasattr(b, "text")).strip()
        st.session_state[f"last_treatment_act{act_number}_raw"] = txt

        # JSON 파싱 시도 — 실패 시 자동 재시도 1회
        try:
            return safe_json_loads(txt)
        except json.JSONDecodeError:
            st.warning(f"⚠️ Treatment {act_label} JSON 파싱 실패 — 자동 재시도 중...")
            retry_system = system_prompt + "\n\n★ 최우선: 반드시 유효한 JSON만 출력. narrative 안에 쌍따옴표 절대 금지. 작은따옴표만. ★"
            response2 = client.messages.create(
                model=ANTHROPIC_MODEL_OPUS, max_tokens=16000, temperature=0.3,
                system=retry_system,
                messages=[{"role": "user", "content": user_prompt}]
            )
            txt2 = "".join(b.text for b in response2.content if hasattr(b, "text")).strip()
            st.session_state[f"last_treatment_act{act_number}_raw"] = txt2
            return safe_json_loads(txt2)

    except Exception as e:
        st.error(f"Treatment {act_label} 실패: {e}")
        raw = st.session_state.get(f"last_treatment_act{act_number}_raw", "")
        if raw:
            with st.expander(f"🔧 Treatment {act_label} Raw (디버그)"):
                st.text_area("Raw", raw, height=400)
        return None


def call_treatment_meta(act1, act2, act3, core_data):
    """Treatment 감정곡선 + 감독포인트 + 투자자요약"""
    try:
        client = get_client()
        all_beats = []
        for act_data in [act1, act2, act3]:
            if act_data:
                for b in act_data.get("beats", []):
                    all_beats.append(f"Beat {b.get('beat_no','')}. {b.get('beat_name','')}")

        # v2.1: 캐릭터 목록 추출 (이름 교차검증용)
        chars = core_data.get("characters", []) + core_data.get("extended_characters", [])
        char_registry = []
        for c in chars:
            name = c.get("name", "")
            role = c.get("role", "")
            age = c.get("age", "")
            if name:
                char_registry.append(f"{name}({age}, {role})")
        char_registry_str = " / ".join(char_registry) if char_registry else "(캐릭터 정보 없음)"

        response = client.messages.create(
            model=ANTHROPIC_MODEL, max_tokens=2000, temperature=0.3,
            system=P.SYSTEM_TREATMENT_META,
            messages=[{"role": "user", "content": f"""[트리트먼트 비트 목록]
{json.dumps(all_beats, ensure_ascii=False)}

[로그라인] {core_data.get("logline_pack",{}).get("washed","")}

[★ 등장 캐릭터 레지스트리 — 감독 포인트에서 이름을 쓸 때 반드시 이 목록의 이름만 사용할 것]
{char_registry_str}

[JSON 스키마]
{{
  "emotion_curve": [
    {{"point": "Beat 라벨", "tension": 0, "emotion": "감정"}}
  ],
  "director_notes": ["감독용 포인트 1", "감독용 포인트 2", "감독용 포인트 3"],
  "investor_summary": "투자자용 요약 3~4문장"
}}

규칙:
- emotion_curve 16개 포인트. tension 1~10. director_notes 3개.
- ★ director_notes에 캐릭터 이름을 언급할 때 반드시 위 [등장 캐릭터 레지스트리]의 이름만 사용하라.
- ★ 캐릭터의 역할(주인공/적대자/조력자 등)과 나이에 맞지 않는 연출 지시를 절대 하지 마라.
  예: 9세 아이에게 '구조적 폭탄', '위협감', '주도권' 같은 어른 연출 지시를 붙이면 안 된다.
- ★ 감독 포인트에서 특정 비트의 연출을 지시할 때는 해당 비트의 실제 주체(누가 그 행동을 하는가)와
  일치하는 이름을 써라. 비트 이름에서 임의로 인물명을 추출하면 안 된다.
- ★ 클라이맥스 질문의 주체가 누구인지(주인공인지, 적대자인지, 조력자인지) 반드시 확인한 후 이름을 지정하라.
"""}]
        )
        txt = "".join(b.text for b in response.content if hasattr(b, "text")).strip()
        return safe_json_loads(txt)
    except Exception as e:
        st.error(f"Treatment 메타 생성 실패: {e}")
        return None


def call_treatment_gate(treatment_data, core_data=None, genre=""):
    """Gate E: Treatment Gate 채점 — v2.3.6 강화판.
    
    이전 문제: beat_summary에 글자 수(chars)만 전달. 실제 내용 미검증.
    v2.3.6: 각 비트의 narrative 내용 요약 + 장르 + Core Build 참조를 전달하여
           장르 기대 충족, 엔딩 정합, 논리 일관성을 실제로 검증.
    """
    try:
        client = get_client()
        
        # v2.3.6: 실제 narrative 내용도 전달 (압축 요약)
        beat_details = []
        for act_key in ["act1", "act2", "act3"]:
            act_data = treatment_data.get(act_key, {})
            if act_data:
                for b in act_data.get("beats", []):
                    narrative = b.get("narrative", "")
                    # 각 비트 narrative를 500자로 압축 (앞 400자 + 뒤 100자)
                    if len(narrative) > 500:
                        narrative_compact = narrative[:400] + " ... [중략] ... " + narrative[-100:]
                    else:
                        narrative_compact = narrative
                    beat_details.append({
                        "beat_no": b.get("beat_no"),
                        "beat_name": b.get("beat_name"),
                        "length": len(narrative),
                        "narrative_summary": narrative_compact,
                    })
        
        # v2.3.6: Core Build 핵심 정보 전달 (엔딩 정합 검증용)
        # v2.3.7: world_build.time 등 시대 고증 검증용 정보 추가
        core_context = ""
        if core_data:
            gns = core_data.get("goal_need_strategy", {})
            lp = core_data.get("logline_pack", {})
            wb = core_data.get("world_build", {})
            ending_payoff = gns.get("ending_payoff", "")
            strategy = gns.get("strategy", "")
            
            # 시대/공간 정보 추출
            time_setting = wb.get("time", "")
            space_setting = wb.get("space", "")
            rules = wb.get("rules", "")
            
            core_context = f"""
[Core Build 참조 — 이 Treatment가 일관되게 구현해야 할 원본 설계]
장르: {genre}
로그라인: {lp.get("washed", "")}
주인공 Strategy: {strategy}
Ending Payoff: {ending_payoff}

[★ 세계관·시대 설정 — 시대 고증 검증용 ★]
시대/시간: {time_setting}
공간/장소: {space_setting}
세계관 규칙: {rules}

위 시대에 존재하지 않는 요소가 Treatment에 포함되었는지 반드시 확인하라.
(예: 1980년대 설정인데 스마트폰/SNS/AI, 1970년대인데 PC, 조선시대인데 사진기 등)
"""

        response = client.messages.create(
            model=ANTHROPIC_MODEL, max_tokens=2500, temperature=0.2,
            system=P.SYSTEM_TREATMENT_GATE,
            messages=[{"role": "user", "content": f"""{core_context}
[Treatment 비트 구성 — 실제 내용 요약]
{json.dumps(beat_details, ensure_ascii=False)}

총 {len(beat_details)}개 비트.

[JSON 스키마]
{{
  "gate_e_treatment": {{
    "cinematic_reading": 0.0,
    "scene_emotion_match": 0.0,
    "beat_completeness": 0.0,
    "screenplay_ready": 0.0,
    "genre_expectation_alignment": 0.0,
    "ending_coherence": 0.0,
    "logic_consistency": 0.0,
    "period_consistency": 0.0,
    "average": 0.0,
    "genre_check_feedback": "장르 체크리스트 기준 충족 여부 1~2문장",
    "ending_check_feedback": "Core Build의 Ending Payoff와 실제 엔딩 비트의 정합성 1문장",
    "logic_check_feedback": "논리 비약/세계관 일관성 이슈가 있으면 구체 비트 번호와 함께 1~2문장",
    "period_check_feedback": "시대 고증 오류(작품 시대에 없던 기술/제도/문화 등) 구체 비트 번호와 함께 1~2문장 (없으면 '이슈 없음')",
    "feedback": "Gate E 종합 피드백 1문장",
    "critical_issues": ["문제가 있는 비트 번호와 짧은 진단 리스트 (없으면 빈 배열)"]
  }}
}}
규칙: average = 8항목 평균 (소수점 1자리).
특히 genre_expectation_alignment, ending_coherence, logic_consistency, period_consistency는 
Core Build 설계와 Treatment 실제 구현의 일치도를 엄격 평가."""}]
        )
        txt = "".join(b.text for b in response.content if hasattr(b, "text")).strip()
        return safe_json_loads(txt)
    except Exception as e:
        st.error(f"Gate E 실패: {e}")
        return None


# ─── API Call: Tone Document ───
def call_tone_document(core_data, structure_data, scene_data, treatment_data, char_bible, genre, fmt, locked_block="",
                        fact_based=False, historical=False, film_type=""):
    """톤 & 연출 문서 — Writer Engine의 스타일 가이드"""
    try:
        client = get_client()
        lp = core_data.get("logline_pack", {})
        gns = core_data.get("goal_need_strategy", {})
        nd = core_data.get("narrative_drive", {})
        wb = core_data.get("world_build", {})

        # 캐릭터 이름/역할만 추출
        char_names = []
        if char_bible:
            for c in char_bible.get("characters", []):
                char_names.append(f"{c.get('name','')}({c.get('role','')})")
        char_str = ", ".join(char_names)

        # Treatment 요약
        treatment_summary = ""
        for act_key in ["act1", "act2", "act3"]:
            act = treatment_data.get(act_key, {})
            if act:
                for b in act.get("beats", []):
                    treatment_summary += f"Beat {b.get('beat_no','')}: {b.get('beat_name','')}. "

        # v2.4.0: 시대극 자동 감지 컨텍스트 추출
        _locked_text, _idea_text = _get_period_scan_context({
            "genre": genre,
            "core_data": core_data or {},
            "locked_items": [locked_block] if locked_block else [],
        })
        system_prompt = P.build_system_tone_document(genre, fmt,
                                                       fact_based=fact_based, historical=historical, film_type=film_type,
                                                       locked_text=_locked_text, idea_text=_idea_text)

        user_prompt = f"""[로그라인]
{lp.get("washed","")}

[캐릭터]
{char_str}

[세계관]
{json.dumps(wb, ensure_ascii=False)}

[Treatment 비트 구성]
{treatment_summary}
{locked_block}

[JSON 스키마]
{P.TONE_DOC_SCHEMA}"""

        response = client.messages.create(
            model=ANTHROPIC_MODEL_OPUS, max_tokens=16000, temperature=0.3,
            system=system_prompt,
            messages=[{"role": "user", "content": user_prompt}]
        )
        if response.stop_reason == "max_tokens":
            # 분량 줄여서 재시도
            retry_prompt = user_prompt.replace(
                "카메라 철학 2~3문장", "카메라 철학 1~2문장"
            ).replace(
                "이 작품에서 절대 하지 말아야 할 연출/톤/대사 규칙 5개",
                "이 작품에서 절대 하지 말아야 할 규칙 3개"
            )
            response = client.messages.create(
                model=ANTHROPIC_MODEL_OPUS, max_tokens=16000, temperature=0.3,
                system=system_prompt,
                messages=[{"role": "user", "content": retry_prompt}]
            )
        txt = "".join(b.text for b in response.content if hasattr(b, "text")).strip()
        st.session_state["last_tone_doc_raw"] = txt

        # JSON 파싱 시도 — 실패 시 자동 재시도 1회
        try:
            return safe_json_loads(txt)
        except json.JSONDecodeError:
            st.warning("⚠️ Tone Document JSON 파싱 실패 — 자동 재시도 중...")
            retry_system = system_prompt + "\n\n★ 최우선: 반드시 유효한 JSON만 출력. 따옴표는 작은따옴표만. 줄바꿈·역슬래시 금지. ★"
            response2 = client.messages.create(
                model=ANTHROPIC_MODEL_OPUS, max_tokens=16000, temperature=0.2,
                system=retry_system,
                messages=[{"role": "user", "content": user_prompt}]
            )
            txt2 = "".join(b.text for b in response2.content if hasattr(b, "text")).strip()
            st.session_state["last_tone_doc_raw"] = txt2
            return safe_json_loads(txt2)

    except Exception as e:
        st.error(f"Tone Document 생성 실패: {e}")
        raw = st.session_state.get("last_tone_doc_raw", "")
        if raw:
            with st.expander("🔧 Tone Document Raw (디버그)"):
                st.text_area("Raw", raw, height=400)
        return None


# ─── DOCX 생성 ───
def generate_docx(project):
    """기획개발보고서 DOCX 생성 — BLUE JEANS 기획서 스타일"""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
    import io

    NAVY = RGBColor(0x19, 0x19, 0x70)
    YELLOW = RGBColor(0xFF, 0xCB, 0x05)
    DIM = RGBColor(0x88, 0x88, 0x99)
    BLACK = RGBColor(0x1A, 0x1A, 0x2E)
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)

    doc = Document()

    # ── 기본 스타일 ──
    style = doc.styles['Normal']
    style.font.name = 'Pretendard'
    style.font.size = Pt(10)
    style.font.color.rgb = BLACK
    style.paragraph_format.line_spacing = 1.4
    style.paragraph_format.space_after = Pt(4)

    # 페이지 여백
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)

    # ── 헬퍼 함수 ──
    def add_yellow_header(kr_text, en_text=""):
        """노란 하이라이트 섹션 헤더 (한글 + ENGLISH 병기)"""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(10)
        # 노란 배경 셰이딩
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="FFCB05" w:val="clear"/>')
        p.paragraph_format.element.get_or_add_pPr().append(shading)
        # 한글
        run_kr = p.add_run(f"  {kr_text}")
        run_kr.font.size = Pt(13)
        run_kr.font.bold = True
        run_kr.font.color.rgb = NAVY
        run_kr.font.name = 'Pretendard'
        # ENGLISH
        if en_text:
            run_en = p.add_run(f"  {en_text}")
            run_en.font.size = Pt(9)
            run_en.font.bold = True
            run_en.font.color.rgb = NAVY
            run_en.font.name = 'Pretendard'
        return p

    def add_sub_header(text):
        """서브 헤더 (네이비 좌측 볼드)"""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = NAVY
        return p

    def add_body(text, size=10):
        """본문 텍스트"""
        if not text:
            return
        p = doc.add_paragraph()
        run = p.add_run(str(text))
        run.font.size = Pt(size)
        run.font.color.rgb = BLACK
        p.paragraph_format.line_spacing = 1.5
        return p

    def add_labeled(label, text, bold_label=True):
        """라벨 + 텍스트"""
        if not text:
            return
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        run_l = p.add_run(f"[{label}]  ")
        run_l.font.size = Pt(9)
        run_l.font.bold = bold_label
        run_l.font.color.rgb = NAVY
        run_t = p.add_run(str(text))
        run_t.font.size = Pt(10)
        run_t.font.color.rgb = BLACK
        return p

    def add_quote(character, dialogue):
        """캐릭터 대사 인용구"""
        if not dialogue:
            return
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        # 왼쪽 보더 효과: 인용부호로 대체
        run = p.add_run(f'"{dialogue}"')
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = NAVY
        if character:
            run2 = p.add_run(f"\n— {character}")
            run2.font.size = Pt(9)
            run2.font.color.rgb = DIM

    def add_spacer(height=6):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(height)
        p.paragraph_format.space_after = Pt(0)

    # ═══════════════════════════════════
    #  COVER PAGE
    # ═══════════════════════════════════
    add_spacer(60)

    # "작품 기획안"
    p0 = doc.add_paragraph()
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p0.add_run("작 품 기 획 안")
    r0.font.size = Pt(11)
    r0.font.color.rgb = DIM
    r0.font.name = 'Pretendard'

    add_spacer(12)

    # 작품 제목 (대형)
    title_text = project.get("title", "제목 없음")
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(title_text)
    r1.font.size = Pt(32)
    r1.font.bold = True
    r1.font.color.rgb = NAVY
    r1.font.name = 'Pretendard'

    # 장르 · 타겟 · 포맷
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_text = f"{project.get('genre','')}  ·  {project.get('target_market','')}  ·  {project.get('format','')}"
    r2 = p2.add_run(meta_text)
    r2.font.size = Pt(10)
    r2.font.color.rgb = DIM

    add_spacer(40)

    # 노란 구분선
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_line = p_line.add_run("━" * 30)
    r_line.font.size = Pt(10)
    r_line.font.color.rgb = YELLOW

    add_spacer(12)

    # 로그라인 미리보기 (있으면)
    core = project.get("core", {})
    lp = core.get("logline_pack", {}) if core else {}
    washed = lp.get("washed", "")
    if washed:
        p_log = doc.add_paragraph()
        p_log.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_log = p_log.add_run(washed)
        r_log.font.size = Pt(10)
        r_log.font.italic = True
        r_log.font.color.rgb = BLACK

    add_spacer(50)

    # 기획/제작 크레딧
    p_credit = doc.add_paragraph()
    p_credit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_c1 = p_credit.add_run("기획 · 공동제작\n")
    r_c1.font.size = Pt(9)
    r_c1.font.color.rgb = DIM
    r_c2 = p_credit.add_run("BLUE JEANS PICTURES")
    r_c2.font.size = Pt(14)
    r_c2.font.bold = True
    r_c2.font.color.rgb = NAVY
    r_c2.font.name = 'Pretendard'

    doc.add_page_break()

    # ═══════════════════════════════════
    #  SECTION 1: 로그라인 LOGLINE
    # ═══════════════════════════════════
    if core:
        add_yellow_header("로그라인", "LOGLINE")
        for label, key in [("Original","original"),("Washed","washed"),("투자자용","investor"),("감독용","director"),("캐릭터 훅","character_hook")]:
            val = lp.get(key, "")
            if val:
                add_labeled(label, val)

        # ── 기획의도 KEY POINTS ──
        add_yellow_header("기획의도", "KEY POINTS")
        pi = core.get("project_intent", {})
        add_labeled("소재", pi.get("subject", ""))
        add_labeled("장르", pi.get("genre_approach", ""))
        add_labeled("시장", pi.get("market_rationale", ""))
        add_labeled("Pitch", pi.get("pitch", ""))
        add_labeled("Theme", pi.get("theme", ""))

        # ── G/N/S ──
        add_yellow_header("드라마 구조", "GOAL / NEED / STRATEGY")
        gns = core.get("goal_need_strategy", {})
        add_labeled("Goal", gns.get("goal", ""))
        add_labeled("Need", gns.get("need", ""))
        add_labeled("Strategy", gns.get("strategy", ""))
        add_labeled("Risk", gns.get("risk", ""))
        add_labeled("Ending Payoff", gns.get("ending_payoff", ""))

        # ── 서사동력 NARRATIVE DRIVE ──
        nd = core.get("narrative_drive", {})
        if nd:
            add_spacer(6)
            add_sub_header("▎ 서사동력  ·  Narrative Drive")
            origin = nd.get("desire_origin", "")
            origin_kr = "상실(Loss)" if origin == "loss" else "결핍(Lack)" if origin == "lack" else origin
            add_labeled("발생요인", f"{origin_kr} — {nd.get('origin_detail', '')}")
            add_labeled("아크 방향", nd.get("arc_direction", ""))
            add_labeled("해결전략", nd.get("resolution_strategy", ""))
            add_labeled("Goal↔Need 간극", nd.get("goal_need_gap", ""))

        # ── 오프닝 전략 OPENING STRATEGY (v2.2 신규) ──
        os_data = core.get("opening_strategy", {})
        if os_data:
            add_yellow_header("오프닝 전략", "OPENING STRATEGY")
            add_labeled("오프닝 타입", os_data.get("opening_type", ""))
            add_labeled("오프닝 의도", os_data.get("opening_intent", ""))
            add_labeled("도파민 포인트", os_data.get("dopamine_point", ""))
            add_labeled("1막 연결 방식", os_data.get("opening_to_act1_link", ""))
            add_labeled("훅 라인/이미지", os_data.get("opening_hook_line", ""))
            add_labeled("장르 DNA 체크", os_data.get("genre_dna_check", ""))

        # ── 세계관 WORLD ──
        add_yellow_header("세계관", "WORLD BUILDING")
        wb = core.get("world_build", {})
        for label, key in [("시간","time"),("공간","space"),("규칙","rules"),("금기","taboo"),("권력구조","power_structure")]:
            add_labeled(label, wb.get(key, ""))

        doc.add_page_break()

        # ── 캐릭터 CHARACTER ──
        add_yellow_header("캐릭터", "CHARACTER")
        role_labels = {"protagonist":"주인공","antagonist":"적대자","ally":"조력자","mirror":"거울","catalyst":"촉매자","subplot_lead":"서브플롯 리드"}
        all_core_chars = core.get("characters", []) + core.get("extended_characters", [])
        for ch in all_core_chars:
            role = role_labels.get(ch.get("role",""), ch.get("role",""))
            name = ch.get("name", "")

            # 캐릭터 이름 헤더
            add_sub_header(f"▎ {role}  ·  {name}")

            add_body(ch.get("description", ""))
            add_labeled("욕망", ch.get("goal", ""))
            add_labeled("결핍", ch.get("need", ch.get("flaw", "")))
            add_labeled("대사톤", ch.get("dialogue_tone", ""))

            # 대사 인용구 (있으면)
            key_dialogue = ch.get("key_dialogue", ch.get("signature_line", ""))
            if key_dialogue:
                add_quote(name, key_dialogue)

            add_spacer(8)

    # ── Character Bible (확장) ──
    char_bible = project.get("char_bible", {})
    if char_bible:
        doc.add_page_break()
        add_yellow_header("캐릭터 바이블", "CHARACTER BIBLE")

        for ch in char_bible.get("characters", []):
            role_kr = {"protagonist":"주인공","antagonist":"적대자","ally":"조력자","mirror":"거울","catalyst":"촉매자","subplot_lead":"서브플롯 리드"}
            role = role_kr.get(ch.get("role",""), ch.get("role",""))
            name = ch.get("name","")
            age = ch.get("age","")

            add_sub_header(f"▎ {role} · {name} ({age})")

            add_labeled("외형·첫인상", ch.get("appearance",""))
            add_labeled("직업/위치", ch.get("occupation",""))
            add_body(ch.get("backstory",""))

            add_labeled("비밀", ch.get("secret",""))
            add_labeled("신념", ch.get("belief",""))
            add_labeled("두려움", ch.get("fear",""))

            habits = ch.get("habits", [])
            if habits:
                h_str = " · ".join(habits) if isinstance(habits, list) else str(habits)
                add_labeled("반복 습관", h_str)

            sp = ch.get("speech_pattern", [])
            if sp:
                sp_list = sp if isinstance(sp, list) else [sp]
                for i, s in enumerate(sp_list, 1):
                    add_labeled(f"말투 {i}", s)

            sl = ch.get("sample_lines", {})
            if sl:
                sl_labels = {"normal":"평상시","angry":"분노","vulnerable":"취약"}
                for k, v in sl.items():
                    label = sl_labels.get(k, k)
                    add_quote(f"{name} — {label}", v)

            ra = ch.get("relationship_attitudes", [])
            if ra:
                ra_list = ra if isinstance(ra, list) else [ra]
                for r in ra_list:
                    add_labeled("관계", r)

            arc = ch.get("arc_detail", {})
            if arc:
                add_labeled("1막 끝", arc.get("act1_end",""))
                add_labeled("미드포인트", arc.get("midpoint",""))
                add_labeled("클라이맥스", arc.get("climax",""))

            add_spacer(10)

    doc.add_page_break()

    # ═══════════════════════════════════
    #  SECTION 2: 시놉시스 SYNOPSIS
    # ═══════════════════════════════════
    story = project.get("structure_story", {})
    diag = project.get("structure_diag", {})

    if story:
        add_yellow_header("시놉시스", "SYNOPSIS")
        syn = story.get("synopsis_1p", {})
        for label, key in [("시작","opening"),("촉발사건","catalyst"),("전개","development"),("미드포인트","midpoint"),("붕괴","collapse"),("결전","climax"),("결말","ending")]:
            add_labeled(label, syn.get(key, ""))

        # 줄글 시놉시스
        prose = project.get("structure_prose", {})
        if prose and prose.get("prose"):
            add_spacer(6)
            add_sub_header("기승전결 시놉시스")
            add_body(prose["prose"])

        doc.add_page_break()

        # ── 스토리라인 STORYLINE ──
        add_yellow_header("스토리라인", "STORYLINE")
        for seq in story.get("storyline", []):
            seq_no = seq.get("seq", "")
            label = seq.get("label", "")
            add_sub_header(f"SEQ {seq_no}  ·  {label}")
            add_body(seq.get("summary", ""))

    if diag:
        doc.add_page_break()

        # ── 3막 구조 진단 ──
        add_yellow_header("3막 구조 진단", "THREE-ACT STRUCTURE")
        ta = diag.get("three_act", {})
        for label, key in [("1막 끝","act1_end"),("미드포인트","act2_midpoint"),("All Is Lost","act2_end"),("클라이맥스","act3_climax")]:
            add_labeled(label, ta.get(key, ""))

        add_spacer(6)
        add_sub_header("15-Beat Sheet")
        for bt in diag.get("beat_sheet", []):
            beat = bt.get("beat", "")
            status = bt.get("status", "")
            note = bt.get("note", "")
            status_mark = "✓" if status in ["있음","O","✓"] else "△" if status in ["약함","△"] else "✗"
            add_labeled(beat, f"{status_mark} {note}")

    doc.add_page_break()

    # ═══════════════════════════════════
    #  SECTION 3: 장면 설계 SCENE DESIGN
    # ═══════════════════════════════════
    scene_design = project.get("scene_design", {})
    if scene_design:
        add_yellow_header("장면 설계", "SCENE DESIGN")
        sms = scene_design.get("scene_map_summary", {})
        if sms.get("must_see_scenes"):
            add_labeled("Must-See 장면", sms["must_see_scenes"])

        add_spacer(6)

        for sc in scene_design.get("key_scenes", []):
            scene_no = sc.get("scene_no", "")
            title = sc.get("title", "")
            add_sub_header(f"S#{scene_no}  {title}")

            add_labeled("장소", sc.get("location", ""))
            add_labeled("인물", sc.get("characters", ""))
            add_labeled("상황", sc.get("setup", ""))
            add_labeled("행동 (Show!)", sc.get("dramatic_action", ""))

            tp = sc.get("turning_point", "")
            if tp:
                add_labeled("전환", tp)
            di = sc.get("dramatic_irony", "")
            if di:
                add_labeled("극적 아이러니", di)

            add_labeled("감정 변화", sc.get("emotion_shift", ""))
            add_labeled("시각 연출", sc.get("visual_direction", ""))
            add_labeled("판돈", sc.get("stakes", ""))

            kl = sc.get("key_line", "")
            if kl:
                # 핵심 대사는 인용구 스타일
                char_name = ""
                if ":" in kl:
                    char_name, dialogue = kl.split(":", 1)
                    add_quote(char_name.strip(), dialogue.strip())
                else:
                    add_quote("", kl)

            add_spacer(4)

        doc.add_page_break()

    # ═══════════════════════════════════
    #  SECTION 4: 트리트먼트 TREATMENT
    # ═══════════════════════════════════
    treatment = project.get("treatment", {})
    if treatment:
        add_yellow_header("트리트먼트", "TREATMENT")

        p_info = doc.add_paragraph()
        r_info = p_info.add_run("16비트 줄글 트리트먼트  ·  서술체 현재형  ·  대사 포함")
        r_info.font.size = Pt(9)
        r_info.font.color.rgb = DIM
        r_info.font.italic = True

        act_headers = {
            1: ("1막 — 설정", "ACT 1: SET-UP"),
            2: ("2막 — 대결", "ACT 2: CONFRONTATION"),
            3: ("3막 — 해결", "ACT 3: RESOLUTION"),
        }

        for act_num in [1, 2, 3]:
            act_data = treatment.get(f"act{act_num}")
            if act_data:
                kr, en = act_headers[act_num]
                add_spacer(10)
                # 막 헤더 (네이비 배경)
                p_act = doc.add_paragraph()
                p_act.paragraph_format.space_before = Pt(14)
                p_act.paragraph_format.space_after = Pt(8)
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="191970" w:val="clear"/>')
                p_act.paragraph_format.element.get_or_add_pPr().append(shading)
                r_act = p_act.add_run(f"  {kr}  ·  {en}")
                r_act.font.size = Pt(12)
                r_act.font.bold = True
                r_act.font.color.rgb = WHITE

                for b in act_data.get("beats", []):
                    beat_no = b.get("beat_no", "")
                    beat_name = b.get("beat_name", "")
                    narrative = b.get("narrative", "")
                    episode = b.get("episode", "")
                    event_s = b.get("event_summary", "")
                    decision_s = b.get("decision_summary", "")
                    consequence_s = b.get("consequence_summary", "")
                    b_story = b.get("b_story_beat", "")
                    cliff = b.get("cliffhanger", "")
                    villain = b.get("villain_beat", "")
                    pp = b.get("plant_payoff", "")

                    # 비트 번호 + 에피소드 태그 + 이름
                    p_beat = doc.add_paragraph()
                    p_beat.paragraph_format.space_before = Pt(12)
                    p_beat.paragraph_format.space_after = Pt(4)
                    if episode:
                        r_ep = p_beat.add_run(f"[{episode}] ")
                        r_ep.font.size = Pt(9)
                        r_ep.font.bold = True
                        r_ep.font.color.rgb = NAVY
                    r_num = p_beat.add_run(f"Beat {beat_no}")
                    r_num.font.size = Pt(10)
                    r_num.font.bold = True
                    r_num.font.color.rgb = YELLOW
                    r_name = p_beat.add_run(f"  {beat_name}")
                    r_name.font.size = Pt(10)
                    r_name.font.bold = True
                    r_name.font.color.rgb = NAVY

                    # 사건/선택/결과 요약 블록
                    meta_parts = []
                    if event_s:
                        meta_parts.append(f"사건: {event_s}")
                    if decision_s:
                        meta_parts.append(f"선택: {decision_s}")
                    if consequence_s:
                        meta_parts.append(f"결과: {consequence_s}")
                    if b_story:
                        meta_parts.append(f"B-Story: {b_story}")
                    if villain:
                        meta_parts.append(f"빌런: {villain}")
                    if pp:
                        meta_parts.append(f"Plant/Payoff: {pp}")
                    if cliff:
                        meta_parts.append(f"CLIFFHANGER: {cliff}")
                    if meta_parts:
                        p_meta = doc.add_paragraph()
                        p_meta.paragraph_format.space_before = Pt(2)
                        p_meta.paragraph_format.space_after = Pt(4)
                        shading_meta = parse_xml(f'<w:shd {nsdecls("w")} w:fill="EEEEF6" w:val="clear"/>')
                        p_meta.paragraph_format.element.get_or_add_pPr().append(shading_meta)
                        r_meta = p_meta.add_run(" | ".join(meta_parts))
                        r_meta.font.size = Pt(8)
                        r_meta.font.color.rgb = RGBColor(0x4A, 0x4A, 0x5A)
                        r_meta.font.italic = True

                    # 줄글 narrative
                    if narrative:
                        p_n = doc.add_paragraph()
                        p_n.paragraph_format.line_spacing = 1.6
                        p_n.paragraph_format.first_line_indent = Cm(0.5)
                        r_n = p_n.add_run(narrative)
                        r_n.font.size = Pt(10)
                        r_n.font.color.rgb = BLACK

        # 투자자용 요약
        meta = treatment.get("meta", {})
        if meta.get("investor_summary"):
            doc.add_page_break()
            add_yellow_header("투자자용 요약", "INVESTOR SUMMARY")
            add_body(meta["investor_summary"])

        # 감독 포인트
        if meta.get("director_notes"):
            add_spacer(8)
            add_sub_header("감독 포인트")
            for i, note in enumerate(meta["director_notes"], 1):
                add_labeled(f"Point {i}", note)

    doc.add_page_break()

    # ═══════════════════════════════════
    #  SECTION 5: 점수 DEVELOPMENT SCORE
    # ═══════════════════════════════════
    add_yellow_header("개발 적합도", "DEVELOPMENT FIT SCORE")

    cg = project.get("core_gate", {})
    fa = cg.get("five_axis_scores", {})
    if fa:
        final = fa.get("final_score", "")
        verdict = fa.get("verdict", "")

        # 큰 점수
        p_score = doc.add_paragraph()
        p_score.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_score.paragraph_format.space_before = Pt(12)
        r_score = p_score.add_run(str(final))
        r_score.font.size = Pt(36)
        r_score.font.bold = True
        r_score.font.color.rgb = NAVY

        p_verdict = doc.add_paragraph()
        p_verdict.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_v = p_verdict.add_run(str(verdict))
        r_v.font.size = Pt(11)
        r_v.font.color.rgb = DIM

        # 5축 상세 (테이블)
        axes = ["originality","market_fit","character","structure","theme"]
        axis_kr = {"originality":"독창성","market_fit":"시장성","character":"캐릭터","structure":"구조","theme":"테마"}
        scores_exist = any(fa.get(a) for a in axes)
        if scores_exist:
            add_spacer(8)
            table = doc.add_table(rows=1, cols=3)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr = table.rows[0].cells
            for i, text in enumerate(["항목", "점수", "상태"]):
                hdr[i].text = text
                for p in hdr[i].paragraphs:
                    for r in p.runs:
                        r.font.bold = True
                        r.font.size = Pt(9)
                        r.font.color.rgb = NAVY
            for a in axes:
                sc = fa.get(a, "")
                if sc:
                    row = table.add_row().cells
                    row[0].text = axis_kr.get(a, a)
                    row[1].text = str(sc)
                    sc_val = float(sc) if sc else 0
                    row[2].text = "●" if sc_val >= 7 else "◐" if sc_val >= 5 else "○"

    # ═══════════════════════════════════
    #  SECTION 6: 톤 문서 TONE DOCUMENT
    # ═══════════════════════════════════
    tone_doc = project.get("tone_doc", {})
    if tone_doc:
        doc.add_page_break()
        add_yellow_header("톤 & 연출 문서", "TONE DOCUMENT")

        vs = tone_doc.get("visual_style", {})
        if vs:
            add_sub_header("비주얼 스타일")
            for label, key in [("카메라 철학","camera_philosophy"),("색감 팔레트","color_palette"),("조명 규칙","lighting_rule"),("시그니처 쇼트","signature_shot")]:
                add_labeled(label, vs.get(key, ""))

        pc = tone_doc.get("pacing", {})
        if pc:
            add_sub_header("페이싱")
            for label, key in [("전체","overall"),("1막","act1_tempo"),("2막","act2_tempo"),("3막","act3_tempo"),("대사 비율","dialogue_density")]:
                add_labeled(label, pc.get(key, ""))

        dr = tone_doc.get("dialogue_rules", {})
        if dr:
            add_sub_header("대사 규칙")
            for label, key in [("전체 톤","overall_tone"),("서브텍스트","subtext_rule"),("침묵 활용","silence_usage")]:
                add_labeled(label, dr.get(key, ""))
            fp = dr.get("forbidden_phrases", [])
            if fp:
                fp_list = fp if isinstance(fp, list) else [fp]
                for f in fp_list:
                    add_labeled("금지", f)

        mt = tone_doc.get("motifs", {})
        if mt:
            add_sub_header("모티프")
            for label, key in [("반복 소품","recurring_objects"),("반복 장소","recurring_locations")]:
                val = mt.get(key, [])
                if val:
                    items = val if isinstance(val, list) else [val]
                    add_labeled(label, " · ".join(items))
            add_labeled("날씨/계절", mt.get("weather_mood",""))

        forbidden = tone_doc.get("forbidden", [])
        if forbidden:
            add_sub_header("금기 사항")
            fb_list = forbidden if isinstance(forbidden, list) else [forbidden]
            for f in fb_list:
                add_labeled("🚫", f)

        refs = tone_doc.get("reference_films", [])
        if refs:
            add_sub_header("참고 작품")
            for ref in refs:
                if isinstance(ref, dict):
                    add_labeled(ref.get("title",""), ref.get("reason",""))
                else:
                    add_body(str(ref))

        wi = tone_doc.get("writer_instruction", "")
        if wi:
            add_spacer(8)
            add_sub_header("Writer Engine 최종 지시")
            add_body(wi)

    # ═══════════════════════════════════
    #  FOOTER
    # ═══════════════════════════════════
    doc.add_paragraph("")
    p_line2 = doc.add_paragraph()
    p_line2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_line2 = p_line2.add_run("━" * 20)
    r_line2.font.size = Pt(8)
    r_line2.font.color.rgb = YELLOW

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_f1 = footer.add_run("© 2026 BLUE JEANS PICTURES\n")
    r_f1.font.size = Pt(8)
    r_f1.font.color.rgb = DIM
    r_f2 = footer.add_run(f"Creator Engine {ENGINE_VERSION}")
    r_f2.font.size = Pt(7)
    r_f2.font.color.rgb = DIM

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ═══════════════════════════════════════════════════
#  UI 렌더링
# ═══════════════════════════════════════════════════

# ─── 공통 헤더 ───
st.markdown(
    '<div style="text-align:center;padding:1rem 0 0 0">'
    '<div class="header">B L U E &nbsp; J E A N S &nbsp; P I C T U R E S</div>'
    '<div class="brand-title">CREATOR ENGINE</div>'
    '<div class="sub">Y O U N G &nbsp; · &nbsp; V I N T A G E &nbsp; · &nbsp; F R E E &nbsp; · &nbsp; I N N O V A T I V E</div>'
    '</div>',
    unsafe_allow_html=True
)

# ─── 뒤로가기 ───
if st.session_state.view in ("project", "core", "char_bible", "structure", "scene_design", "treatment", "tone_doc") and st.session_state.cur:
    if st.session_state.view == "tone_doc":
        col_nav1, col_nav2 = st.columns(2)
        with col_nav1:
            if st.button("← 프로젝트 목록"):
                st.session_state.view = "home"
                st.rerun()
        with col_nav2:
            if st.button("← Treatment"):
                st.session_state.view = "treatment"
                st.rerun()
    elif st.session_state.view == "treatment":
        col_nav1, col_nav2 = st.columns(2)
        with col_nav1:
            if st.button("← 프로젝트 목록"):
                st.session_state.view = "home"
                st.rerun()
        with col_nav2:
            if st.button("← Scene Design"):
                st.session_state.view = "scene_design"
                st.rerun()
    elif st.session_state.view == "scene_design":
        col_nav1, col_nav2 = st.columns(2)
        with col_nav1:
            if st.button("← 프로젝트 목록"):
                st.session_state.view = "home"
                st.rerun()
        with col_nav2:
            if st.button("← Structure"):
                st.session_state.view = "structure"
                st.rerun()
    elif st.session_state.view == "structure":
        col_nav1, col_nav2 = st.columns(2)
        with col_nav1:
            if st.button("← 프로젝트 목록"):
                st.session_state.view = "home"
                st.rerun()
        with col_nav2:
            if st.button("← Character Bible"):
                st.session_state.view = "char_bible"
                st.rerun()
    elif st.session_state.view == "char_bible":
        col_nav1, col_nav2 = st.columns(2)
        with col_nav1:
            if st.button("← 프로젝트 목록"):
                st.session_state.view = "home"
                st.rerun()
        with col_nav2:
            if st.button("← Core Build"):
                st.session_state.view = "core"
                st.rerun()
    elif st.session_state.view == "core":
        col_nav1, col_nav2 = st.columns(2)
        with col_nav1:
            if st.button("← 프로젝트 목록"):
                st.session_state.view = "home"
                st.rerun()
        with col_nav2:
            if st.button("← Brainstorm"):
                st.session_state.view = "project"
                st.rerun()
    else:
        if st.button("← 프로젝트 목록"):
            st.session_state.view = "home"
            st.rerun()


# ═══════════════════════════════════════════════════
#  HOME
# ═══════════════════════════════════════════════════
if st.session_state.view == "home":

    # 새 프로젝트 생성
    with st.expander("➕ 새 프로젝트", expanded=not bool(st.session_state.projects)):
        col1, col2 = st.columns([2, 1])

        with col1:
            title_input = st.text_input(
                "프로젝트 제목",
                placeholder="예: 인도네시아 물귀신 프로젝트"
            )
            idea_input = st.text_area(
                "💡 아이디어",
                height=120,
                placeholder=(
                    "자유롭게 입력\n"
                    "예: 인도네시아용 물귀신 이야기\n"
                    "예: 은퇴한 킬러가 다시 돌아오는 이야기\n"
                    "예: 40대 여형사, 연쇄살인범이 딸의 담임교사"
                )
            )

        # ── LOCKED 설정 (v2.3: OPEN 필드 폐지) ──
        with st.expander("🔒 설정 잠금 (LOCKED)", expanded=False):
            st.caption(
                "LOCKED = 파이프라인 전 과정에서 절대 변경 불가. "
                "여기에 명시되지 않은 모든 디테일은 엔진이 자유롭게 창작합니다. "
                "한 줄에 하나씩 입력하세요."
            )
            locked_input = st.text_area(
                "🔒 LOCKED (반드시 지킬 것만 명시)",
                height=180,
                placeholder=(
                    "한 줄에 하나씩:\n"
                    "장르: 로맨틱 코미디 (다른 장르 전환 금지)\n"
                    "주인공: 강유진 (여, 미혼)\n"
                    "핵심 관계: 아버지의 상속 조건\n"
                    "공간: 쿠킹 클래스 (아이반 + 성인반)\n"
                    "테마: 사람은 요리처럼 섞을 수 없다\n"
                    "엔딩 방향: 한 명 선택 아닌 자기 발견"
                )
            )
            # v2.3: OPEN 필드 폐지 — 하위 호환성을 위해 빈 리스트 유지
            open_input = ""


        with col2:
            genre_input = st.selectbox(
                "🎬 장르",
                ["미지정", "범죄/스릴러", "드라마", "액션", "로맨스", "코미디",
                 "로맨틱 코미디", "호러/공포", "SF", "판타지", "시대극/사극", "느와르",
                 "미스터리", "전쟁", "뮤지컬", "다큐/논픽션"]
            )

            market_type = st.selectbox(
                "🌏 타겟 시장",
                ["미지정", "한국", "북미/미국", "일본", "중국",
                 "동남아", "유럽", "중동", "글로벌", "직접 입력"]
            )

            market_custom = ""
            if market_type == "직접 입력":
                market_custom = st.text_input(
                    "시장 직접 입력",
                    placeholder="예: 인도네시아+한국 공동제작"
                )

            format_input = st.selectbox(
                "📐 포맷",
                ["미지정", "영화", "시리즈", "미니시리즈(4~8화)",
                 "웹툰", "웹소설", "숏폼", "다큐멘터리", "애니메이션"]
            )

            st.markdown("##### 🏛️ 특수 규칙")
            fact_based_input = st.checkbox(
                "실화 배경 작품",
                value=False,
                help="실제 사건/인물/시대를 배경으로 하는 작품. 실명 비사용 + 사실성 균형, 명예훼손·인격권 차단 규칙이 자동 적용됨."
            )
            historical_input = st.checkbox(
                "역사영화",
                value=False,
                help="시대극/사극. 시대 언어의 균형, 공간의 주인공화, 선악 이원 구도 회피 등 역사영화 전용 규칙 적용."
            )
            film_type_input = ""
            if historical_input:
                film_type_input = st.selectbox(
                    "역사영화 유형",
                    ["정통", "팩션", "퓨전"],
                    help="정통: 〈남한산성〉〈1987〉〈서울의 봄〉 / 팩션: 〈왕의 남자〉〈광해〉〈암살〉 / 퓨전: 〈조선명탐정〉〈전우치〉"
                )

        if st.button("🚀 프로젝트 생성", use_container_width=True, disabled=not idea_input.strip()):
            market_final = market_custom if market_type == "직접 입력" else market_type
            project_id = f"p_{datetime.now().strftime('%Y%m%d_%H%M%S')}"

            # LOCKED 파싱 (줄 단위). v2.3: OPEN 필드는 폐지되어 빈 리스트로 유지 (하위 호환성)
            locked_list = [line.strip() for line in locked_input.strip().split("\n") if line.strip()] if locked_input.strip() else []
            open_list = []  # v2.3: OPEN 필드 폐지

            st.session_state.projects[project_id] = {
                "project_id": project_id,
                "title": title_input or "새 프로젝트",
                "idea_text": idea_input,
                "genre": genre_input,
                "target_market": market_final,
                "format": format_input,
                "fact_based": fact_based_input,
                "historical": historical_input,
                "film_type": film_type_input,
                "locked_items": locked_list,
                "open_items": open_list,
                "created_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
                "updated_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
                "research": None,
                "brainstorm_cards": None,
                "brainstorm_analysis": None,
                "core": None,
                "core_gate": None,
                "char_bible": None,
                "structure_story": None,
                "structure_diag": None,
                "structure_gate": None,
                "structure_prose": None,
                "scene_design": None,
                "treatment": None,
                "treatment_gate": None,
                "tone_doc": None,
                "final_score": None,
            }

            st.session_state.cur = project_id
            st.session_state.view = "project"
            st.rerun()

    # ─── 프로젝트 JSON 불러오기 (v2.3.5 신규) ───
    with st.expander("📂 프로젝트 JSON 불러오기 (세션 복구)", expanded=False):
        st.caption(
            "이전 세션에서 저장한 프로젝트 JSON 파일을 업로드하면 해당 시점으로 복원됩니다. "
            "Scene Design 완료 시점에 저장했다면 Treatment부터 이어서 진행 가능합니다."
        )
        uploaded_json = st.file_uploader(
            "프로젝트 JSON 파일 선택",
            type=["json"],
            key="project_json_uploader",
            help="Scene 완료 / Treatment 완료 / 최종완성 시점에 저장한 JSON 파일"
        )

        if uploaded_json is not None:
            try:
                json_str = uploaded_json.read().decode("utf-8")
                result = load_project_from_json(json_str)
                restored_project = result["project"]
                meta = result["meta"]
                warnings = result["warnings"]

                # 정보 표시
                st.markdown(
                    f'<div class="callout">'
                    f'<div class="cl">📋 불러올 프로젝트</div>'
                    f'<b>{restored_project.get("title", "제목 없음")}</b><br>'
                    f'{restored_project.get("genre", "")} · '
                    f'{restored_project.get("target_market", "")} · '
                    f'{restored_project.get("format", "")}<br>'
                    f'<span style="font-size:.75rem;color:var(--dim)">'
                    f'저장 시점: {meta.get("saved_at", "알 수 없음")} / '
                    f'단계: {meta.get("stage", "알 수 없음")} / '
                    f'엔진 버전: {meta.get("engine_version", "알 수 없음")}'
                    f'</span></div>',
                    unsafe_allow_html=True
                )

                # 경고 표시
                for w in warnings:
                    st.warning(w)

                # 중복 ID 체크 및 복원 버튼
                original_id = restored_project.get("project_id", "")
                id_conflict = original_id in st.session_state.projects

                if id_conflict:
                    st.info(
                        f"⚠️ 동일한 프로젝트 ID가 이미 존재합니다. "
                        f"복원 시 새 ID를 부여하여 별도 프로젝트로 추가됩니다."
                    )

                if st.button("✅ 이 프로젝트 복원하기", type="primary", use_container_width=True):
                    # 중복 방지: ID 충돌 시 새 ID 생성
                    if id_conflict:
                        new_id = f"proj_{int(datetime.now().timestamp() * 1000)}"
                        restored_project["project_id"] = new_id
                        restored_project["title"] = restored_project.get("title", "복원") + " (복원본)"
                        restored_id = new_id
                    else:
                        restored_id = original_id

                    # updated_at 갱신
                    restored_project["updated_at"] = datetime.now().strftime("%Y-%m-%d %H:%M")

                    # 세션 상태에 주입
                    st.session_state.projects[restored_id] = restored_project
                    st.session_state.cur = restored_id

                    # 단계별 적절한 view로 이동 (완료된 가장 높은 단계로)
                    if restored_project.get("tone_doc"):
                        st.session_state.view = "tone_doc"
                    elif restored_project.get("treatment"):
                        st.session_state.view = "treatment"
                    elif restored_project.get("scene_design"):
                        st.session_state.view = "scene_design"
                    elif restored_project.get("structure_story"):
                        st.session_state.view = "structure"
                    elif restored_project.get("char_bible"):
                        st.session_state.view = "char_bible"
                    elif restored_project.get("core"):
                        st.session_state.view = "core"
                    else:
                        st.session_state.view = "project"

                    st.success(f"✅ '{restored_project['title']}' 복원 완료. 해당 단계로 이동합니다.")
                    st.rerun()

            except ValueError as e:
                st.error(f"❌ 불러오기 실패: {e}")
            except Exception as e:
                st.error(f"❌ 예상치 못한 오류: {e}")

    # 프로젝트 목록
    if st.session_state.projects:
        st.markdown("---")
        st.markdown("### 📁 프로젝트")

        for project_id, project in sorted(
            st.session_state.projects.items(),
            key=lambda x: x[1]["updated_at"],
            reverse=True
        ):
            col1, col2 = st.columns([5, 1])

            with col1:
                has_cards = "✅" if project.get("brainstorm_cards") else "—"
                has_analysis = "✅" if project.get("brainstorm_analysis") else "—"

                st.markdown(
                    f'<div class="card">'
                    f'<b>{project["title"]}</b><br>'
                    f'<span style="font-size:.75rem;color:var(--dim)">'
                    f'{project["genre"]} · {project["target_market"]} · {project["format"]} · '
                    f'{project["updated_at"]}'
                    f'<br>카드 {has_cards} · 분석 {has_analysis}'
                    f'</span></div>',
                    unsafe_allow_html=True
                )

            with col2:
                if st.button("열기 →", key=f"open_{project_id}"):
                    st.session_state.cur = project_id
                    st.session_state.view = "project"
                    st.rerun()


# ═══════════════════════════════════════════════════
#  PROJECT (단일 페이지 스크롤)
# ═══════════════════════════════════════════════════
elif st.session_state.view == "project" and st.session_state.cur:

    project = st.session_state.projects[st.session_state.cur]

    # ─── 프로젝트 헤더 ───
    st.markdown(f"## {project['title']}")
    st.caption(f"{project['genre']} · {project['target_market']} · {project['format']}")

    # ─── 단계 표시 ───
    render_stepper("project", project)

    st.markdown(
        f'<div class="callout">'
        f'<div class="cl">IDEA</div>'
        f'{project["idea_text"]}'
        f'</div>',
        unsafe_allow_html=True
    )

    # ─── LOCKED 표시 + 편집 (v2.3: OPEN 필드 폐지) ───
    locked = project.get("locked_items", [])
    open_items = project.get("open_items", [])  # 하위 호환성 - 기존 프로젝트 로딩 시에만 사용

    if locked:
        locked_html = "".join(f"<li>{item}</li>" for item in locked)
        st.markdown(
            f'<div style="background:#FFF3CD;padding:8px 12px;border-radius:6px;border-left:4px solid #FFCB05;font-size:.82rem">'
            f'<b>🔒 LOCKED</b> — 변경 불가<ul style="margin:4px 0 0 0;padding-left:18px">{locked_html}</ul></div>',
            unsafe_allow_html=True
        )

    with st.expander("🔒 LOCKED 편집", expanded=False):
        st.caption("LOCKED = 반드시 지킬 것만 명시. 여기에 없는 것은 엔진이 자유롭게 창작합니다.")
        edited_locked = st.text_area(
            "🔒 LOCKED (변경 불가)",
            value="\n".join(locked),
            height=180,
            key="edit_locked",
            placeholder="한 줄에 하나씩"
        )
        if st.button("💾 LOCKED 저장", key="save_locked"):
            project["locked_items"] = [l.strip() for l in edited_locked.strip().split("\n") if l.strip()]
            project["open_items"] = []  # v2.3: OPEN 필드 폐지 - 빈 리스트 유지
            project["updated_at"] = datetime.now().strftime("%Y-%m-%d %H:%M")
            st.success("저장 완료. 이후 생성되는 모든 단계에 반영됩니다.")
            st.rerun()

    st.markdown("---")

    # ═══════════════════════════════════════
    # STEP 1: 리서치 (선택)
    # ═══════════════════════════════════════
    st.markdown('<div class="section-header">🔍 리서치 <span class="en">RESEARCH</span></div>', unsafe_allow_html=True)
    st.caption("실화/뉴스 + 기존 작품 정보 검색. 건너뛰어도 됩니다.")

    if st.button("🔍 리서치 실행"):
        with st.spinner("리서치 정리 중..."):
            result = call_research(
                project["idea_text"],
                project["genre"],
                project["target_market"]
            )
            if result:
                project["research"] = result
                project["updated_at"] = datetime.now().strftime("%Y-%m-%d %H:%M")
                st.rerun()

    # 리서치 결과 표시
    if project.get("research"):
        research_data = project["research"]
        summary = research_data.get("research_summary", {})

        with st.expander(
            f"📰 리서치 결과 — "
            f"실화 {summary.get('total_real_events', 0)}건 · "
            f"작품 {summary.get('total_existing_works', 0)}건",
            expanded=True
        ):
            # 실화/뉴스
            if research_data.get("real_events"):
                st.markdown("**📰 실화 / 뉴스**")
                for event in research_data["real_events"]:
                    st.markdown(
                        f'<div class="ri">'
                        f'<div class="rl">#{event.get("id", "")} '
                        f'[{event.get("year", "")}] {event.get("source", "")}</div>'
                        f'<b>{event.get("title", "")}</b><br>'
                        f'{event.get("summary", "")}<br>'
                        f'<span style="color:var(--navy)">→ {event.get("story_potential", "")}</span>'
                        f'</div>',
                        unsafe_allow_html=True
                    )

            # 기존 작품
            if research_data.get("existing_works"):
                st.markdown("**🎬 기존 작품**")
                for work in research_data["existing_works"]:
                    st.markdown(
                        f'<div class="ri">'
                        f'<div class="rl">#{work.get("id", "")} '
                        f'{work.get("type", "")} · {work.get("country", "")} · '
                        f'{work.get("year", "")}</div>'
                        f'<b>{work.get("title", "")}</b><br>'
                        f'유사: {work.get("similarity", "")}<br>'
                        f'<span style="color:var(--navy)">→ 차별화: '
                        f'{work.get("difference_opportunity", "")}</span>'
                        f'</div>',
                        unsafe_allow_html=True
                    )

            # 핵심 시사점
            if summary.get("key_insight"):
                st.markdown(
                    f'<div class="callout">'
                    f'<div class="cl">💡 핵심 시사점</div>'
                    f'{summary["key_insight"]}'
                    f'</div>',
                    unsafe_allow_html=True
                )

    st.markdown("---")

    # ═══════════════════════════════════════
    # STEP 2: Brainstorm (2단계 분할 호출)
    # ═══════════════════════════════════════
    st.markdown('<div class="section-header">🧠 Brainstorm <span class="en">CONCEPT IDEATION</span></div>', unsafe_allow_html=True)

    if st.button("🧠 Brainstorm 실행", type="primary"):

        # ── 1단계: 카드 생성 ──
        with st.spinner("① 컨셉 카드 생성 중... (약 20~30초)"):
            cards_result = call_brainstorm_cards(
                project["idea_text"],
                project["genre"],
                project["target_market"],
                project["format"],
                project.get("research"),
                locked_block=_build_project_locked_block(project)
            )

        if cards_result:
            project["brainstorm_cards"] = cards_result
            project["updated_at"] = datetime.now().strftime("%Y-%m-%d %H:%M")

            # Top 3 카드 추출
            all_cards = cards_result.get("idea_cards", [])
            cards_map = {c["id"]: c for c in all_cards}

            top3_data = []
            for t in cards_result.get("top3", []):
                card = cards_map.get(t["card_id"])
                if card:
                    top3_data.append(card)

            # ── 2단계: 분석 + Gate A ──
            if top3_data:
                with st.spinner("② 시장 분석 + Gate A 채점 중... (약 10~20초)"):
                    analysis_result = call_brainstorm_analysis(
                        project["idea_text"],
                        project["genre"],
                        project["target_market"],
                        project["format"],
                        top3_data,
                        project.get("research"),
                        locked_block=_build_project_locked_block(project)
                    )

                if analysis_result:
                    project["brainstorm_analysis"] = analysis_result
                    project["updated_at"] = datetime.now().strftime("%Y-%m-%d %H:%M")
            else:
                st.warning("Top 3 카드를 추출하지 못했습니다. 카드만 표시합니다.")

            st.rerun()

    # ═══════════════════════════════════════
    # 결과 표시
    # ═══════════════════════════════════════
    if project.get("brainstorm_cards"):
        bc = project["brainstorm_cards"]
        ba = project.get("brainstorm_analysis", {})

        # ── 아이디어 유형 ──
        idea_type = bc.get("idea_type", "").upper()
        diagnosis = bc.get("idea_type_diagnosis", "")
        action = bc.get("idea_type_action", "")
        type_color = {"STORY": "var(--g)", "MOOD": "var(--navy)", "HYBRID": "#FF8C00"}.get(idea_type, "var(--dim)")
        action_html = (
            f'<div style="margin-top:.5rem;padding:.4rem .7rem;background:#FFF8E1;'
            f'border-radius:6px;font-size:.85rem;font-weight:700;color:#8B6914">→ {action}</div>'
        ) if action else ""
        st.markdown(
            f'<div class="callout" style="border-left:4px solid {type_color}">'
            f'<div class="cl" style="color:{type_color}">아이디어 유형: {idea_type}</div>'
            f'<div style="margin:.3rem 0;font-size:.9rem">{diagnosis}</div>'
            f'{action_html}'
            f'</div>',
            unsafe_allow_html=True
        )

        # ── 시장 · 포맷 맥락 (분석 있을 때) ──
        if ba:
            market_ctx = ba.get("market_context", {})
            format_ctx = ba.get("format_context", {})

            col1, col2 = st.columns(2)

            with col1:
                ref_titles = ", ".join(market_ctx.get("reference_titles", []))
                st.markdown(
                    f'<div class="callout">'
                    f'<div class="cl">🌏 시장 — {market_ctx.get("target_market", "")}</div>'
                    f'기회: {market_ctx.get("market_insight", "")}<br>'
                    f'문화코드: {market_ctx.get("cultural_code", "")}<br>'
                    f'리스크: {market_ctx.get("market_risk", "")}<br>'
                    f'참고작: {ref_titles}'
                    f'</div>',
                    unsafe_allow_html=True
                )

            with col2:
                st.markdown(
                    f'<div class="callout">'
                    f'<div class="cl">📐 포맷 — {format_ctx.get("selected_format", "")}</div>'
                    f'적합성: {format_ctx.get("format_rationale", "")}<br>'
                    f'구조: {format_ctx.get("structure_note", "")}'
                    f'</div>',
                    unsafe_allow_html=True
                )

            # ── 훅 문장 ──
            hook = ba.get("hook_sentence", "")
            if hook:
                st.markdown(
                    f'<div style="text-align:center;padding:1.5rem 0">'
                    f'<div style="font-size:.7rem;color:var(--y);font-weight:600">HOOK</div>'
                    f'<div style="font-size:1.15rem;font-weight:600;color:var(--t);line-height:1.5">'
                    f'"{hook}"</div>'
                    f'</div>',
                    unsafe_allow_html=True
                )

        st.markdown("---")

        # ── Top 3 컨셉 ──
        st.markdown("#### 🏆 Top 3 컨셉")

        cards_map = {c["id"]: c for c in bc.get("idea_cards", [])}

        for top_item in bc.get("top3", []):
            card = cards_map.get(top_item["card_id"], {})
            if card:
                col1, col2 = st.columns([5, 1])

                with col1:
                    st.markdown(
                        f'<div class="card">'
                        f'<span style="color:var(--navy);font-weight:700">#{top_item["rank"]}</span> '
                        f'<b>{card.get("title", "")}</b><br>'
                        f'<span style="color:#444">{card.get("logline_seed", "")}</span><br>'
                        f'<span style="font-size:.8rem;color:#666">'
                        f'👤 {card.get("protagonist", "")}<br>'
                        f'⚔️ {card.get("conflict", "")}<br>'
                        f'✨ {card.get("hook", "")}<br>'
                        f'🎬 {card.get("visual_image", "")}'
                        f'</span><br>'
                        f'<span style="font-size:.75rem;color:var(--dim)">'
                        f'이유: {top_item.get("reason", "")}</span>'
                        f'</div>',
                        unsafe_allow_html=True
                    )

                with col2:
                    st.markdown(
                        f'<div class="big">{card.get("total_score", 0.0)}</div>'
                        f'<div class="sm">{card.get("genre", "")}</div>',
                        unsafe_allow_html=True
                    )

        # ── 차별화 포인트 (분석 있을 때) ──
        if ba:
            differentiation = ba.get("differentiation", [])
            if differentiation:
                st.markdown("#### 💎 차별화 포인트")
                for idx, item in enumerate(differentiation, 1):
                    st.markdown(f"**{idx}.** {item}")

            # ── 개발 우선순위 ──
            dev_priority = ba.get("development_priority", {})
            if dev_priority:
                st.markdown("#### 🧭 개발 우선순위")
                col1, col2, col3 = st.columns(3)

                col1.markdown(
                    f'<div class="callout">'
                    f'<div class="cl">추천 방향</div>'
                    f'{dev_priority.get("recommended_direction", "")}'
                    f'</div>',
                    unsafe_allow_html=True
                )
                col2.markdown(
                    f'<div class="callout">'
                    f'<div class="cl">Core Build 집중</div>'
                    f'{dev_priority.get("next_step", "")}'
                    f'</div>',
                    unsafe_allow_html=True
                )
                col3.markdown(
                    f'<div class="callout" style="border-left-color:var(--r)">'
                    f'<div class="cl" style="color:var(--r)">리스크</div>'
                    f'{dev_priority.get("risk", "")}'
                    f'</div>',
                    unsafe_allow_html=True
                )

            st.markdown("---")

            # ── Gate A: Concept Gate ──
            gate = ba.get("gate_a_scores", {})
            avg = gate.get("average", 0)
            passed = avg >= 7.0

            st.markdown("#### 🚪 Gate A: Concept Gate")

            col1, col2 = st.columns([1, 2])

            with col1:
                gate_color = "var(--g)" if passed else "var(--r)"
                gate_label = "PASS" if passed else "FAIL"
                st.markdown(
                    f'<div style="text-align:center">'
                    f'<div class="big" style="color:{gate_color}">{avg}</div>'
                    f'<div class="sm" style="color:{gate_color};'
                    f'font-size:1rem;font-weight:700">{gate_label}</div>'
                    f'</div>',
                    unsafe_allow_html=True
                )

            with col2:
                gate_items = [
                    ("주인공", gate.get("protagonist_visible", 0)),
                    ("갈등", gate.get("conflict_one_line", 0)),
                    ("차별점", gate.get("differentiation", 0)),
                    ("포스터", gate.get("poster_image", 0)),
                    ("시장성", gate.get("market_potential", 0)),
                ]
                for name, score in gate_items:
                    bar_pct = score * 10
                    st.markdown(
                        f'<div style="display:flex;align-items:center;'
                        f'margin:.2rem 0;font-size:.8rem">'
                        f'<div style="width:60px;color:var(--dim)">{name}</div>'
                        f'<div style="flex:1;background:#E0E0E8;'
                        f'border-radius:4px;height:8px;margin:0 .5rem">'
                        f'<div style="width:{bar_pct}%;background:var(--y);'
                        f'height:100%;border-radius:4px"></div>'
                        f'</div>'
                        f'<div style="width:30px;text-align:right">{score}</div>'
                        f'</div>',
                        unsafe_allow_html=True
                    )

            if passed:
                st.success("✅ Gate A 통과. Core Build 진행 가능.")
                if st.button("🎯 Core Build 진행 →", type="primary", use_container_width=True):
                    project["stage"] = "core"
                    st.session_state.view = "core"
                    st.rerun()
            else:
                st.warning(
                    f"⚠️ Gate A 미통과 (평균 {avg}). "
                    f"아이디어 보강 또는 재실행 권장."
                )
                col_o1, col_o2 = st.columns(2)
                with col_o1:
                    if st.button("🔓 Override (강제 통과)"):
                        project["stage"] = "core"
                        st.session_state.view = "core"
                        st.rerun()
                with col_o2:
                    if st.button("🔄 Brainstorm 재실행"):
                        project["brainstorm_cards"] = None
                        project["brainstorm_analysis"] = None
                        st.rerun()

        # ── 전체 아이디어 카드 (접이식) ──
        all_cards = bc.get("idea_cards", [])
        if all_cards:
            with st.expander(f"📋 전체 아이디어 카드 ({len(all_cards)}개)"):
                for card in sorted(
                    all_cards,
                    key=lambda x: x.get("total_score", 0),
                    reverse=True
                ):
                    st.markdown(
                        f"**#{card['id']} {card.get('title', '')}** — "
                        f"{card.get('total_score', 0.0)}점 &nbsp; "
                        f"{card.get('logline_seed', '')}"
                    )

        # ── 분석 없을 때 fallback 네비게이션 ──
        if not ba:
            st.markdown("---")
            st.warning("⚠️ 시장 분석 + Gate A 채점이 아직 완료되지 않았습니다.")

            col_fb1, col_fb2 = st.columns(2)
            with col_fb1:
                if st.button("🔄 분석 재시도", use_container_width=True):
                    # 2단계만 재실행
                    cards_map = {c["id"]: c for c in bc.get("idea_cards", [])}
                    top3_data = [
                        cards_map[t["card_id"]]
                        for t in bc.get("top3", [])
                        if t["card_id"] in cards_map
                    ]
                    if top3_data:
                        with st.spinner("② 시장 분석 + Gate A 채점 중..."):
                            ar = call_brainstorm_analysis(
                                project["idea_text"],
                                project["genre"],
                                project["target_market"],
                                project["format"],
                                top3_data,
                                project.get("research"),
                                locked_block=_build_project_locked_block(project)
                            )
                        if ar:
                            project["brainstorm_analysis"] = ar
                            project["updated_at"] = datetime.now().strftime("%Y-%m-%d %H:%M")
                            st.rerun()
            with col_fb2:
                if st.button("🎯 분석 건너뛰고 Core Build →", use_container_width=True):
                    project["stage"] = "core"
                    st.session_state.view = "core"
                    st.rerun()

    # ─── 푸터 ───
    st.markdown("---")
    st.caption(ENGINE_FOOTER)


# ═══════════════════════════════════════════════════
#  CORE BUILD
# ═══════════════════════════════════════════════════
elif st.session_state.view == "core" and st.session_state.cur:

    project = st.session_state.projects[st.session_state.cur]
    bc = project.get("brainstorm_cards", {})
    ba = project.get("brainstorm_analysis", {})

    st.markdown(f"## {project['title']}")
    st.caption(f"{project['genre']} · {project['target_market']} · {project['format']}")

    # ─── 단계 표시 ───
    render_stepper("core", project)

    # Brainstorm에서 인계받은 정보 표시
    selected_concept = None
    if bc:
        cards_map = {c["id"]: c for c in bc.get("idea_cards", [])}
        top3 = bc.get("top3", [])
        if top3:
            selected_concept = cards_map.get(top3[0].get("card_id"), {})
            if selected_concept:
                st.markdown(
                    f'<div class="callout">'
                    f'<div class="cl">🏆 선정 컨셉: {selected_concept.get("title", "")}</div>'
                    f'{selected_concept.get("logline_seed", "")}'
                    f'</div>',
                    unsafe_allow_html=True
                )

    if ba:
        dp = ba.get("development_priority", {})
        if dp.get("next_step"):
            st.markdown(
                f'<div class="callout">'
                f'<div class="cl">Core Build 집중 포인트</div>'
                f'{dp["next_step"]}'
                f'</div>',
                unsafe_allow_html=True
            )

    st.markdown("---")

    st.markdown('<div class="section-header">🎯 Core Build <span class="en">CORE DEVELOPMENT</span></div>', unsafe_allow_html=True)
    st.caption("로그라인 · 기획의도 · 세계관 · 캐릭터 · Goal/Need/Strategy를 고정합니다.")

    if st.button("🎯 Core Build 실행", type="primary"):
        if not selected_concept:
            st.error("선정된 컨셉이 없습니다. Brainstorm을 먼저 실행해주세요.")
        else:
            with st.spinner("① Core Build 생성 중... (약 30~40초)"):
                core_result = call_core_build_main(
                    project["idea_text"], project["genre"],
                    project["target_market"], project["format"],
                    selected_concept, project.get("research"),
                    locked_block=_build_project_locked_block(project),
                    fact_based=project.get("fact_based", False),
                    historical=project.get("historical", False),
                    film_type=project.get("film_type", ""),
                )
            if core_result:
                project["core"] = core_result
                project["updated_at"] = datetime.now().strftime("%Y-%m-%d %H:%M")
                with st.spinner("② Gate B + Gate C 채점 중... (약 10~20초)"):
                    gate_result = call_core_gate(core_result)
                if gate_result:
                    project["core_gate"] = gate_result
                    project["final_score"] = gate_result.get("five_axis_scores", {}).get("final_score")
                    project["updated_at"] = datetime.now().strftime("%Y-%m-%d %H:%M")
                st.rerun()

    # ── Core Build 결과 표시 ──
    if project.get("core"):
        core = project["core"]

        # Logline Pack
        st.markdown("#### 📝 Logline Pack")
        lp = core.get("logline_pack", {})
        for label, key in [("Original", "original"), ("Washed", "washed"),
                           ("투자자용", "investor"), ("감독용", "director"),
                           ("캐릭터 훅", "character_hook")]:
            val = lp.get(key, "")
            if val:
                st.markdown(f'<div class="callout"><div class="cl">{label}</div>{val}</div>', unsafe_allow_html=True)

        # Goal / Need / Strategy
        st.markdown("#### 🎯 Goal / Need / Strategy")
        gns = core.get("goal_need_strategy", {})
        c1, c2, c3 = st.columns(3)
        c1.markdown(f'<div class="callout"><div class="cl">GOAL</div>{gns.get("goal","")}</div>', unsafe_allow_html=True)
        c2.markdown(f'<div class="callout"><div class="cl">NEED</div>{gns.get("need","")}</div>', unsafe_allow_html=True)
        c3.markdown(f'<div class="callout"><div class="cl">STRATEGY</div>{gns.get("strategy","")}</div>', unsafe_allow_html=True)
        cr, ce = st.columns(2)
        cr.markdown(f'<div class="callout" style="border-left-color:var(--r)"><div class="cl" style="color:var(--r)">RISK</div>{gns.get("risk","")}</div>', unsafe_allow_html=True)
        ce.markdown(f'<div class="callout" style="border-left-color:var(--g)"><div class="cl" style="color:var(--g)">ENDING PAYOFF</div>{gns.get("ending_payoff","")}</div>', unsafe_allow_html=True)

        # Narrative Drive (BJND)
        nd = core.get("narrative_drive", {})
        if nd and nd.get("desire_origin"):
            origin = nd.get("desire_origin", "")
            origin_kr = "상실(Loss)" if origin == "loss" else "결핍(Lack)" if origin == "lack" else origin
            st.markdown(
                f'<div class="callout" style="border-left-color:#9B59B6">'
                f'<div class="cl" style="color:#9B59B6">🔮 서사동력 — Narrative Drive</div>'
                f'<b>{origin_kr}</b>: {nd.get("origin_detail","")}<br>'
                f'아크: {nd.get("arc_direction","")} · 해결전략: {nd.get("resolution_strategy","")}<br>'
                f'<span style="font-size:.8rem;color:#666">Goal↔Need 간극: {nd.get("goal_need_gap","")}</span>'
                f'</div>',
                unsafe_allow_html=True
            )

        # Project Intent (소재 · 장르 · 시장분석)
        st.markdown("#### 📋 기획의도 — 왜 이 작품을 기획하는가")
        pi = core.get("project_intent", {})

        if pi.get("subject"):
            st.markdown(f'<div class="callout"><div class="cl">소재 — 왜 이 소재인가</div>{pi["subject"]}</div>', unsafe_allow_html=True)

        if pi.get("genre_approach"):
            st.markdown(f'<div class="callout"><div class="cl">장르 — 왜 이 장르인가</div>{pi["genre_approach"]}</div>', unsafe_allow_html=True)

        if pi.get("market_rationale"):
            st.markdown(f'<div class="callout"><div class="cl">시장 — 왜 지금 이 시장인가</div>{pi["market_rationale"]}</div>', unsafe_allow_html=True)

        cp1, cp2 = st.columns(2)
        if pi.get("pitch"):
            cp1.markdown(f'<div class="callout"><div class="cl">Elevator Pitch</div>{pi["pitch"]}</div>', unsafe_allow_html=True)
        if pi.get("theme"):
            cp2.markdown(f'<div class="callout"><div class="cl">Theme</div>{pi["theme"]}</div>', unsafe_allow_html=True)
        if pi.get("tone_manner"):
            st.markdown(f"**Tone & Manner:** {', '.join(pi['tone_manner'])}")

        # World Build
        st.markdown("#### 🌍 세계관")
        wb = core.get("world_build", {})
        cw1, cw2 = st.columns(2)
        with cw1:
            for label, key in [("시간","time"),("공간","space"),("규칙","rules"),("금기","taboo"),("권력구조","power_structure")]:
                val = wb.get(key, "")
                if val:
                    st.markdown(f"**{label}:** {val}")
        with cw2:
            if wb.get("visual_keywords"):
                st.markdown(f"**시각 키워드:** {', '.join(wb['visual_keywords'])}")
            if wb.get("conflict_points"):
                st.markdown("**충돌 포인트:**")
                for cp in wb["conflict_points"]:
                    st.markdown(f"- {cp}")

        # Characters
        st.markdown("#### 🎭 캐릭터")
        role_labels = {"protagonist":"주인공","antagonist":"적대자","ally":"조력자","mirror":"거울","catalyst":"촉매자","subplot_lead":"서브플롯 리드"}
        all_display_chars = core.get("characters", []) + core.get("extended_characters", [])
        for ch in all_display_chars:
            role = role_labels.get(ch.get("role",""), ch.get("role",""))
            st.markdown(
                f'<div class="card"><div class="cl">{role}: {ch.get("name","")}</div>'
                f'{ch.get("description","")}<br>'
                f'<span style="font-size:.8rem;color:#666">'
                f'🎯 {ch.get("goal","")}<br>💔 {ch.get("need",ch.get("flaw",""))}<br>'
                f'⚡ {ch.get("strategy","")}<br>🗣️ {ch.get("dialogue_tone","")}</span></div>',
                unsafe_allow_html=True
            )

        # Relationship
        rel = core.get("relationship_map", [])
        if rel:
            st.markdown("#### 🔗 관계도")
            for r in rel:
                st.markdown(f"- {r}")

        # ── 매력 설계도 (Attraction Design) ──
        ad = core.get("attraction_design", {})
        if ad:
            st.markdown("---")
            st.markdown('<div class="section-header">⚡ 매력 설계도 <span class="en">ATTRACTION DESIGN</span></div>', unsafe_allow_html=True)
            st.caption("이 이야기가 멈출 수 없게 만드는 핵심 설계 — Writer/Series/Novel Engine에 자동 전달됩니다.")

            # 첫 장면
            oh = ad.get("opening_hook", {})
            if oh:
                hook_type = oh.get("type", "")
                hook_desc = oh.get("description", "")
                hook_check = oh.get("forbidden_check", "")
                st.markdown(
                    f'<div class="card" style="border-left:4px solid var(--y)">'
                    f'<div class="cl">🎬 첫 장면 — OPENING HOOK</div>'
                    f'<div style="font-size:.75rem;color:var(--navy);font-weight:700;margin-bottom:.4rem">[{hook_type}]</div>'
                    f'<p style="line-height:1.7;margin:.3rem 0">{hook_desc}</p>'
                    f'<div style="font-size:.75rem;color:var(--g);margin-top:.4rem">✓ {hook_check}</div>'
                    f'</div>',
                    unsafe_allow_html=True
                )

            # 배반 포인트
            tp = ad.get("twist_point", {})
            if tp:
                c1, c2 = st.columns(2)
                with c1:
                    st.markdown(
                        f'<div class="ri"><div class="rl">👀 관객이 예상하는 전개</div>'
                        f'{tp.get("expected_direction", "")}</div>',
                        unsafe_allow_html=True
                    )
                with c2:
                    st.markdown(
                        f'<div class="ri" style="border-left:3px solid var(--y)">'
                        f'<div class="rl">💥 배반 포인트 — 이 이야기만의 방향</div>'
                        f'{tp.get("betrayal", "")}</div>',
                        unsafe_allow_html=True
                    )
                if tp.get("why_more_true"):
                    st.markdown(
                        f'<div class="callout"><div class="cl">배반이 더 진실된 이유</div>'
                        f'{tp["why_more_true"]}</div>',
                        unsafe_allow_html=True
                    )

            # Water Cooler Moment
            wc = ad.get("water_cooler_moment", {})
            if wc:
                st.markdown(
                    f'<div class="card" style="background:var(--light-bg)">'
                    f'<div class="cl">💬 Water Cooler Moment — 말하고 싶어지는 장면</div>'
                    f'<p style="font-size:.95rem;font-weight:700;margin:.3rem 0">{wc.get("scene_or_setup","")}</p>'
                    f'<div style="font-size:.8rem;color:var(--dim)">{wc.get("why_memorable","")}</div>'
                    f'</div>',
                    unsafe_allow_html=True
                )

            # 한국 구체성 + 빌런 논리
            c1, c2 = st.columns(2)
            with c1:
                ks = ad.get("korean_specificity", [])
                if ks:
                    ks_html = "<br>".join([f"• {k}" for k in ks])
                    st.markdown(
                        f'<div class="ri"><div class="rl">🇰🇷 한국 구체성</div>{ks_html}</div>',
                        unsafe_allow_html=True
                    )
            with c2:
                vl = ad.get("villain_logic", "")
                if vl:
                    st.markdown(
                        f'<div class="ri"><div class="rl">🖤 빌런의 논리</div>{vl}</div>',
                        unsafe_allow_html=True
                    )

            # 감정 폭발 설계
            ee = ad.get("emotional_explosion", {})
            if ee:
                c1, c2 = st.columns(2)
                with c1:
                    st.markdown(
                        f'<div class="ri"><div class="rl">🔒 감정 억압 구간</div>'
                        f'{ee.get("suppression","")}</div>',
                        unsafe_allow_html=True
                    )
                with c2:
                    st.markdown(
                        f'<div class="ri" style="border-left:3px solid var(--r)">'
                        f'<div class="rl">💣 폭발 장면</div>'
                        f'{ee.get("explosion_moment","")}</div>',
                        unsafe_allow_html=True
                    )

            # 금지 방향
            fd = ad.get("forbidden_directions", [])
            if fd:
                fd_html = " &nbsp;|&nbsp; ".join([f'<span style="color:var(--r)">✗ {f}</span>' for f in fd])
                st.markdown(
                    f'<div class="ri"><div class="rl">🚫 절대 가면 안 되는 방향</div>{fd_html}</div>',
                    unsafe_allow_html=True
                )

        st.markdown("---")

        # Gate B + C
        if project.get("core_gate"):
            cg = project["core_gate"]
            gb = cg.get("gate_b_drive", {})
            gc = cg.get("gate_c_character", {})
            fa = cg.get("five_axis_scores", {})

            st.markdown("#### 🚪 Gate B: Drive Gate")
            c1, c2 = st.columns([1, 2])
            with c1:
                gb_avg = gb.get("average", 0)
                cl = "var(--g)" if gb_avg >= 7.0 else "var(--r)"
                lb = "PASS" if gb_avg >= 7.0 else "FAIL"
                st.markdown(f'<div style="text-align:center"><div class="big" style="color:{cl}">{gb_avg}</div><div class="sm" style="color:{cl};font-size:1rem;font-weight:700">{lb}</div></div>', unsafe_allow_html=True)
            with c2:
                for nm, sc in [("Goal 선명도",gb.get("goal_clarity",0)),("Need 자연스러움",gb.get("need_from_loss",0)),("Strategy 창의성",gb.get("strategy_creative",0)),("실패 대가",gb.get("failure_cost",0))]:
                    st.markdown(f'<div style="display:flex;align-items:center;margin:.2rem 0;font-size:.8rem"><div style="width:110px;color:var(--dim)">{nm}</div><div style="flex:1;background:#E0E0E8;border-radius:4px;height:8px;margin:0 .5rem"><div style="width:{sc*10}%;background:var(--y);height:100%;border-radius:4px"></div></div><div style="width:30px;text-align:right">{sc}</div></div>', unsafe_allow_html=True)
            if gb.get("feedback"):
                st.caption(gb["feedback"])

            st.markdown("#### 🚪 Gate C: Character Gate")
            c1, c2 = st.columns([1, 2])
            with c1:
                gc_avg = gc.get("average", 0)
                cl = "var(--g)" if gc_avg >= 7.0 else "var(--r)"
                lb = "PASS" if gc_avg >= 7.0 else "FAIL"
                st.markdown(f'<div style="text-align:center"><div class="big" style="color:{cl}">{gc_avg}</div><div class="sm" style="color:{cl};font-size:1rem;font-weight:700">{lb}</div></div>', unsafe_allow_html=True)
            with c2:
                for nm, sc in [("주인공/적대자 논리",gc.get("protagonist_antagonist_logic",0)),("조연 입체성",gc.get("supporting_not_functional",0)),("관계축 갈등",gc.get("relationship_produces_conflict",0))]:
                    st.markdown(f'<div style="display:flex;align-items:center;margin:.2rem 0;font-size:.8rem"><div style="width:110px;color:var(--dim)">{nm}</div><div style="flex:1;background:#E0E0E8;border-radius:4px;height:8px;margin:0 .5rem"><div style="width:{sc*10}%;background:var(--y);height:100%;border-radius:4px"></div></div><div style="width:30px;text-align:right">{sc}</div></div>', unsafe_allow_html=True)
            if gc.get("feedback"):
                st.caption(gc["feedback"])

            # 5축 스코어
            st.markdown("---")
            st.markdown("#### 📊 Development Fit Score")
            final = fa.get("final_score", 0)
            verdict = fa.get("verdict", "")
            vc = {"개발 진행":"var(--g)","개발 보류":"var(--y)","구조 재설계":"var(--r)"}.get(verdict,"var(--dim)")
            c1, c2 = st.columns([1, 2])
            with c1:
                st.markdown(f'<div style="text-align:center"><div class="big">{final}</div><div class="sm" style="color:{vc};font-size:1rem;font-weight:700">{verdict}</div></div>', unsafe_allow_html=True)
            with c2:
                for nm, sc in [("Goal",fa.get("goal",0)),("Need",fa.get("need",0)),("Strategy",fa.get("strategy",0)),("Structure",fa.get("structure",0)),("Character/Concept",fa.get("character_concept",0))]:
                    st.markdown(f'<div style="display:flex;align-items:center;margin:.2rem 0;font-size:.8rem"><div style="width:110px;color:var(--dim)">{nm}</div><div style="flex:1;background:#E0E0E8;border-radius:4px;height:8px;margin:0 .5rem"><div style="width:{sc*10}%;background:var(--y);height:100%;border-radius:4px"></div></div><div style="width:30px;text-align:right">{sc}</div></div>', unsafe_allow_html=True)

            if verdict == "개발 진행":
                st.success(f"✅ {verdict}. Character Bible 진행 가능.")
                if st.button("📖 Character Bible 진행 →", type="primary", use_container_width=True):
                    st.session_state.view = "char_bible"
                    st.rerun()
            elif verdict == "개발 보류":
                st.warning(f"⚠️ {verdict}. Core Build 보강 필요.")
                col_cv1, col_cv2 = st.columns(2)
                with col_cv1:
                    if st.button("🔓 Override → Character Bible"):
                        st.session_state.view = "char_bible"
                        st.rerun()
                with col_cv2:
                    if st.button("🔄 Core Build 재실행"):
                        project["core"] = None
                        project["core_gate"] = None
                        st.rerun()
            else:
                st.error(f"🔴 {verdict}. Brainstorm 재검토 필요.")

        else:
            st.warning("⚠️ Gate B + C 채점이 완료되지 않았습니다.")
            if st.button("🔄 Gate 재시도"):
                with st.spinner("Gate B + C 채점 중..."):
                    gate_result = call_core_gate(core)
                if gate_result:
                    project["core_gate"] = gate_result
                    project["final_score"] = gate_result.get("five_axis_scores", {}).get("final_score")
                    st.rerun()

    else:
        st.markdown(
            '<div style="text-align:center;padding:3rem 0;color:var(--dim)">'
            '🎯 Core Build를 실행하면 여기에 결과가 표시됩니다.<br>'
            'Logline Pack · Project Intent · World Build · Character Build'
            '</div>',
            unsafe_allow_html=True
        )

    st.markdown("---")
    st.caption(ENGINE_FOOTER)


# ═══════════════════════════════════════════════════
#  CHARACTER BIBLE
# ═══════════════════════════════════════════════════
elif st.session_state.view == "char_bible" and st.session_state.cur:

    project = st.session_state.projects[st.session_state.cur]
    core = project.get("core", {})

    st.markdown(f"## {project['title']}")
    st.caption(f"{project['genre']} · {project['target_market']} · {project['format']}")
    render_stepper("char_bible", project)

    lp = core.get("logline_pack", {})
    if lp.get("washed"):
        st.markdown(f'<div class="callout"><div class="cl">Logline</div>{lp["washed"]}</div>', unsafe_allow_html=True)

    st.markdown("---")
    st.markdown('<div class="section-header">📖 Character Bible <span class="en">CHARACTER DESIGN BIBLE</span></div>', unsafe_allow_html=True)
    st.caption("백스토리 · 비밀 · 말투 규칙 · 대사 샘플 · 관계 태도 · 변화 궤적 — Writer Engine이 일관된 인물을 쓰기 위한 설계서")

    if not project.get("char_bible"):
        if st.button("📖 Character Bible 생성", type="primary"):
            with st.spinner("캐릭터 바이블 설계 중... (캐릭터당 약 30초, 4~8명 기준 약 2~4분)"):
                result = call_character_bible(
                    core, project["genre"], project["format"],
                    locked_block=_build_project_locked_block(project),
                    fact_based=project.get("fact_based", False),
                    historical=project.get("historical", False),
                    film_type=project.get("film_type", ""),
                )
            if result:
                project["char_bible"] = result
                project["updated_at"] = datetime.now().strftime("%Y-%m-%d %H:%M")
                st.rerun()

    bible = project.get("char_bible", {})
    if bible:
        chars = bible.get("characters", [])
        role_labels = {"protagonist":"주인공","antagonist":"적대자","ally":"조력자","mirror":"거울","catalyst":"촉매자","subplot_lead":"서브플롯 리드"}
        role_emoji = {"protagonist":"🔥","antagonist":"🖤","ally":"💙","mirror":"🪞","catalyst":"⚡","subplot_lead":"🌐"}

        for ch in chars:
            role = role_labels.get(ch.get("role",""), ch.get("role",""))
            emoji = role_emoji.get(ch.get("role",""), "👤")
            name = ch.get("name", "")
            age = ch.get("age", "")

            st.markdown(f'<div class="section-header">{emoji} {role}: {name} ({age}) <span class="en">{ch.get("role","").upper()}</span></div>', unsafe_allow_html=True)

            # 외형 & 직업
            st.markdown(f'<div class="callout"><div class="cl">외형 · 첫인상</div>{ch.get("appearance","")}</div>', unsafe_allow_html=True)
            if ch.get("occupation"):
                st.markdown(f'<div class="ri"><div class="rl">직업/위치</div>{ch.get("occupation","")}</div>', unsafe_allow_html=True)

            # 욕망/결핍/결점
            c1, c2, c3 = st.columns(3)
            c1.markdown(f'<div class="callout"><div class="cl">GOAL</div>{ch.get("goal","")}</div>', unsafe_allow_html=True)
            c2.markdown(f'<div class="callout"><div class="cl">NEED</div>{ch.get("need","")}</div>', unsafe_allow_html=True)
            c3.markdown(f'<div class="callout"><div class="cl">FLAW</div>{ch.get("flaw","")}</div>', unsafe_allow_html=True)

            # 백스토리
            st.markdown(f'<div class="card"><div class="cl">📜 백스토리</div><p style="line-height:1.7">{ch.get("backstory","")}</p></div>', unsafe_allow_html=True)

            # 비밀 & 신념 & 두려움
            c4, c5, c6 = st.columns(3)
            c4.markdown(f'<div class="ri"><div class="rl">🔒 비밀</div>{ch.get("secret","")}</div>', unsafe_allow_html=True)
            c5.markdown(f'<div class="ri"><div class="rl">⚖️ 신념</div>{ch.get("belief","")}</div>', unsafe_allow_html=True)
            c6.markdown(f'<div class="ri"><div class="rl">😰 두려움</div>{ch.get("fear","")}</div>', unsafe_allow_html=True)

            # 습관
            habits = ch.get("habits", [])
            if habits:
                habits_html = " · ".join(habits) if isinstance(habits, list) else str(habits)
                st.markdown(f'<div class="ri"><div class="rl">🔄 반복 습관</div>{habits_html}</div>', unsafe_allow_html=True)

            # 말투 규칙
            sp = ch.get("speech_pattern", [])
            if sp:
                sp_list = sp if isinstance(sp, list) else [sp]
                sp_html = "<br>".join([f"• {s}" for s in sp_list])
                st.markdown(f'<div class="card"><div class="cl">🗣️ 말투 규칙 (SPEECH PATTERN)</div>{sp_html}</div>', unsafe_allow_html=True)

            # 대사 샘플
            sl = ch.get("sample_lines", {})
            if sl:
                sl_labels = {"normal":"평상시","angry":"분노","vulnerable":"취약"}
                sl_html = ""
                for k, v in sl.items():
                    label = sl_labels.get(k, k)
                    sl_html += f'<div style="margin:.3rem 0"><b style="color:var(--navy);font-size:.75rem">[{label}]</b> <i>\u2018{v}\u2019</i></div>'
                st.markdown(f'<div class="card"><div class="cl">💬 대사 샘플 (SAMPLE LINES)</div>{sl_html}</div>', unsafe_allow_html=True)

            # 관계별 태도
            ra = ch.get("relationship_attitudes", [])
            if ra:
                ra_list = ra if isinstance(ra, list) else [ra]
                ra_html = "<br>".join([f"• {r}" for r in ra_list])
                st.markdown(f'<div class="card"><div class="cl">🔗 관계별 태도</div>{ra_html}</div>', unsafe_allow_html=True)

            # 변화 궤적
            arc = ch.get("arc_detail", {})
            if arc:
                st.markdown(
                    f'<div class="card"><div class="cl">📈 변화 궤적 (ARC)</div>'
                    f'<div style="margin:.3rem 0"><b style="color:var(--navy);font-size:.75rem">[1막 끝]</b> {arc.get("act1_end","")}</div>'
                    f'<div style="margin:.3rem 0"><b style="color:var(--navy);font-size:.75rem">[미드포인트]</b> {arc.get("midpoint","")}</div>'
                    f'<div style="margin:.3rem 0"><b style="color:var(--navy);font-size:.75rem">[클라이맥스]</b> {arc.get("climax","")}</div>'
                    f'</div>',
                    unsafe_allow_html=True
                )

            st.markdown("")  # spacer

        # ── 캐릭터 관계도 ──
        st.markdown("---")
        st.markdown('<div class="section-header">🕸️ 캐릭터 관계도 <span class="en">CHARACTER MAP</span></div>', unsafe_allow_html=True)

        def render_relationship_map(chars):
            import math, html as _html

            role_colors = {
                "protagonist": "#FFCB05",
                "antagonist":  "#D32F2F",
                "ally":        "#2EC484",
                "mirror":      "#7B68EE",
                "catalyst":    "#FF8C00",
                "subplot_lead":"#20B2AA",
            }
            role_labels_kr = {
                "protagonist":"주인공","antagonist":"적대자","ally":"조력자",
                "mirror":"거울","catalyst":"촉매자","subplot_lead":"서브플롯",
            }

            # 관계 키워드 → 유형 매핑
            def classify_rel(text):
                t = text.lower()
                if any(k in t for k in ["죽이","살해","적","증오","혐오","원수","복수","배신","이용","조종","견제"]):
                    return "hostile"
                if any(k in t for k in ["사랑","좋아","연인","설레","끌","호감","연모","연인"]):
                    return "love"
                if any(k in t for k in ["돕","협력","신뢰","지지","동료","우정","친구","함께","의지"]):
                    return "ally"
                if any(k in t for k in ["질투","경쟁","라이벌","갈등","충돌","대립"]):
                    return "rival"
                if any(k in t for k in ["가르","멘토","스승","보호","이끌","안내"]):
                    return "mentor"
                return "neutral"

            rel_style = {
                "hostile": {"color":"#D32F2F","label":"⚔️ 적대","dash":"8,4"},
                "love":    {"color":"#FF69B4","label":"❤️ 사랑","dash":"0"},
                "ally":    {"color":"#2EC484","label":"🤝 협력","dash":"0"},
                "rival":   {"color":"#FF8C00","label":"🥊 경쟁","dash":"6,3"},
                "mentor":  {"color":"#7B68EE","label":"🎓 멘토","dash":"0"},
                "neutral": {"color":"#AAAAAA","label":"— 중립","dash":"4,4"},
            }

            n = len(chars)
            if n == 0:
                return

            # 원형 배치 좌표 계산
            cx, cy, r = 400, 300, 200
            positions = []
            for i in range(n):
                angle = (2 * math.pi * i / n) - math.pi / 2
                x = cx + r * math.cos(angle)
                y = cy + r * math.sin(angle)
                positions.append((x, y))

            # 관계 엣지 파싱
            edges = []
            for i, ch in enumerate(chars):
                ra = ch.get("relationship_attitudes", [])
                if isinstance(ra, list):
                    for rel_text in ra:
                        for j, other in enumerate(chars):
                            if i == j:
                                continue
                            other_name = other.get("name", "")
                            if other_name and other_name in rel_text:
                                rel_type = classify_rel(rel_text)
                                short = rel_text.replace(f"→ {other_name}:", "").replace(f"→{other_name}:", "").strip()
                                short = short[:30] + "…" if len(short) > 30 else short
                                edges.append((i, j, rel_type, short))

            # SVG 생성
            svg_lines = [
                f'<svg viewBox="0 0 800 600" xmlns="http://www.w3.org/2000/svg" '
                f'style="width:100%;max-width:800px;background:#F7F7F5;border-radius:16px;'
                f'border:1px solid #E2E2E0;font-family:Pretendard,sans-serif">',
                '<defs>',
            ]
            # 화살표 마커
            for rtype, rs in rel_style.items():
                svg_lines.append(
                    f'<marker id="arr_{rtype}" markerWidth="8" markerHeight="8" refX="6" refY="3" orient="auto">'
                    f'<path d="M0,0 L0,6 L8,3 z" fill="{rs["color"]}"/></marker>'
                )
            svg_lines.append('</defs>')

            # 엣지 선
            drawn = set()
            for (i, j, rtype, label) in edges:
                key = tuple(sorted([i, j]))
                rs = rel_style[rtype]
                x1, y1 = positions[i]
                x2, y2 = positions[j]
                # 살짝 오프셋 (양방향 선 구분)
                offset = 8 if (i, j) > (j, i) else -8
                mx = (x1 + x2) / 2 + offset * ((y2 - y1) / max(abs(x2 - x1) + abs(y2 - y1), 1))
                my = (y1 + y2) / 2 - offset * ((x2 - x1) / max(abs(x2 - x1) + abs(y2 - y1), 1))

                dash_attr = f'stroke-dasharray="{rs["dash"]}"' if rs["dash"] != "0" else ""
                svg_lines.append(
                    f'<line x1="{x1:.0f}" y1="{y1:.0f}" x2="{x2:.0f}" y2="{y2:.0f}" '
                    f'stroke="{rs["color"]}" stroke-width="2" {dash_attr} '
                    f'marker-end="url(#arr_{rtype})" opacity="0.7"/>'
                )
                # 관계 레이블
                escaped = _html.escape(label)
                svg_lines.append(
                    f'<text x="{mx:.0f}" y="{my:.0f}" text-anchor="middle" '
                    f'font-size="9" fill="{rs["color"]}" font-weight="600" '
                    f'style="paint-order:stroke" stroke="white" stroke-width="3">'
                    f'{escaped}</text>'
                )

            # 노드 원
            node_r = 44
            for i, ch in enumerate(chars):
                x, y = positions[i]
                role = ch.get("role", "neutral")
                color = role_colors.get(role, "#888888")
                name = _html.escape(ch.get("name", f"캐릭터{i+1}"))
                role_kr = role_labels_kr.get(role, role)

                # 외곽 링
                svg_lines.append(
                    f'<circle cx="{x:.0f}" cy="{y:.0f}" r="{node_r+4}" '
                    f'fill="{color}" opacity="0.15"/>'
                )
                # 메인 원
                svg_lines.append(
                    f'<circle cx="{x:.0f}" cy="{y:.0f}" r="{node_r}" '
                    f'fill="white" stroke="{color}" stroke-width="3"/>'
                )
                # 이름
                svg_lines.append(
                    f'<text x="{x:.0f}" y="{y-6:.0f}" text-anchor="middle" '
                    f'font-size="13" font-weight="900" fill="#191970">{name}</text>'
                )
                # 역할
                svg_lines.append(
                    f'<text x="{x:.0f}" y="{y+10:.0f}" text-anchor="middle" '
                    f'font-size="10" fill="{color}" font-weight="700">{role_kr}</text>'
                )

            svg_lines.append('</svg>')

            # 범례
            legend_items = []
            used_types = {e[2] for e in edges}
            for rtype in ["love","ally","mentor","rival","hostile","neutral"]:
                if rtype in used_types:
                    rs = rel_style[rtype]
                    legend_items.append(
                        f'<span style="display:inline-flex;align-items:center;gap:4px;margin-right:12px;">'
                        f'<svg width="24" height="6"><line x1="0" y1="3" x2="24" y2="3" '
                        f'stroke="{rs["color"]}" stroke-width="2.5" '
                        f'stroke-dasharray="{rs["dash"] if rs["dash"] != "0" else ""}"/></svg>'
                        f'<span style="font-size:0.78rem;color:#444">{rs["label"]}</span></span>'
                    )

            svg_html = "\n".join(svg_lines)
            legend_html = "".join(legend_items)

            st.markdown(
                f'<div style="text-align:center">{svg_html}</div>'
                f'<div style="text-align:center;margin-top:8px;padding:8px 0">{legend_html}</div>',
                unsafe_allow_html=True
            )

        render_relationship_map(chars)

        # Structure 진행 버튼
        st.markdown("---")
        st.success("✅ Character Bible 완료. Structure Build 진행 가능.")
        if st.button("🏗️ Structure Build 진행 →", type="primary", use_container_width=True):
            st.session_state.view = "structure"
            st.rerun()

    else:
        st.markdown(
            '<div style="text-align:center;padding:3rem 0;color:var(--dim)">'
            '📖 Character Bible을 생성하면 여기에 결과가 표시됩니다.<br>'
            '백스토리 · 비밀 · 말투 규칙 · 대사 샘플 · 관계 태도 · 변화 궤적'
            '</div>',
            unsafe_allow_html=True
        )

    st.markdown("---")
    st.caption(ENGINE_FOOTER)


# ═══════════════════════════════════════════════════
#  STRUCTURE BUILD
# ═══════════════════════════════════════════════════
elif st.session_state.view == "structure" and st.session_state.cur:

    project = st.session_state.projects[st.session_state.cur]
    core = project.get("core", {})

    st.markdown(f"## {project['title']}")
    st.caption(f"{project['genre']} · {project['target_market']} · {project['format']}")
    render_stepper("structure", project)

    gns = core.get("goal_need_strategy", {})
    lp = core.get("logline_pack", {})
    if lp.get("washed"):
        st.markdown(f'<div class="callout"><div class="cl">Logline</div>{lp["washed"]}</div>', unsafe_allow_html=True)
    if gns:
        c1, c2, c3 = st.columns(3)
        c1.markdown(f'<div class="callout"><div class="cl">GOAL</div>{gns.get("goal","")}</div>', unsafe_allow_html=True)
        c2.markdown(f'<div class="callout"><div class="cl">NEED</div>{gns.get("need","")}</div>', unsafe_allow_html=True)
        c3.markdown(f'<div class="callout"><div class="cl">STRATEGY</div>{gns.get("strategy","")}</div>', unsafe_allow_html=True)

    st.markdown("---")
    st.markdown('<div class="section-header">🏗️ Structure Build <span class="en">STORY STRUCTURE</span></div>', unsafe_allow_html=True)
    st.caption("시놉시스 · 스토리라인 · 3막 구조 · 15비트 · 캐릭터 변화표를 설계합니다.")

    if st.button("🏗️ Structure Build 실행", type="primary"):
        if not core:
            st.error("Core Build가 없습니다.")
        else:
            with st.spinner("① 시놉시스 + 스토리라인... (20~30초)"):
                story = call_structure_story(
                    core, project["genre"], project["target_market"], project["format"],
                    locked_block=_build_project_locked_block(project),
                    fact_based=project.get("fact_based", False),
                    historical=project.get("historical", False),
                    film_type=project.get("film_type", ""),
                )
            if story:
                project["structure_story"] = story
                with st.spinner("② 구조 진단 + 캐릭터 변화표... (20~30초)"):
                    diag = call_structure_diagnosis(
                        core, story, project["genre"], project["format"],
                        locked_block=_build_project_locked_block(project),
                        fact_based=project.get("fact_based", False),
                        historical=project.get("historical", False),
                        film_type=project.get("film_type", ""),
                    )
                if diag:
                    project["structure_diag"] = diag
                    with st.spinner("③ Gate D 채점... (10초)"):
                        gate_d = call_structure_gate(story, diag)
                    if gate_d:
                        project["structure_gate"] = gate_d
                with st.spinner("④ 기승전결 줄글 시놉시스... (10~15초)"):
                    prose = call_structure_prose(core, story)
                if prose:
                    project["structure_prose"] = prose
                project["updated_at"] = datetime.now().strftime("%Y-%m-%d %H:%M")
                st.rerun()

    if project.get("structure_story"):
        story = project["structure_story"]
        st.markdown("#### 📖 Synopsis 1P")
        syn = story.get("synopsis_1p", {})
        for label, key in [("시작","opening"),("촉발사건","catalyst"),("전개","development"),("미드포인트","midpoint"),("붕괴","collapse"),("결전","climax"),("결말","ending")]:
            val = syn.get(key, "")
            if val:
                st.markdown(f'<div class="callout"><div class="cl">{label}</div>{val}</div>', unsafe_allow_html=True)

        st.markdown("#### 📊 스토리라인 (시퀀스 방향)")
        for seq in story.get("storyline", []):
            st.markdown(f'<div class="card"><div class="cl">SEQ {seq.get("seq","")} · {seq.get("label","")} ({seq.get("pages","")})</div>{seq.get("summary","")}<br><span style="font-size:.8rem;color:#666">⚔️ {seq.get("conflict","")} · 💭 {seq.get("emotion","")} · → {seq.get("hook","")}</span></div>', unsafe_allow_html=True)

    if project.get("structure_diag"):
        diag = project["structure_diag"]
        st.markdown("#### 🎬 3막 구조 진단")
        ta = diag.get("three_act", {})
        for label, key in [("1막 끝","act1_end"),("미드포인트","act2_midpoint"),("All Is Lost","act2_end"),("클라이맥스","act3_climax")]:
            val = ta.get(key, "")
            if val:
                st.markdown(f'<div class="callout"><div class="cl">{label}</div>{val}</div>', unsafe_allow_html=True)
        if ta.get("diagnosis"):
            st.caption(f"진단: {ta['diagnosis']}")

        st.markdown("#### 🥁 15-Beat Sheet")
        for bt in diag.get("beat_sheet", []):
            status = bt.get("status", "")
            color = {"있음":"var(--g)","약함":"var(--y)","없음":"var(--r)"}.get(status, "var(--dim)")
            st.markdown(
                f'<div style="margin:.3rem 0;font-size:.85rem">'
                f'<span style="color:{color};font-weight:700">[{status}]</span> '
                f'<b>{bt.get("beat","")}</b> — '
                f'<span style="color:#666">{bt.get("note","")}</span>'
                f'</div>',
                unsafe_allow_html=True
            )

        # 캐릭터 변화표
        arcs = diag.get("character_arcs", [])
        if arcs:
            st.markdown("#### 🎭 캐릭터 변화표")
            role_labels = {"protagonist":"주인공","antagonist":"적대자","ally":"조력자","mirror":"거울","catalyst":"촉매자","subplot_lead":"서브플롯 리드"}
            for arc in arcs:
                role = role_labels.get(arc.get("role",""), arc.get("role",""))
                arc_type = arc.get("arc_type", "")
                st.markdown(
                    f'<div class="card">'
                    f'<div class="cl">{role}: {arc.get("name","")} [{arc_type}]</div>'
                    f'<span style="font-size:.8rem;color:#666">'
                    f'1막: {arc.get("act1_state","")}<br>'
                    f'전환: {arc.get("turning_point","")}<br>'
                    f'3막: {arc.get("act3_state","")}'
                    f'</span></div>',
                    unsafe_allow_html=True
                )

        # 관계 변화표
        rels = diag.get("relationship_changes", [])
        if rels:
            st.markdown("#### 🔗 관계 변화표")
            for rel in rels:
                st.markdown(
                    f'<div class="card">'
                    f'<div class="cl">{rel.get("pair","")}</div>'
                    f'<span style="font-size:.8rem;color:#666">'
                    f'1막: {rel.get("act1","")} → '
                    f'미드: {rel.get("midpoint","")} → '
                    f'3막: {rel.get("act3","")}'
                    f'</span></div>',
                    unsafe_allow_html=True
                )

    # 기승전결 줄글 시놉시스
    if project.get("structure_prose"):
        prose_data = project["structure_prose"]
        if prose_data.get("prose"):
            st.markdown("#### 📝 기승전결 시놉시스 (1P 줄글)")
            st.markdown(f'<div class="callout" style="line-height:1.8;font-size:.9rem">{prose_data["prose"]}</div>', unsafe_allow_html=True)

    if project.get("structure_gate"):
        sg = project["structure_gate"]
        gd = sg.get("gate_d_structure", {})
        gd_avg = gd.get("average", 0)
        gd_ok = gd_avg >= 7.0
        st.markdown("---")
        st.markdown("#### 🚪 Gate D: Structure Gate")
        c1, c2 = st.columns([1, 2])
        with c1:
            cl = "var(--g)" if gd_ok else "var(--r)"
            st.markdown(f'<div style="text-align:center"><div class="big" style="color:{cl}">{gd_avg}</div><div class="sm" style="color:{cl};font-size:1rem;font-weight:700">{"PASS" if gd_ok else "FAIL"}</div></div>', unsafe_allow_html=True)
        with c2:
            for nm, sc in [("전환점",gd.get("turning_points_valid",0)),("Midpoint",gd.get("midpoint_redirects",0)),("All Is Lost",gd.get("all_is_lost_works",0)),("결말",gd.get("ending_inevitable_surprising",0))]:
                st.markdown(f'<div style="display:flex;align-items:center;margin:.2rem 0;font-size:.8rem"><div style="width:80px;color:var(--dim)">{nm}</div><div style="flex:1;background:#E0E0E8;border-radius:4px;height:8px;margin:0 .5rem"><div style="width:{sc*10}%;background:var(--y);height:100%;border-radius:4px"></div></div><div style="width:30px;text-align:right">{sc}</div></div>', unsafe_allow_html=True)
        if gd.get("feedback"):
            st.caption(gd["feedback"])

        if gd_ok:
            st.success("✅ Gate D 통과. Scene Design 진행 가능.")
            if st.button("🎬 Scene Design 진행 →", type="primary", use_container_width=True):
                st.session_state.view = "scene_design"
                st.rerun()
        else:
            st.warning(f"⚠️ Gate D 미통과 (평균 {gd_avg}).")
            col_sd1, col_sd2 = st.columns(2)
            with col_sd1:
                if st.button("🔓 Override → Scene Design"):
                    st.session_state.view = "scene_design"
                    st.rerun()
            with col_sd2:
                if st.button("🔄 Structure 재실행"):
                    for k in ["structure_story","structure_diag","structure_gate","structure_prose"]:
                        project[k] = None
                    st.rerun()

        # DOCX 1차 다운로드
        st.markdown("---")
        st.markdown("#### 📥 기획개발보고서 [1차] 다운로드")
        st.caption("Core Build + Structure Build까지의 기획서입니다. 다운로드 후 직접 수정/보강하세요.")
        st.markdown(
            '<div class="callout">'
            '<div class="cl">💡 1차 기획서 활용법</div>'
            '다운로드한 기획서를 검토하고, 추가하고 싶은 요소(예: 사이렌, 특정 장치, 캐릭터 보강 등)를 '
            '직접 수정하세요. 수정이 끝나면 Scene Design → Treatment로 진행하여 최종 기획서를 완성합니다.'
            '</div>',
            unsafe_allow_html=True
        )
        title_safe = project.get("title", "프로젝트").replace(" ", "_")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M")
        docx_buffer = generate_docx(project)
        st.download_button(
            label="📥 기획개발보고서 [1차] 다운로드 (.docx)",
            data=docx_buffer,
            file_name=f"기획개발보고서_{title_safe}_1차_Blue_{timestamp}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

    elif not project.get("structure_story"):
        st.markdown('<div style="text-align:center;padding:3rem 0;color:var(--dim)">🏗️ Structure Build를 실행하면 여기에 결과가 표시됩니다.</div>', unsafe_allow_html=True)

    st.markdown("---")
    st.caption(ENGINE_FOOTER)


# ═══════════════════════════════════════════════════
#  SCENE DESIGN (장면화)
# ═══════════════════════════════════════════════════
elif st.session_state.view == "scene_design" and st.session_state.cur:

    project = st.session_state.projects[st.session_state.cur]
    core = project.get("core", {})
    story = project.get("structure_story", {})
    diag = project.get("structure_diag", {})

    st.markdown(f"## {project['title']}")
    st.caption(f"{project['genre']} · {project['target_market']} · {project['format']}")
    render_stepper("scene_design", project)

    gns = core.get("goal_need_strategy", {})
    lp = core.get("logline_pack", {})
    if lp.get("washed"):
        st.markdown(f'<div class="callout"><div class="cl">Logline</div>{lp["washed"]}</div>', unsafe_allow_html=True)

    st.markdown("---")
    st.markdown('<div class="section-header">🎬 Scene Design <span class="en">SCENE DESIGN</span></div>', unsafe_allow_html=True)
    st.caption("Show, don't tell — 핵심 장면의 극적 행동 · 반전 · 시각 연출을 설계합니다.")

    if st.button("🎬 Scene Design 실행", type="primary"):
        if not story:
            st.error("Structure Build가 없습니다.")
        else:
            with st.spinner("핵심 장면 설계 중... (약 30~40초)"):
                sd = call_scene_design(
                    core, story, diag, project["genre"], project["format"],
                    locked_block=_build_project_locked_block(project),
                    fact_based=project.get("fact_based", False),
                    historical=project.get("historical", False),
                    film_type=project.get("film_type", ""),
                )
            if sd:
                project["scene_design"] = sd
                project["updated_at"] = datetime.now().strftime("%Y-%m-%d %H:%M")
                st.rerun()

    # ── Scene Design 결과 표시 ──
    if project.get("scene_design"):
        sd = project["scene_design"]
        scenes = sd.get("key_scenes", [])
        sms = sd.get("scene_map_summary", {})

        # 장면 맵 요약
        if sms:
            st.markdown("#### 🗺️ Scene Map")
            c1, c2, c3, c4 = st.columns(4)
            c1.markdown(f'<div class="callout"><div class="cl">1막</div>{sms.get("act1_scenes","")}</div>', unsafe_allow_html=True)
            c2.markdown(f'<div class="callout"><div class="cl">2막 전반</div>{sms.get("act2a_scenes","")}</div>', unsafe_allow_html=True)
            c3.markdown(f'<div class="callout"><div class="cl">2막 후반</div>{sms.get("act2b_scenes","")}</div>', unsafe_allow_html=True)
            c4.markdown(f'<div class="callout"><div class="cl">3막</div>{sms.get("act3_scenes","")}</div>', unsafe_allow_html=True)

            if sms.get("must_see_scenes"):
                st.markdown(f'<div class="callout" style="border-left-color:var(--y)"><div class="cl">⭐ Must-See 장면</div>{sms["must_see_scenes"]}</div>', unsafe_allow_html=True)

        # 핵심 장면 카드
        st.markdown(f"#### 🎬 핵심 장면 ({len(scenes)}개)")
        for sc in scenes:
            tp = sc.get("turning_point", "")
            tp_html = f'<br>🔄 <b>전환:</b> {tp}' if tp else ""
            di = sc.get("dramatic_irony", "")
            di_html = f'<br>👁️ <b>아이러니:</b> <i>{di}</i>' if di else ""
            kl = sc.get("key_line", "")
            kl_html = f'<br><span style="color:var(--navy);font-weight:600">💬 "{kl}"</span>' if kl else ""
            st.markdown(
                f'<div class="card">'
                f'<div class="cl">S#{sc.get("scene_no","")} · {sc.get("sequence","")} · {sc.get("location","")}</div>'
                f'<b>{sc.get("title","")}</b> — {sc.get("characters","")}<br>'
                f'<span style="font-size:.85rem">'
                f'📍 {sc.get("setup","")}<br>'
                f'🎭 <b>행동:</b> {sc.get("dramatic_action","")}'
                f'{tp_html}'
                f'{di_html}<br>'
                f'💭 {sc.get("emotion_shift","")}<br>'
                f'🎥 {sc.get("visual_direction","")}<br>'
                f'⚡ 판돈: {sc.get("stakes","")}'
                f'{kl_html}<br>'
                f'→ {sc.get("connection","")}'
                f'</span></div>',
                unsafe_allow_html=True
            )

        st.markdown("---")

        # JSON 프로젝트 저장 (세션 끊김 대비)
        st.markdown("#### 💾 프로젝트 JSON 저장 (세션 복구용)")
        st.caption(
            "브라우저를 닫거나 세션이 끊겨도 여기서 JSON 파일을 저장하면 "
            "다음에 홈 화면의 '프로젝트 불러오기'로 복원 후 Treatment부터 이어서 진행할 수 있습니다."
        )
        title_safe_json = project.get("title", "프로젝트").replace(" ", "_")
        ts_json = datetime.now().strftime("%Y%m%d_%H%M")
        project_json_str = save_project_to_json(project, engine_version="v2.3.5")
        st.download_button(
            label="💾 프로젝트 저장 (Scene 완료 시점)",
            data=project_json_str.encode("utf-8"),
            file_name=f"{title_safe_json}_Scene완료_{ts_json}.json",
            mime="application/json",
            use_container_width=True,
        )

        st.markdown("---")

        # Treatment 진행
        st.success("✅ Scene Design 완료. Treatment Build 진행 가능.")
        if st.button("📝 Treatment Build 진행 →", type="primary", use_container_width=True):
            st.session_state.view = "treatment"
            st.rerun()

    else:
        st.markdown(
            '<div style="text-align:center;padding:3rem 0;color:var(--dim)">'
            '🎬 Scene Design을 실행하면 여기에 핵심 장면이 표시됩니다.<br>'
            'Show, don\'t tell — 행동 · 반전 · 시각 연출'
            '</div>',
            unsafe_allow_html=True
        )

    st.markdown("---")
    st.caption(ENGINE_FOOTER)


# ═══════════════════════════════════════════════════
#  TREATMENT BUILD
# ═══════════════════════════════════════════════════
elif st.session_state.view == "treatment" and st.session_state.cur:

    project = st.session_state.projects[st.session_state.cur]
    core = project.get("core", {})
    story = project.get("structure_story", {})
    diag = project.get("structure_diag", {})

    st.markdown(f"## {project['title']}")
    st.caption(f"{project['genre']} · {project['target_market']} · {project['format']}")
    render_stepper("treatment", project)

    # 요약 정보
    gns = core.get("goal_need_strategy", {})
    lp = core.get("logline_pack", {})
    if lp.get("washed"):
        st.markdown(f'<div class="callout"><div class="cl">Logline</div>{lp["washed"]}</div>', unsafe_allow_html=True)

    st.markdown("---")
    st.markdown('<div class="section-header">📝 Treatment Build <span class="en">16-BEAT TREATMENT</span></div>', unsafe_allow_html=True)
    st.caption("16비트 구조 × 줄글 트리트먼트. 1막(6비트) + 2막(6비트) + 3막(4비트) = 약 40~50페이지")

    scene_data = project.get("scene_design", {})

    if st.button("📝 Treatment Build 실행", type="primary"):
        if not story:
            st.error("Structure Build가 없습니다.")
        else:
            with st.spinner("① 1막 Treatment (Beat 1~6)... (약 40~50초)"):
                act1 = call_treatment_beats(
                    core, story, scene_data, project["genre"], project["format"], 1,
                    locked_block=_build_project_locked_block(project),
                    fact_based=project.get("fact_based", False),
                    historical=project.get("historical", False),
                    film_type=project.get("film_type", ""),
                    research=project.get("research"),
                )
            with st.spinner("② 2막 Treatment (Beat 7~12)... (약 40~50초)"):
                act2 = call_treatment_beats(
                    core, story, scene_data, project["genre"], project["format"], 2,
                    locked_block=_build_project_locked_block(project),
                    fact_based=project.get("fact_based", False),
                    historical=project.get("historical", False),
                    film_type=project.get("film_type", ""),
                    research=project.get("research"),
                )
            with st.spinner("③ 3막 Treatment (Beat 13~16)... (약 30~40초)"):
                act3 = call_treatment_beats(
                    core, story, scene_data, project["genre"], project["format"], 3,
                    locked_block=_build_project_locked_block(project),
                    fact_based=project.get("fact_based", False),
                    historical=project.get("historical", False),
                    film_type=project.get("film_type", ""),
                    research=project.get("research"),
                )

            if act1 or act2 or act3:
                project["treatment"] = {"act1": act1, "act2": act2, "act3": act3}
                project["updated_at"] = datetime.now().strftime("%Y-%m-%d %H:%M")

                with st.spinner("④ 감정 곡선 + 감독 포인트... (약 10초)"):
                    meta = call_treatment_meta(act1, act2, act3, core)
                if meta:
                    project["treatment"]["meta"] = meta

                with st.spinner("⑤ Gate E 채점... (약 15초, 장르·엔딩·논리 검증 포함)"):
                    gate_e = call_treatment_gate(
                        project["treatment"],
                        core_data=core,
                        genre=project["genre"],
                    )
                if gate_e:
                    project["treatment_gate"] = gate_e

                project["updated_at"] = datetime.now().strftime("%Y-%m-%d %H:%M")
                st.rerun()

    # ── Treatment 결과 표시 (16비트 줄글) ──
    if project.get("treatment"):
        treat = project["treatment"]

        act_labels = {1: "1막 — 설정 (Set-up)", 2: "2막 — 대결 (Confrontation)", 3: "3막 — 해결 (Resolution)"}

        for act_num in [1, 2, 3]:
            act_key = f"act{act_num}"
            act_data = treat.get(act_key)
            if act_data:
                st.markdown(f"#### 📖 {act_labels[act_num]}")

                for b in act_data.get("beats", []):
                    beat_no = b.get("beat_no", "")
                    beat_name = b.get("beat_name", "")
                    episode = b.get("episode", "")
                    narrative = b.get("narrative", "").replace("\n", "<br>")
                    char_count = len(b.get("narrative", ""))
                    event_s = b.get("event_summary", "")
                    decision_s = b.get("decision_summary", "")
                    consequence_s = b.get("consequence_summary", "")
                    status_c = b.get("status_change", "")
                    b_story = b.get("b_story_beat", "")
                    cliff = b.get("cliffhanger", "")
                    villain = b.get("villain_beat", "")
                    pp = b.get("plant_payoff", "")

                    ep_tag = f'<span style="background:var(--y);color:var(--n);padding:1px 6px;border-radius:3px;font-size:.7rem;margin-right:6px">{episode}</span>' if episode else ""

                    meta_lines = []
                    if event_s:
                        meta_lines.append(f'<b>사건</b>: {event_s}')
                    if decision_s:
                        meta_lines.append(f'<b>선택</b>: {decision_s}')
                    if consequence_s:
                        meta_lines.append(f'<b>결과</b>: {consequence_s}')
                    if status_c:
                        meta_lines.append(f'<b>변화</b>: {status_c}')
                    if b_story:
                        meta_lines.append(f'<b>B-Story</b>: {b_story}')
                    if villain:
                        meta_lines.append(f'<b>빌런</b>: {villain}')
                    if pp:
                        meta_lines.append(f'<b>🌱 Plant/Payoff</b>: {pp}')
                    if cliff:
                        meta_lines.append(f'<b style="color:var(--r)">CLIFFHANGER</b>: {cliff}')
                    meta_html = "<br>".join(meta_lines)
                    meta_block = f'<div style="background:var(--mist);padding:8px 10px;border-radius:6px;margin-top:8px;font-size:.78rem;line-height:1.6">{meta_html}</div>' if meta_lines else ""

                    st.markdown(
                        f'<div class="card">'
                        f'<div class="cl">{ep_tag}Beat {beat_no}. {beat_name}</div>'
                        f'<div style="line-height:2.0;font-size:.9rem;margin-top:.5rem">'
                        f'{narrative}'
                        f'</div>'
                        f'{meta_block}'
                        f'<div style="text-align:right;font-size:.65rem;color:var(--dim);margin-top:.3rem">{char_count}자</div>'
                        f'</div>',
                        unsafe_allow_html=True
                    )

                st.markdown("")

        # 감정 곡선
        meta = treat.get("meta", {})
        ec = meta.get("emotion_curve", [])
        if ec:
            st.markdown("#### 📈 감정 곡선")
            import plotly.graph_objects as go
            fig = go.Figure()
            fig.add_trace(go.Scatter(
                x=[p.get("point","") for p in ec],
                y=[p.get("tension",0) for p in ec],
                mode='lines+markers',
                line=dict(color='#191970', width=3),
                marker=dict(size=8, color='#FFCB05', line=dict(color='#191970', width=2)),
                text=[p.get("emotion","") for p in ec],
                hovertemplate='%{x}<br>텐션: %{y}<br>감정: %{text}<extra></extra>'
            ))
            fig.update_layout(
                plot_bgcolor='#FAFAFA', paper_bgcolor='#FAFAFA',
                font_color='#1A1A2E', yaxis_range=[0,10],
                yaxis_title="Tension", height=300,
                margin=dict(l=40,r=20,t=20,b=40),
                yaxis=dict(gridcolor='#E0E0E8'),
                xaxis=dict(gridcolor='#E0E0E8')
            )
            st.plotly_chart(fig, use_container_width=True)

        # 감독용 포인트
        notes = meta.get("director_notes", [])
        if notes:
            st.markdown("#### 🎬 감독용 포인트")
            for i, n in enumerate(notes, 1):
                st.markdown(f"**{i}.** {n}")

        # 투자자 요약
        inv = meta.get("investor_summary", "")
        if inv:
            st.markdown(f'<div class="callout"><div class="cl">💰 투자자용 요약</div>{inv}</div>', unsafe_allow_html=True)

        st.markdown("---")

        # Gate E
        if project.get("treatment_gate"):
            tg = project["treatment_gate"]
            ge = tg.get("gate_e_treatment", {})
            ge_avg = ge.get("average", 0)
            ge_ok = ge_avg >= 7.0

            st.markdown("#### 🚪 Gate E: Treatment Gate")
            c1, c2 = st.columns([1, 2])
            with c1:
                cl = "var(--g)" if ge_ok else "var(--r)"
                st.markdown(f'<div style="text-align:center"><div class="big" style="color:{cl}">{ge_avg}</div><div class="sm" style="color:{cl};font-size:1rem;font-weight:700">{"PASS" if ge_ok else "FAIL"}</div></div>', unsafe_allow_html=True)
            with c2:
                for nm, sc in [("영화적 읽힘",ge.get("cinematic_reading",0)),("씬-감정 일치",ge.get("scene_emotion_match",0)),("비트 충실도",ge.get("beat_completeness",0)),("초고 직행 가능",ge.get("screenplay_ready",0))]:
                    st.markdown(f'<div style="display:flex;align-items:center;margin:.2rem 0;font-size:.8rem"><div style="width:100px;color:var(--dim)">{nm}</div><div style="flex:1;background:#E0E0E8;border-radius:4px;height:8px;margin:0 .5rem"><div style="width:{sc*10}%;background:var(--y);height:100%;border-radius:4px"></div></div><div style="width:30px;text-align:right">{sc}</div></div>', unsafe_allow_html=True)
            if ge.get("feedback"):
                st.caption(ge["feedback"])

            if ge_ok:
                st.success("✅ Gate E 통과. 기획개발 완료. Screenplay Writer Engine 연동 준비 완료.")
            else:
                st.warning(f"⚠️ Gate E 미통과 (평균 {ge_avg}). Treatment 보강 필요.")
                if st.button("🔄 Treatment 재실행"):
                    project["treatment"] = None
                    project["treatment_gate"] = None
                    st.rerun()

        # DOCX 2차(최종) 다운로드
        st.markdown("---")
        st.markdown("#### 📥 기획개발보고서 [최종] 다운로드")
        st.caption("Core + Structure + Scene Design + Treatment 전체가 포함된 최종 기획서입니다.")
        title_safe = project.get("title", "프로젝트").replace(" ", "_")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M")
        docx_buffer = generate_docx(project)
        st.download_button(
            label="📥 기획개발보고서 [최종] 다운로드 (.docx)",
            data=docx_buffer,
            file_name=f"기획개발보고서_{title_safe}_최종_Blue_{timestamp}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

        # JSON 프로젝트 저장 (세션 복구용)
        st.markdown("#### 💾 프로젝트 JSON 저장 (세션 복구용)")
        project_json_str = save_project_to_json(project, engine_version="v2.3.5")
        st.download_button(
            label="💾 프로젝트 저장 (Treatment 완료 시점)",
            data=project_json_str.encode("utf-8"),
            file_name=f"{title_safe}_Treatment완료_{timestamp}.json",
            mime="application/json",
            use_container_width=True,
        )

        # Tone Document 진행
        st.markdown("---")
        st.success("✅ Treatment 완료. Tone Document 진행 가능.")
        if st.button("🎨 Tone Document 진행 →", type="primary", use_container_width=True):
            st.session_state.view = "tone_doc"
            st.rerun()

    else:
        st.markdown(
            '<div style="text-align:center;padding:3rem 0;color:var(--dim)">'
            '📝 Treatment Build를 실행하면 여기에 결과가 표시됩니다.'
            '</div>',
            unsafe_allow_html=True
        )

    st.markdown("---")
    st.caption(ENGINE_FOOTER)


# ═══════════════════════════════════════════════════
#  TONE DOCUMENT
# ═══════════════════════════════════════════════════
elif st.session_state.view == "tone_doc" and st.session_state.cur:

    project = st.session_state.projects[st.session_state.cur]
    core = project.get("core", {})
    treatment = project.get("treatment", {})

    st.markdown(f"## {project['title']}")
    st.caption(f"{project['genre']} · {project['target_market']} · {project['format']}")
    render_stepper("tone_doc", project)

    lp = core.get("logline_pack", {})
    if lp.get("washed"):
        st.markdown(f'<div class="callout"><div class="cl">Logline</div>{lp["washed"]}</div>', unsafe_allow_html=True)

    st.markdown("---")
    st.markdown('<div class="section-header">🎨 Tone Document <span class="en">VISUAL & TONAL GUIDE</span></div>', unsafe_allow_html=True)
    st.caption("카메라 · 색감 · 페이싱 · 대사 규칙 · 모티프 · 금기 — Writer Engine의 톤 일관성 가이드")

    if not project.get("tone_doc"):
        if st.button("🎨 Tone Document 생성", type="primary"):
            with st.spinner("톤 & 연출 문서 설계 중... (최대 40초)"):
                result = call_tone_document(
                    core, project.get("structure_story", {}),
                    project.get("scene_design", {}), treatment,
                    project.get("char_bible", {}),
                    project["genre"], project["format"],
                    locked_block=_build_project_locked_block(project),
                    fact_based=project.get("fact_based", False),
                    historical=project.get("historical", False),
                    film_type=project.get("film_type", ""),
                )
            if result:
                project["tone_doc"] = result
                project["updated_at"] = datetime.now().strftime("%Y-%m-%d %H:%M")
                st.rerun()

    td = project.get("tone_doc", {})
    if td:
        # Visual Style
        vs = td.get("visual_style", {})
        if vs:
            st.markdown('<div class="section-header">🎥 비주얼 스타일 <span class="en">VISUAL STYLE</span></div>', unsafe_allow_html=True)
            for label, key in [("카메라 철학","camera_philosophy"),("색감 팔레트","color_palette"),("조명 규칙","lighting_rule"),("시그니처 쇼트","signature_shot")]:
                val = vs.get(key, "")
                if val:
                    st.markdown(f'<div class="ri"><div class="rl">{label}</div>{val}</div>', unsafe_allow_html=True)

        # Pacing
        pc = td.get("pacing", {})
        if pc:
            st.markdown('<div class="section-header">⏱️ 페이싱 <span class="en">PACING</span></div>', unsafe_allow_html=True)
            for label, key in [("전체 철학","overall"),("1막 템포","act1_tempo"),("2막 템포","act2_tempo"),("3막 템포","act3_tempo"),("대사 비율","dialogue_density")]:
                val = pc.get(key, "")
                if val:
                    st.markdown(f'<div class="ri"><div class="rl">{label}</div>{val}</div>', unsafe_allow_html=True)

        # Dialogue Rules
        dr = td.get("dialogue_rules", {})
        if dr:
            st.markdown('<div class="section-header">💬 대사 규칙 <span class="en">DIALOGUE RULES</span></div>', unsafe_allow_html=True)
            for label, key in [("전체 톤","overall_tone"),("서브텍스트","subtext_rule"),("침묵 활용","silence_usage")]:
                val = dr.get(key, "")
                if val:
                    st.markdown(f'<div class="ri"><div class="rl">{label}</div>{val}</div>', unsafe_allow_html=True)
            forbidden = dr.get("forbidden_phrases", [])
            if forbidden:
                fp_list = forbidden if isinstance(forbidden, list) else [forbidden]
                fp_html = "<br>".join([f"🚫 {f}" for f in fp_list])
                st.markdown(f'<div class="card"><div class="cl">금지 대사 패턴</div>{fp_html}</div>', unsafe_allow_html=True)

        # Motifs
        mt = td.get("motifs", {})
        if mt:
            st.markdown('<div class="section-header">🔄 모티프 <span class="en">RECURRING MOTIFS</span></div>', unsafe_allow_html=True)
            for label, key in [("반복 소품/모티프","recurring_objects"),("반복 장소","recurring_locations")]:
                val = mt.get(key, [])
                if val:
                    items = val if isinstance(val, list) else [val]
                    items_html = "<br>".join([f"• {i}" for i in items])
                    st.markdown(f'<div class="card"><div class="cl">{label}</div>{items_html}</div>', unsafe_allow_html=True)
            wm = mt.get("weather_mood", "")
            if wm:
                st.markdown(f'<div class="ri"><div class="rl">날씨/계절</div>{wm}</div>', unsafe_allow_html=True)

        # Music & Sound
        ms = td.get("music_sound", {})
        if ms:
            st.markdown('<div class="section-header">🎵 사운드 <span class="en">MUSIC & SOUND</span></div>', unsafe_allow_html=True)
            for label, key in [("음악 방향","score_direction"),("무음 활용","silence_scenes")]:
                val = ms.get(key, "")
                if val:
                    st.markdown(f'<div class="ri"><div class="rl">{label}</div>{val}</div>', unsafe_allow_html=True)
            ds = ms.get("diegetic_sounds", [])
            if ds:
                items = ds if isinstance(ds, list) else [ds]
                items_html = " · ".join(items)
                st.markdown(f'<div class="ri"><div class="rl">작품 내 소리</div>{items_html}</div>', unsafe_allow_html=True)

        # Forbidden
        forbidden = td.get("forbidden", [])
        if forbidden:
            st.markdown('<div class="section-header">🚫 금기 <span class="en">FORBIDDEN</span></div>', unsafe_allow_html=True)
            fb_list = forbidden if isinstance(forbidden, list) else [forbidden]
            fb_html = "<br>".join([f"🚫 {f}" for f in fb_list])
            st.markdown(f'<div class="card">{fb_html}</div>', unsafe_allow_html=True)

        # Reference Films
        refs = td.get("reference_films", [])
        if refs:
            st.markdown('<div class="section-header">🎬 참고 작품 <span class="en">REFERENCE FILMS</span></div>', unsafe_allow_html=True)
            for ref in refs:
                title_ref = ref.get("title", "") if isinstance(ref, dict) else str(ref)
                reason = ref.get("reason", "") if isinstance(ref, dict) else ""
                st.markdown(f'<div class="ri"><div class="rl">{title_ref}</div>{reason}</div>', unsafe_allow_html=True)

        # Writer Instruction
        wi = td.get("writer_instruction", "")
        if wi:
            st.markdown(f'<div class="callout" style="border-left-color:var(--y)"><div class="cl">✍️ Writer Engine 최종 지시</div>{wi}</div>', unsafe_allow_html=True)

        # 완료
        st.markdown("---")
        st.success("🎉 기획개발 완료! 전체 9단계 파이프라인이 완성되었습니다.")
        st.info("→ Writer Engine에서 이 기획서를 불러와 시나리오를 생성할 수 있습니다.")

        # 최종 DOCX 다운로드
        title_safe = project.get("title", "프로젝트").replace(" ", "_")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M")
        docx_buffer = generate_docx(project)
        st.download_button(
            label="📥 기획개발보고서 [최종 + Tone] 다운로드 (.docx)",
            data=docx_buffer,
            file_name=f"기획개발보고서_{title_safe}_최종_Blue_{timestamp}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

        # 최종 프로젝트 JSON 저장 (완성 상태 백업)
        st.markdown("#### 💾 프로젝트 JSON 저장 (완성 상태 백업)")
        st.caption("완성된 프로젝트 전체 상태. 나중에 Writer Engine 투입 전 참고용 또는 백업용.")
        project_json_str = save_project_to_json(project, engine_version="v2.3.5")
        st.download_button(
            label="💾 프로젝트 저장 (최종 완성)",
            data=project_json_str.encode("utf-8"),
            file_name=f"{title_safe}_최종완성_{timestamp}.json",
            mime="application/json",
            use_container_width=True,
        )

    else:
        st.markdown(
            '<div style="text-align:center;padding:3rem 0;color:var(--dim)">'
            '🎨 Tone Document를 생성하면 여기에 결과가 표시됩니다.<br>'
            '카메라 · 색감 · 페이싱 · 대사 규칙 · 모티프 · 참고 작품'
            '</div>',
            unsafe_allow_html=True
        )

    st.markdown("---")
    st.caption(ENGINE_FOOTER)
